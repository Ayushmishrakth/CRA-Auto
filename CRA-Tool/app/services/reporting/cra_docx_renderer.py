from __future__ import annotations

import json
import re
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any

from app.services.reporting.cra_pdf_renderer import (
    SEVERITY_RANK,
    SERVICE_NAMES,
    _business_risk,
    _canonical_service,
    _coverage_percent,
    _executive_summary,
    _fail_percent,
    _microsoft_reference,
    _licensing_table,
    _normalized_status,
    _pillar_readiness,
    _priority,
    _readiness_level,
    _risk_counts,
    _risk_rating,
    _service_summary,
    _status_counts,
    _status_label,
)


def render_docx(path: Path, report: dict[str, Any]) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ModuleNotFoundError as exc:
        raise RuntimeError("python-docx is required for CRA executive DOCX reports.") from exc

    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    summary = report["summary"]
    narrative = report.get("narrative") or {}
    rows = _template_aligned_rows(report.get("parameter_rows") or [])
    summary = _summary_for_rows(summary, rows)
    chart_dir = path.parent / f"{path.stem}_charts"
    chart_dir.mkdir(parents=True, exist_ok=True)

    _cover_page(doc, summary, WD_ALIGN_PARAGRAPH)
    doc.add_page_break()
    _toc(doc)
    doc.add_page_break()

    _heading(doc, "1. Executive Summary")
    for paragraph in _executive_summary(summary, narrative, rows):
        doc.add_paragraph(paragraph)
    _metrics_table(doc, [
        ("Overall readiness", f"{summary.get('overall_readiness', 0)}%"),
        ("Readiness status", summary.get("readiness_status", "-")),
        ("Total parameters", summary.get("parameter_total", len(rows))),
        ("Passed", summary.get("pass_total", 0)),
        ("Failed", summary.get("fail_total", 0)),
        ("Licensing required", summary.get("licensing_required_total", 0)),
        ("Manual validation", summary.get("manual_validation_total", 0)),
        ("Coverage", _coverage_percent(summary)),
    ])

    _heading(doc, "2. Purpose")
    for text in [
        "Assess security readiness for Microsoft 365 Copilot by identifying identity, access, collaboration, and data exposure risks.",
        "Assess governance readiness by reviewing control ownership, administrative boundaries, external access, lifecycle controls, and operational monitoring.",
        "Assess compliance readiness by reviewing Purview, audit, retention, labeling, and data protection readiness signals.",
        "Assess Copilot deployment readiness by translating technical control results into business risk and remediation priorities.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    _heading(doc, "3. Assessment Overview")
    _metrics_table(doc, [
        ("Total Parameters", summary.get("parameter_total", len(rows))),
        ("Passed", summary.get("pass_total", 0)),
        ("Failed", summary.get("fail_total", 0)),
        ("Licensing Required", summary.get("licensing_required_total", 0)),
        ("Manual Validation", summary.get("manual_validation_total", 0)),
        ("Coverage %", _coverage_percent(summary)),
        ("Readiness Score", f"{summary.get('overall_readiness', 0)}%"),
    ])
    _chart(doc, chart_dir, "Overall Readiness", {"Readiness": summary.get("overall_readiness", 0)}, chart_type="bar")
    _chart(doc, chart_dir, "Pass vs Fail", _status_counts(rows), chart_type="donut")

    _heading(doc, "4. Readiness Level")
    critical = _risk_counts(rows).get("Critical", 0)
    fail_pct = _fail_percent(rows)
    readiness_level = _readiness_level(summary, rows)
    _metrics_table(doc, [
        ("Readiness Level", readiness_level),
        ("Fail %", f"{fail_pct}%"),
        ("Critical Findings", critical),
        ("Deployment Guidance", summary.get("deployment_recommendation", "-")),
    ])

    _heading(doc, "5. Risk Matrix")
    risk_counts = _risk_counts(rows)
    failed_total = max(1, sum(risk_counts.values()))
    _table(doc, ["Risk Level", "Count", "Percentage"], [
        [level, risk_counts.get(level, 0), f"{round((risk_counts.get(level, 0) / failed_total) * 100, 2)}%"]
        for level in ["Critical", "High", "Medium", "Low"]
    ])
    _chart(doc, chart_dir, "Risk Category of Parameters Assessed", risk_counts, chart_type="donut")
    _chart(doc, chart_dir, "Severity Status Breakdown", _severity_status_counts(rows), chart_type="stacked")

    _heading(doc, "6. Service Assessment Summary")
    service_summary = _service_summary(rows)
    _table(doc, ["Service", "Pass Count", "Fail Count", "Risk Rating", "Readiness %"], [
        [item["service"], item["pass"], item["fail"], item["risk"], f"{item['readiness']}%"]
        for item in service_summary
    ])
    _chart(doc, chart_dir, "M365 Service Readiness", {item["service"]: item["readiness"] for item in service_summary}, chart_type="bar")

    _heading(doc, "7. Pillar Readiness")
    _chart(doc, chart_dir, "Security, Governance, and Best Practice Readiness", _pillar_readiness(rows), chart_type="bar")

    _heading(doc, "8. Detailed Assessment")
    _detailed_assessment(doc, rows)

    _heading(doc, "9. Licensing Readiness")
    _metrics_table(doc, [
        ("Eligible Users", "Derived from assessment evidence where available"),
        ("Required Licenses", "Microsoft 365 Copilot and prerequisite Microsoft 365 service plans"),
        ("Missing Licenses", len([row for row in rows if _normalized_status(row) == "licensing_required"])),
        ("Copilot Eligibility", "Licensing remediation required" if any(_normalized_status(row) == "licensing_required" for row in rows) else "Eligible"),
    ])
    licensing_rows = [row for row in rows if _normalized_status(row) == "licensing_required"]
    if licensing_rows:
        _table(doc, ["Parameter", "Service", "Status", "Recommendation"], [
            [row.get("title", "-"), _canonical_service(row.get("service")), _status_label(row), row.get("recommendation", "-")]
            for row in licensing_rows
        ])

    _heading(doc, "10. Recommendations")
    _recommendations(doc, rows)

    _heading(doc, "11. Remediation Roadmap")
    actionable_counts = Counter(str(row.get("severity", "info")).lower() for row in rows if _normalized_status(row) != "pass")
    _table(doc, ["Phase", "Focus", "Controls", "Outcome"], [
        ["Phase 1", "Critical findings", actionable_counts.get("critical", 0), "Remove blockers to enterprise Copilot readiness."],
        ["Phase 2", "High findings", actionable_counts.get("high", 0), "Reduce material security and governance exposure."],
        ["Phase 3", "Medium findings", actionable_counts.get("medium", 0), "Improve operational consistency and adoption controls."],
        ["Phase 4", "Optimization", actionable_counts.get("low", 0) + actionable_counts.get("info", 0), "Tune controls and establish continuous readiness monitoring."],
    ])

    _heading(doc, "12. Conclusion")
    for paragraph in _conclusion(summary, rows):
        doc.add_paragraph(paragraph)

    _heading(doc, "Appendix A. Complete Parameter Matrix")
    _parameter_matrix(doc, rows)

    doc.save(path)
    return path


def _cover_page(doc, summary: dict[str, Any], align) -> None:
    doc.add_paragraph()
    logo = doc.add_table(rows=1, cols=1)
    logo.cell(0, 0).text = "Company Logo Placeholder"
    doc.add_paragraph()
    title = doc.add_heading("Copilot Readiness Assessment", 0)
    title.alignment = align.CENTER
    doc.add_paragraph()
    _metrics_table(doc, [
        ("Customer Name", summary.get("customer_name") or "-"),
        ("Tenant Name", summary.get("tenant_name") or summary.get("customer_name") or "-"),
        ("Assessment Date", summary.get("assessment_date") or "-"),
        ("Prepared By", "CRA Platform"),
    ])


def _toc(doc) -> None:
    _heading(doc, "Table of Contents")
    for title in [
        "1. Executive Summary",
        "2. Purpose",
        "3. Assessment Overview",
        "4. Readiness Level",
        "5. Risk Matrix",
        "6. Service Assessment Summary",
        "7. Pillar Readiness",
        "8. Detailed Findings",
        "9. Positive Findings",
        "10. Licensing Readiness",
        "11. Recommendations",
        "12. Remediation Roadmap",
        "13. Conclusion",
        "Appendix A. Complete Parameter Matrix",
    ]:
        doc.add_paragraph(title)


def _heading(doc, text: str) -> None:
    doc.add_heading(text, 1)


def _metrics_table(doc, pairs: list[tuple[Any, Any]]) -> None:
    _table(doc, ["Metric", "Value"], [[key, value] for key, value in pairs])


def _chart_data_table(doc, title: str, values: dict[str, Any]) -> None:
    doc.add_heading(title, 2)
    _table(doc, ["Category", "Value"], [[key, value] for key, value in values.items()])


def _template_aligned_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    template = _load_template_data()
    if not template:
        return rows
    rows_by_slug = {}
    for row in rows:
        for value in [row.get("title"), row.get("parameter_key")]:
            rows_by_slug[_param_slug(value)] = row
    aligned: list[dict[str, Any]] = []
    for service in template.get("detailedAssessment") or []:
        summary_by_id = {item.get("sNo"): item for item in service.get("summaryTable") or []}
        for parameter in service.get("parameters") or []:
            row = _find_template_row(parameter, rows_by_slug)
            summary_row = summary_by_id.get(parameter.get("id")) or {}
            if row is None:
                aligned.append(_placeholder_row(service, parameter, summary_row))
                continue
            aligned.append({
                **row,
                "title": parameter.get("name") or row.get("title"),
                "service": service.get("name") or row.get("service"),
                "pillar": summary_row.get("pillar") or row.get("pillar"),
                "severity": row.get("severity") or summary_row.get("severity") or parameter.get("severity") or "info",
                "documentation_link": row.get("documentation_link") or parameter.get("documentationUrl") or "",
            })
    return aligned


def _load_template_data() -> dict[str, Any] | None:
    template_path = Path(__file__).resolve().parents[3] / "out" / "index.html"
    if not template_path.exists():
        return None
    template = template_path.read_text(encoding="utf-8")
    match = re.search(
        r"// START_REPORT_DATA\s*let\s+REPORT_DATA\s*=\s*(.*?)\s*;\s*// END_REPORT_DATA",
        template,
        flags=re.DOTALL,
    )
    if not match:
        return None
    return json.loads(match.group(1))


def _find_template_row(parameter: dict[str, Any], rows_by_slug: dict[str, dict[str, Any]]) -> dict[str, Any] | None:
    aliases = {
        "devicewithoutcompliancepolicies": "deviceswithoutcompliancepolicies",
        "entratenantcreationbynonadmins": "entratenantcreationbynonadmin",
        "tenantcollaborationinvitation": "tenantcollaborationinvitations",
        "administratorconsentworkflows": "adminconsentworkflow",
        "autoexpirationpolicyform365groups": "autoexpirationpolicyforinactivem365groups",
        "numberofaccountsenabled": "accountenabled",
        "mailboxstatusactiveinactive": "mailboxesstatusactiveinactive",
        "activeinactiveteamsusers": "activerinactiveteamsusers",
        "permissionsettingsforanyonelinks": "permissionsettingforanyonelinks",
        "sensitivesharepointsitesexcludedfromcopilot": "gettingallsiteswithsensitivitykeywordsonatenant",
    }
    candidates = [
        _param_slug(parameter.get("name")),
        _param_slug(parameter.get("rawTitle")),
        _param_slug(parameter.get("parameter")),
    ]
    candidates.extend(aliases.get(candidate, candidate) for candidate in list(candidates))
    for candidate in candidates:
        row = rows_by_slug.get(candidate)
        if row is not None:
            return row
    return None


def _placeholder_row(service: dict[str, Any], parameter: dict[str, Any], summary_row: dict[str, Any]) -> dict[str, Any]:
    return {
        "parameter_key": _param_slug(parameter.get("name")),
        "title": parameter.get("name") or "-",
        "service": service.get("name") or "Microsoft 365",
        "pillar": summary_row.get("pillar") or "Best Practice",
        "category": summary_row.get("pillar") or "Best Practice",
        "severity": (summary_row.get("severity") or parameter.get("severity") or "info").lower(),
        "status": "not_collected",
        "finding": "NOT COLLECTED",
        "actual_result": "Tenant evidence was not collected for this template parameter.",
        "expected_result": "",
        "description": parameter.get("description") or "This control is part of the CRA Word template.",
        "recommendation": "Add or repair the collector/registry mapping for this parameter, run the assessment again, and validate the tenant result.",
        "documentation_link": parameter.get("documentationUrl") or "",
        "artifact_status": "missing_registry_mapping",
    }


def _summary_for_rows(summary: dict[str, Any], rows: list[dict[str, Any]]) -> dict[str, Any]:
    statuses = Counter(_normalized_status(row) for row in rows)
    assessed = [row for row in rows if _normalized_status(row) in {"pass", "warning", "fail", "failed"}]
    points = sum(1.0 if _normalized_status(row) == "pass" else 0.5 if _normalized_status(row) == "warning" else 0.0 for row in assessed)
    readiness = round(points / len(assessed) * 100, 2) if assessed else 0
    return {
        **summary,
        "parameter_total": len(rows),
        "pass_total": statuses.get("pass", 0),
        "warning_total": statuses.get("warning", 0),
        "fail_total": statuses.get("fail", 0) + statuses.get("failed", 0),
        "not_collected_total": statuses.get("not_collected", 0),
        "licensing_required_total": statuses.get("licensing_required", 0),
        "manual_validation_total": statuses.get("manual_validation_required", 0),
        "overall_readiness": readiness,
    }


def _param_slug(value: Any) -> str:
    text = str(value or "").lower().replace("–", "-").replace("—", "-").replace("&", "and")
    text = re.sub(r"^\s*\d+\s*[:.)-]\s*", "", text)
    return re.sub(r"[^a-z0-9]+", "", text)


def _chart(doc, chart_dir: Path, title: str, values: dict[str, Any], *, chart_type: str = "bar") -> None:
    from docx.shared import Inches

    if chart_type == "stacked":
        if not values:
            return
        doc.add_heading(title, 2)
        path = chart_dir / (re_slug(title) + ".png")
        _write_stacked_chart(path, title, values)
        doc.add_picture(str(path), width=Inches(6.3))
        return

    clean_values = {str(key): _numeric(value) for key, value in values.items()}
    clean_values = {key: value for key, value in clean_values.items() if value is not None}
    if not clean_values:
        return
    doc.add_heading(title, 2)
    file_name = re_slug(title) + ".png"
    path = chart_dir / file_name
    if chart_type == "donut":
        _write_donut_chart(path, title, clean_values)
    elif chart_type == "stacked":
        _write_stacked_chart(path, title, values)
    else:
        _write_bar_chart(path, title, clean_values)
    doc.add_picture(str(path), width=Inches(6.3))


def _write_bar_chart(path: Path, title: str, values: dict[str, float]) -> None:
    from PIL import Image, ImageDraw, ImageFont

    width, height = 1150, 520
    margin_left, margin_right = 270, 70
    top, row_height = 82, 42
    image = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(image)
    title_font = _font(30, bold=True)
    label_font = _font(20)
    small_font = _font(18)
    draw.text((34, 24), title, fill="#0f172a", font=title_font)
    max_value = max(max(values.values()), 1)
    if max_value <= 100:
        max_value = 100
    for index, (label, value) in enumerate(values.items()):
        y = top + index * row_height
        if y > height - 48:
            break
        draw.text((34, y + 5), _clip(label, 28), fill="#334155", font=label_font)
        draw.rounded_rectangle((margin_left, y, width - margin_right, y + 24), radius=6, fill="#e2e8f0")
        bar_width = int((width - margin_right - margin_left) * max(0, min(value, max_value)) / max_value)
        draw.rounded_rectangle((margin_left, y, margin_left + bar_width, y + 24), radius=6, fill=_chart_color(index))
        suffix = "%" if max_value == 100 else ""
        draw.text((width - margin_right + 10, y + 1), f"{value:g}{suffix}", fill="#0f172a", font=small_font)
    image.save(path)


def _write_donut_chart(path: Path, title: str, values: dict[str, float]) -> None:
    from PIL import Image, ImageDraw

    width, height = 1150, 560
    image = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(image)
    title_font = _font(30, bold=True)
    label_font = _font(20)
    draw.text((34, 24), title, fill="#0f172a", font=title_font)
    total = sum(max(0, value) for value in values.values()) or 1
    box = (90, 105, 430, 445)
    start = -90
    for index, (_label, value) in enumerate(values.items()):
        angle = 360 * max(0, value) / total
        draw.pieslice(box, start, start + angle, fill=_chart_color(index))
        start += angle
    draw.ellipse((180, 195, 340, 355), fill="#ffffff")
    draw.text((212, 253), f"{int(total)}", fill="#0f172a", font=_font(34, bold=True))
    legend_x, legend_y = 520, 125
    for index, (label, value) in enumerate(values.items()):
        y = legend_y + index * 48
        draw.rounded_rectangle((legend_x, y, legend_x + 26, y + 26), radius=4, fill=_chart_color(index))
        pct = round((max(0, value) / total) * 100, 1)
        draw.text((legend_x + 42, y - 1), f"{label}: {value:g} ({pct:g}%)", fill="#334155", font=label_font)
    image.save(path)


def _write_stacked_chart(path: Path, title: str, values: dict[str, Any]) -> None:
    from PIL import Image, ImageDraw

    statuses = ["Pass", "Fail"]
    severity_counts = _severity_status_counts_from_values(values)
    width, height = 1150, 540
    margin_left, margin_right = 210, 90
    image = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(image)
    draw.text((34, 24), title, fill="#0f172a", font=_font(30, bold=True))
    y = 100
    for index, (severity, counts) in enumerate(severity_counts.items()):
        total = sum(counts.get(status, 0) for status in statuses) or 1
        draw.text((34, y + 3), severity, fill="#334155", font=_font(20))
        x = margin_left
        for status_index, status in enumerate(statuses):
            value = counts.get(status, 0)
            segment = int((width - margin_right - margin_left) * value / total)
            fill = "#10b981" if status == "Pass" else _chart_color(index)
            draw.rectangle((x, y, x + segment, y + 26), fill=fill)
            x += segment
        draw.rectangle((margin_left, y, width - margin_right, y + 26), outline="#cbd5e1", width=2)
        draw.text((width - margin_right + 10, y + 1), f"{counts.get('Fail', 0)}/{total} gaps", fill="#0f172a", font=_font(18))
        y += 55
    draw.rounded_rectangle((34, height - 58, 58, height - 34), radius=4, fill="#10b981")
    draw.text((68, height - 60), "Pass", fill="#334155", font=_font(18))
    draw.rounded_rectangle((145, height - 58, 169, height - 34), radius=4, fill="#ef4444")
    draw.text((179, height - 60), "Fail", fill="#334155", font=_font(18))
    image.save(path)


def _severity_status_counts(rows: list[dict[str, Any]]) -> dict[str, dict[str, int]]:
    result = {
        "Critical": {"Pass": 0, "Fail": 0},
        "High": {"Pass": 0, "Fail": 0},
        "Medium": {"Pass": 0, "Fail": 0},
        "Low": {"Pass": 0, "Fail": 0},
        "Informational": {"Pass": 0, "Fail": 0},
    }
    for row in rows:
        severity = str(row.get("severity") or "info").lower()
        label = "Informational" if severity in {"info", "informational"} else severity.title()
        if label not in result:
            result[label] = {"Pass": 0, "Fail": 0}
        status = "Pass" if _normalized_status(row) == "pass" else "Fail"
        result[label][status] += 1
    return result


def _severity_status_counts_from_values(values: dict[str, Any]) -> dict[str, dict[str, int]]:
    return {
        key: {"Pass": int(value.get("Pass", 0)), "Fail": int(value.get("Fail", 0))}
        for key, value in values.items()
        if isinstance(value, dict)
    }


def _table(doc, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = _text(value)


def _detailed_assessment(doc, rows: list[dict[str, Any]]) -> None:
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in rows:
        grouped[_canonical_service(row.get("service"))].append(row)
    ordered_services = [*SERVICE_NAMES, "Licensing", "Microsoft 365"]
    for service in ordered_services:
        items = grouped.get(service, [])
        if not items:
            continue
        doc.add_heading(service.upper(), 2)
        _table(doc, ["S. No", "Parameter", "CRA Pillar", "Finding", "Severity"], [
            [
                f"{index:02d}",
                row.get("title", "-"),
                _pillar_label(row),
                _status_label(row),
                _severity_label(row),
            ]
            for index, row in enumerate(items, start=1)
        ])
        for index, row in enumerate(items, start=1):
            doc.add_heading(f"{index:02d}: {_text(row.get('title') or 'Parameter')}", 3)
            doc.add_paragraph(f"Risk Rating: {_severity_label(row)} - {_status_label(row)}")
            _metrics_table(doc, [
                ("Description / Tenant Result", row.get("actual_result") or row.get("finding") or "Evidence was not collected for this assessment."),
                ("Expected Result", row.get("expected_result") or row.get("pass_criteria") or row.get("expected_output") or "-"),
                ("Risk", _business_risk(row)),
                ("Recommendation", row.get("recommendation") or "-"),
                ("Microsoft Documentation", row.get("documentation_link") or _microsoft_reference(row)),
                ("Evidence Source", row.get("source_script") or row.get("source_csv") or row.get("artifact_status") or "-"),
            ])


def _detailed_findings(doc, rows: list[dict[str, Any]]) -> None:
    failed_rows = [row for row in rows if _normalized_status(row) in {"fail", "failed"}]
    failed_rows.sort(key=lambda row: SEVERITY_RANK.get(str(row.get("severity", "info")).lower(), 1), reverse=True)
    if not failed_rows:
        doc.add_paragraph("No failed parameters were identified.")
        return
    for row in failed_rows:
        doc.add_heading(_text(row.get("title") or "Finding"), 2)
        _metrics_table(doc, [
            ("CRA Pillar", row.get("pillar", "-")),
            ("Severity", str(row.get("severity", "-")).title()),
            ("Status", _status_label(row)),
            ("Description", row.get("description") or "This control was assessed for Copilot readiness."),
            ("Business Risk", _business_risk(row)),
            ("Actual Result", row.get("actual_result") or row.get("finding") or "-"),
            ("Expected Result", row.get("expected_result") or row.get("pass_criteria") or "-"),
            ("Recommendation", row.get("recommendation") or "-"),
            ("Microsoft Documentation Link", row.get("documentation_link") or "Refer to Microsoft 365 admin documentation for this control."),
        ])


def _positive_findings(doc, rows: list[dict[str, Any]]) -> None:
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in rows:
        if _normalized_status(row) == "pass":
            grouped[_canonical_service(row.get("service"))].append(row)
    if not any(grouped.values()):
        doc.add_paragraph("No passing parameters were identified.")
        return
    for service in SERVICE_NAMES:
        items = grouped.get(service, [])
        if not items:
            continue
        doc.add_heading(service, 2)
        _table(doc, ["Parameter", "Actual Result"], [
            [row.get("title", "-"), row.get("actual_result") or row.get("finding") or "-"]
            for row in items
        ])


def _recommendations(doc, rows: list[dict[str, Any]]) -> None:
    actionable = [row for row in rows if _normalized_status(row) != "pass"]
    if not actionable:
        doc.add_paragraph("No remediation recommendations are required for passing controls.")
        return
    for severity in ["critical", "high", "medium", "low"]:
        items = [row for row in actionable if str(row.get("severity", "info")).lower() == severity]
        if not items:
            continue
        doc.add_heading(severity.title(), 2)
        _table(doc, ["Issue", "Recommendation", "Business Impact", "Implementation Priority"], [
            [row.get("title", "-"), row.get("recommendation", "-"), _business_risk(row), _priority(severity)]
            for row in items
        ])


def _parameter_matrix(doc, rows: list[dict[str, Any]]) -> None:
    headers = ["#", "Parameter", "Service", "Status", "Severity", "Actual Result", "Expected Result", "Recommendation"]
    matrix = []
    for index, row in enumerate(rows, start=1):
        matrix.append([
            index,
            row.get("title", "-"),
            _canonical_service(row.get("service")),
            _status_label(row),
            str(row.get("severity", "-")).title(),
            row.get("actual_result") or row.get("finding") or "-",
            row.get("expected_result") or row.get("pass_criteria") or "-",
            row.get("recommendation") or "-",
        ])
    _table(doc, headers, matrix)


def _conclusion(summary: dict[str, Any], rows: list[dict[str, Any]]) -> list[str]:
    failed = len([row for row in rows if _normalized_status(row) in {"fail", "failed"}])
    licensing = len([row for row in rows if _normalized_status(row) == "licensing_required"])
    manual = len([row for row in rows if _normalized_status(row) == "manual_validation_required"])
    return [
        (
            f"The current Copilot readiness score is {summary.get('overall_readiness', 0)}%, with {failed} failed controls, "
            f"{licensing} licensing-dependent controls, and {manual} controls requiring manual validation."
        ),
        (
            "The organization should remediate critical and high-risk findings first, validate licensing eligibility, "
            "and re-run the assessment before expanding Copilot deployment beyond a controlled pilot."
        ),
        f"Deployment recommendation: {summary.get('deployment_recommendation', '-')}",
    ]


def _text(value: Any) -> str:
    return " ".join(str(value if value is not None else "-").split())


def _numeric(value: Any) -> float | None:
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value or "").strip().replace("%", "")
    try:
        return float(text)
    except ValueError:
        return None


def _font(size: int, *, bold: bool = False):
    from PIL import ImageFont

    candidates = [
        r"C:\Windows\Fonts\aptos-bold.ttf" if bold else r"C:\Windows\Fonts\aptos.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def _chart_color(index: int) -> str:
    colors = ["#ef4444", "#f97316", "#fbbf24", "#2563eb", "#64748b", "#7c3aed", "#10b981", "#0f53c5"]
    return colors[index % len(colors)]


def _clip(value: Any, length: int) -> str:
    text = _text(value)
    return text if len(text) <= length else text[: length - 3].rstrip() + "..."


def re_slug(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-") or "chart"


def _pillar_label(row: dict[str, Any]) -> str:
    value = row.get("pillar") or row.get("category") or "-"
    text = str(value).replace("_", " ").title()
    if text.lower() in {"Best Practices", "Best Practice"}:
        return "Best Practice"
    return text


def _severity_label(row: dict[str, Any]) -> str:
    severity = str(row.get("severity") or "info").lower()
    return "Informational" if severity in {"info", "informational"} else severity.title()
