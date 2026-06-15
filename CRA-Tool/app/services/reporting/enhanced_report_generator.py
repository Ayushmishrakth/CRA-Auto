"""
Enhanced Professional CRA Report Generator with Real Data and Charts
Generates comprehensive Word and PDF reports with:
- Real assessment data from database
- Professional charts (Severity, Pillars, Services)
- Color-coded severity indicators
- Exact structure matching CRA requirements
"""

import io
from datetime import datetime
from typing import Optional, Dict, List
from collections import defaultdict

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from io import BytesIO
import base64


def shade_cell(cell, color):
    """Shade table cell with RGB color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


class EnhancedReportGenerator:
    """Generate professional CRA reports with real data and charts."""

    # Severity colors
    SEVERITY_COLORS = {
        'Critical': {'hex': 'DC143C', 'rgb': (220, 20, 60)},
        'High': {'hex': 'FF8C00', 'rgb': (255, 140, 0)},
        'Medium': {'hex': 'FFD700', 'rgb': (255, 215, 0)},
        'Low': {'hex': 'FFFF99', 'rgb': (255, 255, 153)},
        'Informational': {'hex': 'D3D3D3', 'rgb': (211, 211, 211)},
    }

    SERVICE_MAPPING = {
        'entra': 'ENTRA ID',
        'exchange': 'EXCHANGE ONLINE',
        'purview': 'MICROSOFT PURVIEW',
        'teams': 'MICROSOFT TEAMS',
        'onedrive': 'ONEDRIVE FOR BUSINESS',
        'sharepoint': 'SHAREPOINT ONLINE',
    }

    def __init__(self, assessment_data: dict, logo_path: str = None):
        """Initialize with assessment data and optional logo path."""
        import logging
        logger = logging.getLogger(__name__)

        self.assessment = assessment_data
        self.logo_path = logo_path
        self.doc = Document()
        self.findings = assessment_data.get('findings', [])
        self.summary = assessment_data.get('summary', {})

        # Log initialization
        logger.info(f"[INIT] EnhancedReportGenerator initialized")
        logger.info(f"[INIT] logo_path received: {logo_path}")
        logger.info(f"[INIT] Assessment data keys: {list(assessment_data.keys())}")
        if 'logo_path' in assessment_data:
            logger.info(f"[INIT] assessment_data['logo_path']: {assessment_data.get('logo_path')}")
        logger.info(f"[INIT] Company name: {assessment_data.get('tenant_name')}")
        logger.info(f"[INIT] Company address: {assessment_data.get('company_address')}")

    def _generate_severity_chart(self) -> bytes:
        """Generate severity distribution chart."""
        severity_counts = defaultdict(int)
        for finding in self.findings:
            severity = finding.get('severity', 'Informational')
            severity_counts[severity] += 1

        # Order by severity
        order = ['Critical', 'High', 'Medium', 'Low', 'Informational']
        severities = [s for s in order if s in severity_counts]
        counts = [severity_counts[s] for s in severities]
        colors = [self.SEVERITY_COLORS[s]['rgb'] for s in severities]

        fig, ax = plt.subplots(figsize=(10, 6))
        bars = ax.bar(severities, counts, color=colors, edgecolor='black', linewidth=1.5)

        # Add value labels on bars
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height,
                   f'{int(height)}',
                   ha='center', va='bottom', fontweight='bold', fontsize=11)

        ax.set_ylabel('Number of Parameters', fontsize=12, fontweight='bold')
        ax.set_title('Risk Category of Parameters Assessed', fontsize=14, fontweight='bold')
        ax.set_ylim(0, max(counts) * 1.2)
        ax.grid(axis='y', alpha=0.3)

        img_stream = BytesIO()
        plt.tight_layout()
        fig.savefig(img_stream, format='png', dpi=150, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def _generate_pillar_chart(self) -> bytes:
        """Generate pillar distribution chart."""
        pillar_counts = defaultdict(int)
        for finding in self.findings:
            pillar = finding.get('pillar', 'Best Practices')
            pillar_counts[pillar] += 1

        pillars = list(pillar_counts.keys())
        counts = list(pillar_counts.values())
        colors = ['#FF6B6B', '#4ECDC4', '#45B7D1']  # Red, Teal, Blue

        fig, ax = plt.subplots(figsize=(10, 6))
        bars = ax.bar(pillars, counts, color=colors, edgecolor='black', linewidth=1.5)

        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height,
                   f'{int(height)}',
                   ha='center', va='bottom', fontweight='bold', fontsize=11)

        ax.set_ylabel('Number of Parameters', fontsize=12, fontweight='bold')
        ax.set_title('3 Pillars of Microsoft 365 Copilot Readiness Assessment', fontsize=14, fontweight='bold')
        ax.set_ylim(0, max(counts) * 1.2)
        ax.grid(axis='y', alpha=0.3)

        img_stream = BytesIO()
        plt.tight_layout()
        fig.savefig(img_stream, format='png', dpi=150, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def _generate_service_chart(self) -> bytes:
        """Generate service breakdown chart."""
        service_counts = defaultdict(int)
        service_failed = defaultdict(int)

        for finding in self.findings:
            category = finding.get('category', 'unknown').lower()
            service_counts[category] += 1
            if finding.get('status') == 'fail':
                service_failed[category] += 1

        services = [self.SERVICE_MAPPING.get(k, k.upper()) for k in sorted(service_counts.keys())]
        passed = [service_counts[k] - service_failed[k] for k in sorted(service_counts.keys())]
        failed = [service_failed[k] for k in sorted(service_counts.keys())]

        fig, ax = plt.subplots(figsize=(12, 6))
        x = range(len(services))
        width = 0.35

        bars1 = ax.bar([i - width/2 for i in x], passed, width, label='Passed', color='#2ECC71', edgecolor='black')
        bars2 = ax.bar([i + width/2 for i in x], failed, width, label='Failed', color='#E74C3C', edgecolor='black')

        ax.set_ylabel('Number of Parameters', fontsize=12, fontweight='bold')
        ax.set_title('M365 Services Assessment Summary', fontsize=14, fontweight='bold')
        ax.set_xticks(x)
        ax.set_xticklabels(services, rotation=45, ha='right')
        ax.legend(fontsize=10)
        ax.grid(axis='y', alpha=0.3)

        img_stream = BytesIO()
        plt.tight_layout()
        fig.savefig(img_stream, format='png', dpi=150, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream

    def _add_heading(self, text: str, level: int = 1):
        """Add formatted heading."""
        heading = self.doc.add_heading(text, level=level)
        if level == 1:
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)

    def _add_page_break(self):
        """Add page break."""
        self.doc.add_page_break()

    def generate(self) -> io.BytesIO:
        """Generate complete report and return as BytesIO object."""
        try:
            # Build all sections
            self._add_cover_page()
            self._add_toc()
            self._add_executive_summary()
            self._add_purpose()
            self._add_evaluation_summary()
            self._add_summary_assessment()
            self._add_key_observations()
            self._add_detailed_assessment_by_service()
            self._add_summary_tables()
            self._add_conclusion()

            # Save to bytes
            output = io.BytesIO()
            self.doc.save(output)
            output.seek(0)
            return output
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Report generation failed: {e}")
            raise

    def save(self, filepath: str):
        """Save report to file."""
        data = self.generate()
        with open(filepath, 'wb') as f:
            f.write(data)

    def _add_cover_page(self):
        """Add cover page with logo and company details."""
        import logging
        import os
        from pathlib import Path

        logger = logging.getLogger(__name__)

        # Log what we received
        logger.info(f"[LOGO] _add_cover_page called with logo_path: {self.logo_path}")
        logger.info(f"[LOGO] logo_path type: {type(self.logo_path)}")
        logger.info(f"[LOGO] logo_path is None: {self.logo_path is None}")
        logger.info(f"[LOGO] logo_path is empty string: {self.logo_path == ''}")

        # Add logo if provided
        if self.logo_path and str(self.logo_path).strip():
            try:
                logo_path_str = str(self.logo_path).strip()
                logger.info(f"[LOGO] Processing logo from: {logo_path_str}")

                # Check absolute path
                logo_file = Path(logo_path_str)
                logger.info(f"[LOGO] Absolute path: {logo_file.absolute()}")
                logger.info(f"[LOGO] Exists: {logo_file.exists()}")

                if logo_file.exists():
                    size = logo_file.stat().st_size
                    logger.info(f"[LOGO] File size: {size} bytes")

                    if size > 0:
                        # Add logo centered at top
                        logger.info(f"[LOGO] Adding paragraph for logo...")
                        logo_para = self.doc.add_paragraph()
                        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        # Add picture with detailed error handling
                        try:
                            logger.info(f"[LOGO] Creating run and adding picture...")
                            run = logo_para.add_run()
                            run.add_picture(logo_path_str, width=Inches(1.5))
                            logger.info(f"[LOGO] ✅ Logo added successfully!")
                            self.doc.add_paragraph()  # Spacing after logo
                        except Exception as pic_err:
                            logger.error(f"[LOGO] ❌ Picture insertion failed: {pic_err}", exc_info=True)
                            # Don't add fallback text - just log and continue
                    else:
                        logger.error(f"[LOGO] ❌ File is empty: {logo_path_str}")
                else:
                    # Try creating the path as fallback
                    logger.error(f"[LOGO] ❌ File does not exist: {logo_path_str}")
                    logger.error(f"[LOGO] Current working directory: {os.getcwd()}")
                    logger.error(f"[LOGO] Looking for file at: {logo_file.absolute()}")

                    # List contents of storage/logos if it exists
                    logos_dir = Path("storage/logos")
                    if logos_dir.exists():
                        logger.info(f"[LOGO] Contents of storage/logos:")
                        for f in logos_dir.iterdir():
                            logger.info(f"[LOGO]   - {f.name} ({f.stat().st_size} bytes)")
                    else:
                        logger.error(f"[LOGO] storage/logos directory does not exist")

            except Exception as e:
                logger.exception(f"[LOGO] ❌ Unexpected error: {e}")

        # Title
        title = self.doc.add_heading('Microsoft 365 Copilot Readiness Assessment Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Organization name (use customized name if provided)
        org_name = self.assessment.get('tenant_name', 'Organization')
        org = self.doc.add_paragraph(org_name)
        org.alignment = WD_ALIGN_PARAGRAPH.CENTER
        org.runs[0].font.size = Pt(20)
        org.runs[0].font.bold = True

        # Date
        date_str = self.assessment.get('created_at', datetime.now()).strftime('%d %B %Y') \
            if hasattr(self.assessment.get('created_at'), 'strftime') \
            else str(self.assessment.get('created_at'))

        date = self.doc.add_paragraph(f"Assessment Date: {date_str}")
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Company address if provided
        company_address = self.assessment.get('company_address')
        if company_address:
            address = self.doc.add_paragraph(company_address)
            address.alignment = WD_ALIGN_PARAGRAPH.CENTER
            address.runs[0].font.size = Pt(10)

        self._add_page_break()

    def _add_toc(self):
        """Add table of contents."""
        self._add_heading('Table of Contents', level=1)

        toc_items = [
            ('Executive Summary', 1),
            ('Purpose', 1),
            ('Evaluation Summary', 1),
            ('3 Pillars of Microsoft 365 Copilot Readiness Assessment', 1),
            ('M365 Services assessed in CRA', 1),
            ('Risk Category of Parameters Assessed', 1),
            ('Summary of Assessment', 1),
            ('Key Observations', 1),
            ('Risks of Immediate Deployment', 1),
            ('Recommendations', 1),
            ('Detailed Assessment', 1),
        ]

        # Add services
        services_set = set(f.get('category', 'unknown').lower() for f in self.findings)
        for service_key in sorted(services_set):
            service_name = self.SERVICE_MAPPING.get(service_key, service_key.upper())
            toc_items.append((service_name, 2))

        toc_items.append(('Conclusion', 1))

        for item, level in toc_items:
            p = self.doc.add_paragraph(item, style='List Number' if level == 1 else 'List Bullet 2')

        self._add_page_break()

    def _add_executive_summary(self):
        """Add executive summary."""
        self._add_heading('Executive Summary', level=1)

        tenant = self.assessment.get('tenant_name', 'Organization')
        partner = self.assessment.get('partner_name', 'Assessment Team')

        summary_text = f"""As part of its digital transformation strategy, {tenant} engaged {partner} for a Copilot Readiness Assessment. The purpose of this engagement was to evaluate the Client's Microsoft 365 environment across areas including security, governance, and best practices to determine readiness for the secure and responsible adoption of Microsoft 365 Copilot.

The assessment covered critical services including Entra ID, Exchange Online, Microsoft Teams, SharePoint Online, OneDrive for Business, and Microsoft Purview. It aimed to identify configuration gaps, policy misalignments, and potential vulnerabilities that could impact the responsible use of AI-powered tools like Copilot. By benchmarking the current environment against industry standards and Microsoft's Copilot deployment criteria, the assessment provides a clear roadmap for remediation and optimization.

The findings serve as a strategic foundation for {tenant} to enhance its digital workplace, mitigate operational and compliance risks, and unlock the full potential of Microsoft 365 Copilot. With targeted improvements, the organization can ensure a secure and scalable AI integration that aligns with its long-term business goals."""

        self.doc.add_paragraph(summary_text)
        self._add_page_break()

    def _add_purpose(self):
        """Add purpose section."""
        self._add_heading('Purpose', level=1)

        purposes = [
            'Evaluate the organization\'s environment for alignment with industry best practices.',
            'Assess the environment across Microsoft 365 products and services.',
            'Identify gaps that could pose security or compliance risks upon integrating Copilot.',
            'Establish a baseline for future audits and compliance tracking.',
            'Highlight licensing readiness and user eligibility for Microsoft 365 Copilot deployment.',
            'Provide a risk-based prioritization of remediation efforts.',
            'Offer actionable insights to strengthen governance, data protection, and identity management.',
            'Support strategic decision-making by outlining deployment prerequisites.',
        ]

        for purpose in purposes:
            self.doc.add_paragraph(purpose, style='List Bullet')

        self._add_page_break()

    def _add_evaluation_summary(self):
        """Add evaluation summary."""
        self._add_heading('Evaluation Summary', level=1)

        self._add_heading('3 Pillars of Microsoft 365 Copilot Readiness Assessment', level=2)
        p = self.doc.add_paragraph()
        p.add_run('Security').bold = True
        p.add_run('  •  ')
        p.add_run('Governance').bold = True
        p.add_run('  •  ')
        p.add_run('Best Practices').bold = True

        self._add_heading('M365 Services assessed in CRA', level=2)
        services = ['Entra ID', 'Exchange Online', 'Microsoft Purview', 'Microsoft Teams', 'OneDrive for Business', 'SharePoint Online']
        for svc in services:
            self.doc.add_paragraph(svc, style='List Bullet')

        self._add_heading('Risk Score Matrix', level=2)
        self.doc.add_paragraph('The findings presented in this report are graded according to the following levels of severity:')

        # Add chart
        self.doc.add_paragraph()
        try:
            img_stream = self._generate_pillar_chart()
            self.doc.add_picture(img_stream, width=Inches(5.5))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            self.doc.add_paragraph(f'[Chart generation: {str(e)}]')

        self._add_page_break()

    def _add_summary_assessment(self):
        """Add summary of assessment."""
        self._add_heading('Summary of Assessment', level=1)

        total = self.summary.get('total_parameters', 0)
        failed = self.summary.get('fail_count', 0)
        score = self.assessment.get('overall_score', 0)
        readiness = "Not Ready" if score < 80 else "Ready"

        assessment_text = f"""The Copilot Readiness Assessment uncovered several configuration gaps and policy deficiencies that could impact the secure and compliant adoption of Microsoft Copilot. Each finding has been categorized by severity and mapped to specific areas of risk within the Microsoft 365 environment.

Overall Readiness:
Based on the findings, the Client's current readiness level for Copilot integration is assessed as:

Readiness Level: {readiness}
Readiness Gaps: {failed} out of {total}

{"Significant remediation is required prior to enabling Copilot in the production environment." if score < 80 else "The environment demonstrates good readiness for Copilot deployment."}"""

        self.doc.add_paragraph(assessment_text)

        # Add chart
        self.doc.add_paragraph()
        try:
            img_stream = self._generate_severity_chart()
            self.doc.add_picture(img_stream, width=Inches(5.5))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            self.doc.add_paragraph(f'[Chart generation: {str(e)}]')

        self._add_page_break()

    def _add_key_observations(self):
        """Add key observations."""
        self._add_heading('Key Observations', level=1)

        total = self.summary.get('total_parameters', 0)
        failed = self.summary.get('fail_count', 0)

        observations = [
            f'A total of {failed} gaps out of {total} parameters were identified.',
            'Medium to Critical severity issues make up most findings.',
            'Critical need for immediate remediation across all areas.',
            'Licensing and user eligibility gaps identified.',
            'User profile information incomplete across organization.',
            'Activity metrics show varying engagement levels.',
        ]

        for obs in observations:
            self.doc.add_paragraph(obs, style='List Bullet')

        self._add_heading('Risks of Immediate Deployment', level=2)
        risks = [
            'Data exposure through inadequate classification.',
            'Compliance violations due to policy gaps.',
            'Operational risks from misconfigured environments.',
            'Security vulnerabilities from weak identity controls.',
        ]
        for risk in risks:
            self.doc.add_paragraph(risk, style='List Bullet')

        self._add_heading('Recommendations', level=2)
        recs = [
            'Remediation of identified gaps to meet security baseline standards.',
            'Postpone deployment until critical issues are resolved.',
            'Implement recommended controls for regulatory compliance.',
        ]
        for rec in recs:
            self.doc.add_paragraph(rec, style='List Bullet')

        self._add_page_break()

    def _add_detailed_assessment_by_service(self):
        """Add detailed assessment organized by service."""
        self._add_heading('Detailed Assessment', level=1)

        # Group by service
        findings_by_service = defaultdict(list)
        for finding in self.findings:
            service = finding.get('category', 'unknown').lower()
            findings_by_service[service].append(finding)

        # Process each service
        for service_key in sorted(findings_by_service.keys()):
            service_name = self.SERVICE_MAPPING.get(service_key, service_key.upper())
            self._add_heading(service_name, level=2)

            service_findings = findings_by_service[service_key]

            for idx, finding in enumerate(sorted(service_findings, key=lambda f: f.get('severity', 'Low')), 1):
                param_name = finding.get('parameter_name', f'Finding {idx}')
                severity = finding.get('severity', 'Informational')
                status = finding.get('status', 'not_collected')
                description = finding.get('description', '')
                risk = finding.get('risk', '')

                # Parameter heading
                self._add_heading(f'{idx:02d}: {param_name}', level=3)

                # Risk rating with color
                p = self.doc.add_paragraph()
                run = p.add_run(f'Risk Rating: {severity} - {status.capitalize()}')
                run.bold = True
                run.font.size = Pt(11)
                color = self.SEVERITY_COLORS.get(severity, {}).get('rgb', (0, 0, 0))
                run.font.color.rgb = RGBColor(*color)

                # Description
                p = self.doc.add_paragraph()
                p.add_run('Description: ').bold = True
                p.add_run(description)

                # Risk
                p = self.doc.add_paragraph()
                p.add_run('Risk: ').bold = True
                p.add_run(risk)

            self._add_page_break()

    def _add_summary_tables(self):
        """Add summary tables by service."""
        self._add_heading('Assessment Summary Tables', level=1)

        findings_by_service = defaultdict(list)
        for finding in self.findings:
            service = finding.get('category', 'unknown').lower()
            findings_by_service[service].append(finding)

        for service_key in sorted(findings_by_service.keys()):
            service_name = self.SERVICE_MAPPING.get(service_key, service_key.upper())
            self._add_heading(service_name, level=2)

            findings = findings_by_service[service_key]

            # Create table
            table = self.doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'

            # Header
            header_cells = table.rows[0].cells
            headers = ['S. No', 'Parameter', 'CRA Pillar', 'Finding', 'Severity']
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                shade_cell(header_cells[i], '003366')
                for para in header_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True

            # Data rows
            for idx, finding in enumerate(findings, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = f'{idx:02d}'
                row_cells[1].text = finding.get('parameter_name', '')

                pillar = finding.get('pillar', '')
                row_cells[2].text = pillar
                # Color-code pillar column
                if pillar == 'Security':
                    shade_cell(row_cells[2], 'FFE6E6')
                elif pillar == 'Governance':
                    shade_cell(row_cells[2], 'E6F2FF')
                elif pillar == 'Best Practices':
                    shade_cell(row_cells[2], 'E6F9E6')

                status = finding.get('status', 'not_collected').capitalize()
                row_cells[3].text = status
                # Color-code status column
                if status == 'Pass':
                    shade_cell(row_cells[3], '90EE90')
                elif status == 'Fail':
                    shade_cell(row_cells[3], 'FFB6C6')

                severity = finding.get('severity', 'Informational')
                row_cells[4].text = severity
                severity_color = self.SEVERITY_COLORS.get(severity, {}).get('hex', 'E0E0E0')
                shade_cell(row_cells[4], severity_color)

                # Add text color for better readability on dark backgrounds
                for cell in [row_cells[2], row_cells[3], row_cells[4]]:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True

            self.doc.add_paragraph()

        self._add_page_break()

    def _add_conclusion(self):
        """Add conclusion."""
        self._add_heading('Conclusion', level=1)

        tenant = self.assessment.get('tenant_name', 'Organization')
        partner = self.assessment.get('partner_name', 'Assessment Team')
        total = self.summary.get('total_parameters', 0)
        failed = self.summary.get('fail_count', 0)
        score = self.assessment.get('overall_score', 0)

        conclusion = f"""The Copilot Readiness Assessment for {tenant} reveals that the current Microsoft 365 environment {"is not yet prepared" if score < 80 else "is well-positioned"} for the secure and compliant deployment of Microsoft 365 Copilot. With {failed} out of {total} parameters {"failing to meet readiness standards" if failed > 0 else "meeting readiness standards"}, there is {"a clear need for immediate remediation" if failed > 0 else "excellent alignment with best practices"}.

Key gaps were identified across all three foundational pillars: Security, Governance, and Best Practices. Critical vulnerabilities such as the absence of sensitivity labels, permissive external sharing configurations, and insufficient audit logging significantly elevate the risk of data exposure and non-compliance.

To mitigate these risks and ensure a successful Copilot rollout, {partner} strongly recommends a phased remediation strategy. This should prioritize the resolution of critical and high-severity issues, followed by medium and low-risk items. Only after these gaps are addressed should the organization consider enabling Copilot in the production environment.

By aligning with the recommendations outlined in this report, {tenant} can enhance its security posture, ensure regulatory compliance, and fully leverage the transformative potential of Microsoft 365 Copilot."""

        self.doc.add_paragraph(conclusion)
