"""
Professional CRA Report Generator - creates Word and PDF reports from assessment data.
Uses python-docx for Word generation with full styling and formatting.
"""

import io
from datetime import datetime
from typing import Optional
from collections import defaultdict

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, color):
    """Shade a table cell with the given RGB color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


class ProfessionalReportGenerator:
    """Generate professional Word and PDF reports from assessment data."""

    SEVERITY_COLORS = {
        'Critical': 'FF0000',
        'High': 'FF6600',
        'Medium': 'FFCC00',
        'Low': 'FFFF00',
        'Informational': 'E0E0E0',
    }

    SERVICE_MAPPING = {
        'entra': 'ENTRA ID',
        'exchange': 'EXCHANGE ONLINE',
        'purview': 'MICROSOFT PURVIEW',
        'teams': 'MICROSOFT TEAMS',
        'onedrive': 'ONEDRIVE FOR BUSINESS',
        'sharepoint': 'SHAREPOINT ONLINE',
    }

    def __init__(self, assessment_data: dict):
        """
        Initialize with assessment data.

        assessment_data should contain:
        - id: Assessment ID
        - tenant_id: Tenant ID
        - tenant_name: Organization name
        - partner_name: Partner/Assessor name
        - created_at: Assessment date
        - overall_score: Overall readiness score
        - findings: List of finding objects
        - summary: Summary statistics
        """
        self.assessment = assessment_data
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Setup document styles."""
        styles = self.doc.styles

        # Title style
        try:
            title_style = styles['Heading 1']
        except:
            title_style = styles.add_style('Heading 1', 1)
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 51, 102)

    def _add_heading(self, text: str, level: int = 1):
        """Add a formatted heading."""
        heading = self.doc.add_heading(text, level=level)
        if level == 1:
            heading.style = 'Heading 1'

    def _add_table_of_contents(self):
        """Add table of contents."""
        self._add_heading('Table of Contents', level=1)

        services = self._group_findings_by_service()
        toc_items = [
            ('Executive Summary', 2),
            ('Purpose', 2),
            ('Evaluation Summary', 2),
            ('3 Pillars of Microsoft 365 Copilot Readiness Assessment', 2),
            ('M365 Services assessed in CRA', 2),
            ('Risk Category of Parameters Assessed', 2),
            ('Summary of Assessment', 2),
            ('Key Observations', 2),
            ('Risks of Immediate Deployment', 2),
            ('Recommendations', 2),
            ('Detailed Assessment', 2),
        ]

        for service_key in services.keys():
            service_name = self.SERVICE_MAPPING.get(service_key, service_key.upper())
            toc_items.append((service_name, 3))

        toc_items.append(('Conclusion', 2))

        for item, indent in toc_items:
            p = self.doc.add_paragraph(item, style=f'List Bullet {indent}')
            p.paragraph_format.left_indent = Inches(0.25 * (indent - 1))

        self.doc.add_page_break()

    def _add_executive_summary(self):
        """Add executive summary section."""
        self._add_heading('Executive Summary', level=1)

        tenant_name = self.assessment.get('tenant_name', 'Organization')
        partner_name = self.assessment.get('partner_name', 'Assessment Team')
        assessment_date = self.assessment.get('created_at', datetime.now())

        intro = f"""As part of its digital transformation strategy, {tenant_name} engaged {partner_name} for a Copilot Readiness Assessment. The purpose of this engagement was to evaluate the Client's Microsoft 365 environment across areas including security, governance, and best practices to determine readiness for the secure and responsible adoption of Microsoft 365 Copilot.

The assessment covered critical services including Entra ID, Exchange Online, Microsoft Teams, SharePoint Online, OneDrive for Business, and Microsoft Purview. It aimed to identify configuration gaps, policy misalignments, and potential vulnerabilities that could impact the responsible use of AI-powered tools like Copilot. By benchmarking the current environment against industry standards and Microsoft's Copilot deployment criteria, the assessment provides a clear roadmap for remediation and optimization.

The findings serve as a strategic foundation for {tenant_name} to enhance its digital workplace, mitigate operational and compliance risks, and unlock the full potential of Microsoft 365 Copilot. With targeted improvements, the organization can ensure a secure and scalable AI integration that aligns with its long-term business goals."""

        self.doc.add_paragraph(intro)
        self.doc.add_page_break()

    def _add_purpose(self):
        """Add purpose section."""
        self._add_heading('Purpose', level=1)

        purposes = [
            'Evaluate the organization\'s environment for alignment with industry best practices.',
            'Assess the environment across Microsoft 365 products and services like SharePoint, Teams, OneDrive for business etc.',
            'Identify gaps that could pose security or compliance risks upon integrating Copilot.',
            'Establish a baseline for future audits and compliance tracking related to AI usage within Microsoft 365.',
            'Highlight licensing readiness and user eligibility for Microsoft 365 Copilot deployment.',
            'Provide a risk-based prioritization of remediation efforts to guide Copilot enablement planning.',
            'Offer actionable insights to strengthen governance, data protection, and identity management in preparation for AI integration.',
            'Support strategic decision-making by outlining Copilot deployment prerequisites and dependencies.',
        ]

        for purpose in purposes:
            self.doc.add_paragraph(purpose, style='List Bullet')

        self.doc.add_page_break()

    def _add_evaluation_summary(self):
        """Add evaluation summary."""
        self._add_heading('Evaluation Summary', level=1)

        self._add_heading('3 Pillars of Microsoft 365 Copilot Readiness Assessment', level=2)
        p = self.doc.add_paragraph()
        p.add_run('Security').bold = True
        p.add_run(' • ')
        p.add_run('Governance').bold = True
        p.add_run(' • ')
        p.add_run('Best Practices').bold = True

        self._add_heading('M365 Services assessed in CRA', level=2)
        services_list = [
            'Entra ID',
            'Exchange Online',
            'Microsoft Purview',
            'Microsoft Teams',
            'OneDrive for Business',
            'SharePoint Online',
        ]
        for service in services_list:
            self.doc.add_paragraph(service, style='List Bullet')

        self.doc.add_page_break()

    def _add_summary_assessment(self):
        """Add summary of assessment section."""
        self._add_heading('Summary of Assessment', level=1)

        summary = self.assessment.get('summary', {})
        overall_score = self.assessment.get('overall_score', 0)
        fail_count = summary.get('fail_count', 0)
        total_count = summary.get('total_parameters', 65)

        assessment_text = f"""The Copilot Readiness Assessment uncovered several configuration gaps and policy deficiencies that could impact the secure and compliant adoption of Microsoft Copilot. Each finding has been categorized by severity and mapped to specific areas of risk within the Microsoft 365 environment.

Overall Readiness:
Based on the findings, the Client's current readiness level for Copilot integration is assessed as:

Readiness Level: {"Ready" if overall_score >= 80 else "Not Ready"}
Readiness Gaps: {fail_count} out of {total_count}

{"Significant remediation is required prior to enabling Copilot in the production environment." if overall_score < 80 else "The environment is prepared for Copilot deployment with minor considerations."}"""

        self.doc.add_paragraph(assessment_text)
        self.doc.add_page_break()

    def _add_key_observations(self):
        """Add key observations section."""
        self._add_heading('Key Observations', level=1)

        summary = self.assessment.get('summary', {})
        fail_count = summary.get('fail_count', 0)
        total_count = summary.get('total_parameters', 65)

        observations = [
            f'A total of {fail_count} gaps out of {total_count} parameters were identified, distributed across Security, Governance, and Best Practice categories.',
            'Medium to Critical severity issues make up most of the findings, indicating substantial exposure to operational and compliance risks.',
            'Gap findings reveal that the percentage of failed parameters indicates a critical need for immediate remediation in those areas.',
            'There are specific findings related to user eligibility for M365 Copilot licenses. Copilot requires a base Microsoft 365 subscription, such as Microsoft 365 E3, E5, Business Standard, or Business Premium.',
            'User information completeness and organizational hierarchy documentation are essential for accurate Copilot functionality.',
            'Activity metrics across Microsoft 365 services show varying levels of engagement, which impacts the scope and risks of Copilot deployment.',
        ]

        for obs in observations:
            self.doc.add_paragraph(obs, style='List Bullet')

        self._add_heading('Risks of Immediate Deployment', level=2)
        self.doc.add_paragraph('Proceeding with Copilot activation in the current state may lead to:')
        risks = [
            'Data exposure through inadequate classification and access controls.',
            'Compliance violations due to policy gaps and insufficient audit logging.',
            'Operational risks from unmanaged or misconfigured M365 environments.',
            'Security vulnerabilities from weak identity and governance controls.',
        ]
        for risk in risks:
            self.doc.add_paragraph(risk, style='List Bullet')

        self._add_heading('Recommendations', level=2)
        recommendations = [
            'Remediation of identified gaps: Address all findings regardless of severity to meet cybersecurity baseline standards.',
            'Postpone Deployment: Due to the current maturity level of the environment, it is recommended to adopt Copilot deployment after all critical and high-priority gaps are resolved.',
            'Futureproofing: Implementing the recommendations provided will reduce security risks and ensure regulatory compliance during and after Copilot integration.',
        ]
        for rec in recommendations:
            self.doc.add_paragraph(rec, style='List Bullet')

        self.doc.add_page_break()

    def _group_findings_by_service(self) -> dict:
        """Group findings by service/category."""
        findings = self.assessment.get('findings', [])
        grouped = defaultdict(list)

        for finding in findings:
            category = finding.get('category', 'unknown').lower()
            grouped[category].append(finding)

        return grouped

    def _add_detailed_assessment(self):
        """Add detailed assessment by service."""
        self._add_heading('Detailed Assessment', level=1)

        services = self._group_findings_by_service()

        for idx, (service_key, findings) in enumerate(sorted(services.items())):
            service_name = self.SERVICE_MAPPING.get(service_key, service_key.upper())
            self._add_heading(service_name, level=2)

            # Add service findings
            for finding_idx, finding in enumerate(findings, 1):
                param_name = finding.get('parameter_name', f'Finding {finding_idx}')
                severity = finding.get('severity', 'Informational')
                status = finding.get('status', 'not_collected')
                description = finding.get('description', 'No description available.')
                risk = finding.get('risk', 'No risk information available.')

                # Heading for each parameter
                self._add_heading(f'{finding_idx:02d}: {param_name}', level=3)

                # Risk rating
                p = self.doc.add_paragraph()
                p.add_run(f'Risk Rating: {severity} - {status.capitalize()}').bold = True

                # Description
                p = self.doc.add_paragraph()
                p.add_run('Description: ').bold = True
                p.add_run(description)

                # Risk
                p = self.doc.add_paragraph()
                p.add_run('Risk: ').bold = True
                p.add_run(risk)

                self.doc.add_paragraph()  # Spacing

            if idx < len(services) - 1:
                self.doc.add_page_break()

    def _add_summary_tables(self):
        """Add summary tables for each service."""
        self._add_heading('Assessment Summary Tables', level=1)

        services = self._group_findings_by_service()

        for service_key, findings in sorted(services.items()):
            service_name = self.SERVICE_MAPPING.get(service_key, service_key.upper())
            self._add_heading(service_name, level=2)

            # Create table
            table = self.doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'

            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'S. No'
            header_cells[1].text = 'Parameter'
            header_cells[2].text = 'CRA Pillar'
            header_cells[3].text = 'Finding'
            header_cells[4].text = 'Severity'

            # Shade header
            for cell in header_cells:
                shade_cell(cell, '003366')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)

            # Data rows
            for idx, finding in enumerate(findings, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = f'{idx:02d}'
                row_cells[1].text = finding.get('parameter_name', '')
                row_cells[2].text = finding.get('pillar', 'N/A')
                row_cells[3].text = finding.get('status', 'not_collected').capitalize()
                row_cells[4].text = finding.get('severity', 'Informational')

                # Shade severity cell
                severity = finding.get('severity', 'Informational')
                color = self.SEVERITY_COLORS.get(severity, 'E0E0E0')
                shade_cell(row_cells[4], color)

            self.doc.add_paragraph()

        self.doc.add_page_break()

    def _add_conclusion(self):
        """Add conclusion section."""
        self._add_heading('Conclusion', level=1)

        tenant_name = self.assessment.get('tenant_name', 'Organization')
        partner_name = self.assessment.get('partner_name', 'Assessment Team')
        summary = self.assessment.get('summary', {})
        fail_count = summary.get('fail_count', 0)
        total_count = summary.get('total_parameters', 65)
        overall_score = self.assessment.get('overall_score', 0)

        conclusion = f"""The Copilot Readiness Assessment for {tenant_name} reveals that the current Microsoft 365 environment {"is not yet prepared" if overall_score < 80 else "is prepared"} for the secure and compliant deployment of Microsoft 365 Copilot. With {fail_count} out of {total_count} parameters {"failing to meet readiness standards" if fail_count > 0 else "meeting readiness standards"}, there is {"a clear need for immediate and comprehensive remediation" if fail_count > 0 else "excellent alignment with best practices"}.

Key gaps were identified across all three foundational pillars: Security, Governance, and Best Practices. Critical vulnerabilities such as the absence of sensitivity labels, permissive external sharing configurations, and insufficient audit logging significantly elevate the risk of data exposure and non-compliance. Furthermore, the lack of complete user profile information and inconsistent policy enforcement could impair Copilot's ability to deliver accurate, context-aware insights.

To mitigate these risks and ensure a successful Copilot rollout, {partner_name} strongly recommends a phased remediation strategy. This should prioritise the resolution of critical and high-severity issues, followed by medium and low-risk items. Only after these gaps are addressed should the organisation consider enabling Copilot in the production environment.

By aligning with the recommendations outlined in this report, {tenant_name} can enhance its security posture, ensure regulatory compliance, and fully leverage the transformative potential of Microsoft 365 Copilot."""

        self.doc.add_paragraph(conclusion)

    def generate_word_report(self) -> bytes:
        """Generate and return Word document as bytes."""
        # Cover page
        title = self.doc.add_heading('Microsoft 365 Copilot Readiness Assessment Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = self.doc.add_paragraph(self.assessment.get('tenant_name', 'Organization'))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(18)

        date_para = self.doc.add_paragraph(f"Assessment Date: {self.assessment.get('created_at', datetime.now()).strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_page_break()

        # Add sections
        self._add_table_of_contents()
        self._add_executive_summary()
        self._add_purpose()
        self._add_evaluation_summary()
        self._add_summary_assessment()
        self._add_key_observations()
        self._add_detailed_assessment()
        self._add_summary_tables()
        self._add_conclusion()

        # Save to bytes
        output = io.BytesIO()
        self.doc.save(output)
        output.seek(0)
        return output.getvalue()

    def save_word_report(self, filepath: str):
        """Save Word report to file."""
        self.doc.save(filepath)
        return filepath
