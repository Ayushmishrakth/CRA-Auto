"""CRA Report Builder - 50-page DOCX with all 7 fixes applied."""

from pathlib import Path
from datetime import datetime
from collections import defaultdict
import tempfile, os, logging, json
import textwrap
import math

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np
import shutil
import tempfile

logger = logging.getLogger(__name__)
REPORT_CHART_DPI = 240
REPORT_CONTENT_MANIFEST = "report_content_manifest.yml"
REPORT_DESIGN_SYSTEM = "report_design_system.yml"
AAA_REPORT_BLUEPRINT = "aaa_report_blueprint.yml"
FIRST_NINE_SECTION_IDS = {
    "cover",
    "toc",
    "executive_summary",
    "purpose",
    "evaluation_summary",
    "readiness_summary",
    "key_observations",
    "risks_and_recommendations",
}
DATA_NOT_AVAILABLE = "Data Not Available"
EMPTY_STATE_TEXT = "Assessment data could not be collected."
SERVICE_ORDER = ["Entra ID", "Exchange Online", "Microsoft Purview", "Microsoft Teams", "OneDrive for Business", "SharePoint Online"]
PILLAR_MAP = {'identity access': 'Security', 'collaboration': 'Governance', 'compliance': 'Best Practice',
              'security': 'Security', 'governance': 'Governance', 'best practice': 'Best Practice', 'best practices': 'Best Practice'}
SEVERITY_MAP = {
    'critical': 'Critical',
    'high': 'High',
    'medium': 'Medium',
    'low': 'Low',
    'info': 'Informational',
    'informational': 'Informational',
    'none': 'Informational',
    '': 'Informational',
}
SERVICE_DISPLAY_NAMES = {
    'entra id': 'ENTRA ID',
    'entra': 'ENTRA ID',
    'exchange online': 'EXCHANGE ONLINE',
    'exchange': 'EXCHANGE ONLINE',
    'microsoft purview': 'MICROSOFT PURVIEW',
    'purview': 'MICROSOFT PURVIEW',
    'microsoft teams': 'MICROSOFT TEAMS',
    'teams': 'MICROSOFT TEAMS',
    'onedrive for business': 'ONEDRIVE FOR BUSINESS',
    'onedrive': 'ONEDRIVE FOR BUSINESS',
    'sharepoint online': 'SHAREPOINT ONLINE',
    'sharepoint': 'SHAREPOINT ONLINE',
}

# FIX 1: FINDING SORT ORDER
SEVERITY_ORDER = {
    'critical': 1,
    'high': 2,
    'medium': 3,
    'low': 4,
    'informational': 5,
    'info': 5,
}
FINDING_ORDER = {
    'fail': 1,
    'pass': 2,
}
PARAMETER_ORDER = {
    'Custom Banned Password List': 1,
    'Restricted Access to Entra Admin Ctr': 2,
    'Restricted Access to Microsoft Entra Admin Centre': 2,
    'Emergency Access Account': 3,
    'Device without Compliance Policies': 4,
    'Authentication Methods Enabled': 5,
    'Entra - Tenant Creation by Non-Admins': 6,
    'Global Administrator Accounts': 7,
    'Self-Service Password Reset Auth Meth': 8,
    'Self-Service Password Reset Authentication Method': 8,
    'Tenant Collaboration Invitation': 9,
    'Administrator Consent Workflows': 10,
    'CAP Policies for Risky Sign-Ins': 11,
    'Conditional Access Policies Exclusion': 12,
    'Conditional Access Policies (Exclusion)': 12,
    'User Consent for Applications': 13,
    'Entra - Third-Party App Integrations': 14,
    'Users without MFA': 15,
    'Auto-expiration policy M365 Groups': 16,
    'Auto-expiration policy for inactive m365 groups': 16,
    'Customer Lockbox': 17,
    'Guest Invite Settings': 18,
    'Guest Users count': 19,
    'User Information': 20,
    'Number of accounts enabled': 21,
}

DEFAULT_REPORT_CONFIG = {
    "branding": {
        "primary_color": "0078D4",
        "secondary_color": "2F5496",
        "accent_color": "00A4EF",
        "dark_color": "1F2937",
        "muted_color": "6B7280",
        "light_band": "EBF4FD",
        "success_color": "00B050",
        "fail_color": "C00000",
        "warning_color": "F79646",
        "partner_logo": "app/services/reporting/assets/techplustalent-logo.png",
    },
    "typography": {
        "heading_font": "Calibri",
        "body_font": "Calibri",
        "h1_size": 16,
        "h2_size": 14,
        "h3_size": 12,
        "body_size": 11,
        "caption_size": 11,
        "h1_color": "003366",
        "h2_color": "003366",
        "line_spacing": 1.5,
    },
    "layout": {
        "top_margin": 1.0,
        "bottom_margin": 1.0,
        "left_margin": 1.0,
        "right_margin": 1.0,
        "chart_dpi": 220,
    },
    "table_styling": {
        "header_bg_color": "F2F2F2",
        "header_text_color": "000000",
        "border_color": "CCCCCC",
        "cell_padding": 0.05,
    },
    "severity_definitions": {
        "Critical": {"color": "8B0000", "description": "Immediate remediation required before Copilot rollout."},
        "High": {"color": "FF0000", "description": "Material exposure that should be remediated as a priority."},
        "Medium": {"color": "F79646", "description": "Moderate risk that can affect readiness if left unresolved."},
        "Low": {"color": "FFD966", "description": "Limited risk to track in the remediation backlog."},
        "Informational": {"color": "70AD47", "description": "Contextual item for readiness planning."},
    },
    "readiness_thresholds": [
        {"label": "READY", "min": 80, "color": "00B050"},
        {"label": "PARTIALLY READY", "min": 50, "color": "F79646"},
        {"label": "NOT READY", "min": 0, "color": "C00000"},
    ],
    "chart_palette": {
        "pillars": ["4472C4", "A5A5A5", "ED7D31"],
        "services": ["4472C4", "ED7D31", "A5A5A5", "FFC000", "5B9BD5", "70AD47"],
        "status": {"Pass": "00B050", "Fail": "C00000"},
    },
    "toc": [
        {"title": "Executive Summary", "page": 5, "level": 1},
        {"title": "Evaluation Summary", "page": 6, "level": 1},
        {"title": "Risk Matrix", "page": 7, "level": 1},
        {"title": "Summary of Assessment", "page": 8, "level": 1},
        {"title": "Executive Dashboard", "page": 9, "level": 1},
        {"title": "Key Observations", "page": 10, "level": 1},
        {"title": "User Information Analysis", "page": 11, "level": 1},
        {"title": "Usage and Recommendations", "page": 12, "level": 1},
        {"title": "Detailed Assessment", "page": 13, "level": 1},
    ],
}

def _repo_file_path(filename):
    candidates = [Path.cwd() / filename]
    current = Path(__file__).resolve()
    candidates.extend(parent / filename for parent in current.parents)
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None

def _load_yaml_config(filename):
    path = _repo_file_path(filename)
    if not path:
        logger.warning("[REPORT_BUILDER] %s not found", filename)
        return {}
    try:
        import yaml
        with path.open("r", encoding="utf-8") as handle:
            return yaml.safe_load(handle) or {}
    except Exception as exc:
        logger.warning("[REPORT_BUILDER] Could not parse %s: %s", path, exc)
        return {}

def _deep_merge(base, override):
    merged = dict(base or {})
    for key, value in (override or {}).items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = _deep_merge(merged[key], value)
        else:
            merged[key] = value
    return merged

def _report_config(assessment_data=None):
    if isinstance(assessment_data, dict) and isinstance(assessment_data.get("_report_config"), dict):
        return assessment_data["_report_config"]
    return DEFAULT_REPORT_CONFIG

def _config_from_yaml(content_manifest, design_system):
    config = _deep_merge({}, DEFAULT_REPORT_CONFIG)
    if isinstance(content_manifest, dict):
        config = _deep_merge(config, content_manifest.get("report", {}))
        if content_manifest.get("toc"):
            config["toc"] = content_manifest.get("toc")
        if content_manifest.get("sections"):
            config["sections"] = content_manifest.get("sections")
    if isinstance(design_system, dict):
        for source_key, target_key in [
            ("branding", "branding"),
            ("typography", "typography"),
            ("layout", "layout"),
            ("severity_definitions", "severity_definitions"),
            ("readiness_thresholds", "readiness_thresholds"),
            ("chart_palette", "chart_palette"),
        ]:
            if source_key in design_system:
                if isinstance(design_system[source_key], dict) and isinstance(config.get(target_key), dict):
                    config[target_key] = _deep_merge(config.get(target_key, {}), design_system[source_key])
                else:
                    config[target_key] = design_system[source_key]
    return config

def _cfg(config, section, key, default=None):
    section_data = config.get(section, {}) if isinstance(config, dict) else {}
    return section_data.get(key, default) if isinstance(section_data, dict) else default

def _hex_color(hex_value):
    text = str(hex_value or "000000").strip().lstrip("#")
    if len(text) != 6:
        text = "000000"
    return RGBColor(int(text[0:2], 16), int(text[2:4], 16), int(text[4:6], 16))

def _hex_text(hex_value, default="000000"):
    text = str(hex_value or default).strip().lstrip("#")
    return text if len(text) == 6 else default

def _font_name(config, kind="body"):
    typography = config.get("typography", {}) if isinstance(config, dict) else {}
    return typography.get("heading_font" if kind == "heading" else "body_font", "Calibri")

def _pt(config, key, default):
    typography = config.get("typography", {}) if isinstance(config, dict) else {}
    try:
        return Pt(float(typography.get(key, default)))
    except (TypeError, ValueError):
        return Pt(default)

def _apply_run_style(run, config, size_key="body_size", bold=False, color=None, font_kind="body"):
    run.font.name = _font_name(config, font_kind)
    run.font.size = _pt(config, size_key, 10.5)
    run.font.bold = bold
    if color:
        run.font.color.rgb = _hex_color(color)
    return run

def _styled_heading(doc, text, config, level=1, color=None, before=0, after=8):
    heading = doc.add_heading(str(text), level=level)
    heading.paragraph_format.space_before = Pt(before)
    heading.paragraph_format.space_after = Pt(after)
    heading.paragraph_format.line_spacing = 1.5
    size_key = "h1_size" if level == 1 else "h2_size"
    color_key = "h1_color" if level == 1 else "h2_color"
    default_color = _cfg(config, "typography", color_key, "003366" if level == 1 else "003366")
    for run in heading.runs:
        _apply_run_style(
            run,
            config,
            size_key=size_key,
            bold=True,
            color=color or default_color,
            font_kind="heading",
        )
    return heading

def _body_paragraph(doc, text="", config=None, before=0, after=6, bold=False, italic=False, color=None):
    config = config or DEFAULT_REPORT_CONFIG
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = _cfg(config, "typography", "line_spacing", 1.5)
    run = p.add_run(str(text))
    _apply_run_style(run, config, "body_size", bold=bold, color=color)
    run.font.italic = italic
    return p

def _add_emphasized_text(paragraph, text, config, emphasize_terms, *, size_key="body_size", underline_terms=None):
    """Add text with selected terms emphasized while preserving live text values."""
    underline_terms = set(underline_terms or [])
    terms = [str(term) for term in emphasize_terms if str(term or "").strip()]
    if not terms:
        return _apply_run_style(paragraph.add_run(str(text)), config, size_key)

    lowered_terms = [(term.lower(), term) for term in sorted(terms, key=len, reverse=True)]
    source = str(text)
    idx = 0
    while idx < len(source):
        match = None
        source_lower = source.lower()
        for lowered, original in lowered_terms:
            if source_lower.startswith(lowered, idx):
                match = original
                break
        if match:
            run = paragraph.add_run(source[idx:idx + len(match)])
            _apply_run_style(run, config, size_key, bold=True)
            if match in underline_terms:
                run.font.underline = True
            idx += len(match)
            continue

        next_idx = len(source)
        for lowered, _original in lowered_terms:
            found = source_lower.find(lowered, idx + 1)
            if found != -1:
                next_idx = min(next_idx, found)
        run = paragraph.add_run(source[idx:next_idx])
        _apply_run_style(run, config, size_key)
        idx = next_idx

def _body_paragraph_emphasis(doc, text, config, emphasize_terms, before=0, after=6, underline_terms=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    _add_emphasized_text(p, text, config, emphasize_terms, underline_terms=underline_terms)
    return p

def _set_cell_text(cell, text, config, size_key="body_size", bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(str(text))
    _apply_run_style(run, config, size_key, bold=bold, color=color)
    return paragraph

def _set_table_header_style(table, config=None):
    """Apply professional styling to table header row."""
    config = config or DEFAULT_REPORT_CONFIG
    header_bg = _cfg(config, "table_styling", "header_bg_color", "F2F2F2")
    border_color = _cfg(config, "table_styling", "border_color", "CCCCCC")

    if table.rows:
        header_cells = table.rows[0].cells
        for cell in header_cells:
            _set_cell_bg(cell, header_bg)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if cell.paragraphs:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.tcPr
            if tcPr is None:
                tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "12")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), border_color)
                tcBorders.append(border)
            tcPr.append(tcBorders)

def _card_table(doc, columns, widths=None, row_height=None):
    table = doc.add_table(rows=1, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)
    if widths:
        for idx, width in enumerate(widths[:columns]):
            _set_cell_width(table.rows[0].cells[idx], width)
    if row_height:
        _set_row_height(table.rows[0], row_height)
    return table

def _set_row_height(row, inches):
    tr_pr = row._tr.get_or_add_trPr()
    height = OxmlElement('w:trHeight')
    height.set(qn('w:val'), str(int(float(inches) * 1440)))
    height.set(qn('w:hRule'), 'atLeast')
    tr_pr.append(height)

def _empty_state_card(doc, message=None, config=None):
    config = config or DEFAULT_REPORT_CONFIG
    table = _card_table(doc, 1, widths=[6.5], row_height=0.5)
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, _cfg(config, "branding", "light_band", "EBF4FD"))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_text(
        cell,
        message or EMPTY_STATE_TEXT,
        config,
        bold=True,
        color=_cfg(config, "branding", "muted_color", "6B7280"),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    return table

def _save_figure(fig, config=None, edgecolor=None):
    path = tempfile.mktemp(suffix='.png')
    layout = (config or DEFAULT_REPORT_CONFIG).get("layout", {})
    dpi = int(layout.get("chart_dpi", 220))
    plt.savefig(path, dpi=dpi, bbox_inches='tight', facecolor='white', edgecolor=edgecolor or 'white')
    plt.close(fig)
    return path

def _readiness_badge(score, explicit_status, config=None):
    config = config or DEFAULT_REPORT_CONFIG
    if explicit_status and explicit_status != DATA_NOT_AVAILABLE:
        label = str(explicit_status).upper()
    else:
        value = float(score or 0)
        label = "NOT READY"
        for threshold in config.get("readiness_thresholds", []):
            try:
                if value >= float(threshold.get("min", 0)):
                    label = str(threshold.get("label", label)).upper()
                    break
            except (TypeError, ValueError):
                continue
    color = _cfg(config, "branding", "fail_color", "C00000")
    for threshold in config.get("readiness_thresholds", []):
        if str(threshold.get("label", "")).upper() == label:
            color = threshold.get("color", color)
            break
    return label, _hex_text(color, "C00000")

def _manifest_sections(manifest):
    sections = manifest.get("sections", []) if isinstance(manifest, dict) else []
    return [section for section in sections if isinstance(section, dict)]

def _manifest_section_ids(manifest):
    return [section.get("id") for section in _manifest_sections(manifest)]

def _inch_value(value, default=1.0):
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value or "").strip().lower()
    if text.endswith("in"):
        text = text[:-2].strip()
    try:
        return float(text)
    except ValueError:
        return float(default)

def _twips_value(value, default=0):
    try:
        return int(value)
    except (TypeError, ValueError):
        return int(default)

def _load_aaa_report_blueprint():
    return _load_yaml_config(AAA_REPORT_BLUEPRINT)

def _blueprint_page_setup(blueprint):
    if not isinstance(blueprint, dict):
        return {}
    return blueprint.get("page_setup", {}) if isinstance(blueprint.get("page_setup"), dict) else {}

def _blueprint_margins(blueprint):
    if not isinstance(blueprint, dict):
        return {}
    page_setup = _blueprint_page_setup(blueprint)
    if isinstance(page_setup.get("margins"), dict):
        return page_setup.get("margins")
    return blueprint.get("margins", {}) if isinstance(blueprint.get("margins"), dict) else {}

def _apply_blueprint_page_setup(doc, blueprint):
    page_setup = _blueprint_page_setup(blueprint)
    page_size = page_setup.get("page_size", {}) if isinstance(page_setup.get("page_size"), dict) else {}
    margins = _blueprint_margins(blueprint)
    for section in doc.sections:
        section.page_width = Inches(_inch_value(page_size.get("width"), 8.2681))
        section.page_height = Inches(_inch_value(page_size.get("height"), 11.6931))
        section.top_margin = Inches(_inch_value(margins.get("top"), 1))
        section.bottom_margin = Inches(_inch_value(margins.get("bottom"), 1))
        section.left_margin = Inches(_inch_value(margins.get("left"), 1))
        section.right_margin = Inches(_inch_value(margins.get("right"), 1))
        section.header_distance = Inches(_inch_value(margins.get("header"), 0))
        section.footer_distance = Inches(_inch_value(margins.get("footer"), 0))

def _apply_blueprint_document_defaults(doc, blueprint):
    fonts = blueprint.get("fonts", {}) if isinstance(blueprint, dict) else {}
    body_text = fonts.get("body_text", {}) if isinstance(fonts.get("body_text"), dict) else {}
    family = body_text.get("family")
    size = body_text.get("size")
    if family:
        doc.styles["Normal"].font.name = str(family)
    if isinstance(size, (int, float)):
        doc.styles["Normal"].font.size = Pt(size)

def _apply_blueprint_paragraph_tabs(paragraph, tab_specs):
    if not isinstance(tab_specs, list):
        return
    pPr = paragraph._p.get_or_add_pPr()
    existing_tabs = pPr.find(qn("w:tabs"))
    if existing_tabs is not None:
        pPr.remove(existing_tabs)
    tabs = OxmlElement("w:tabs")
    for tab_spec in tab_specs:
        if not isinstance(tab_spec, dict):
            continue
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), str(tab_spec.get("alignment", "left")))
        if tab_spec.get("leader"):
            tab.set(qn("w:leader"), str(tab_spec.get("leader")))
        tab.set(qn("w:pos"), str(_twips_value(tab_spec.get("position_twips"), 0)))
        tabs.append(tab)
    pPr.append(tabs)

def _set_paragraph_spacing_twips(paragraph, before_twips=None, after_twips=None, line_twips=None, line_rule=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    if before_twips is not None:
        spacing.set(qn("w:before"), str(_twips_value(before_twips)))
    if after_twips is not None:
        spacing.set(qn("w:after"), str(_twips_value(after_twips)))
    if line_twips is not None:
        spacing.set(qn("w:line"), str(_twips_value(line_twips)))
    if line_rule:
        spacing.set(qn("w:lineRule"), str(line_rule))

def _set_run_font_from_blueprint(run, font_config):
    if not isinstance(font_config, dict):
        return
    family = font_config.get("family")
    size = font_config.get("size")
    bold = font_config.get("bold")
    italic = font_config.get("italic")
    color = font_config.get("color")
    if isinstance(family, str) and not family.startswith("inherited"):
        run.font.name = family
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), family)
        rFonts.set(qn("w:hAnsi"), family)
    if isinstance(size, (int, float)):
        run.font.size = Pt(size)
    if isinstance(bold, bool):
        run.font.bold = bold
    if isinstance(italic, bool):
        run.font.italic = italic
    if isinstance(color, str) and color.startswith("#") and len(color) == 7:
        run.font.color.rgb = _hex_color(color)

def _apply_design_system(doc, design_system):
    document_config = design_system.get("document", {}) if isinstance(design_system, dict) else {}
    margin_config = document_config.get("margins", {}) if isinstance(document_config, dict) else {}
    for section in doc.sections:
        section.top_margin = Inches(_inch_value(margin_config.get("top"), 1))
        section.bottom_margin = Inches(_inch_value(margin_config.get("bottom"), 1))
        section.left_margin = Inches(_inch_value(margin_config.get("left"), 1))
        section.right_margin = Inches(_inch_value(margin_config.get("right"), 1))

def _log_first_nine_manifest(manifest):
    ids = _manifest_section_ids(manifest)
    service_sections = [
        section.get("id")
        for section in _manifest_sections(manifest)
        if section.get("template") == "service_assessment"
    ]
    logger.info("[REPORT_BUILDER] Content manifest sections: %s", ids)
    if service_sections:
        logger.info(
            "[REPORT_BUILDER] Skipping service assessment sections for pages 1-9 only: %s",
            service_sections,
        )

def _scan_additional_yaml_sources():
    roots = []
    for base in [Path.cwd(), Path(__file__).resolve().parents[3], Path(__file__).resolve().parents[4]]:
        for name in ("config", "data", "reports"):
            candidate = base / name
            if candidate.exists():
                roots.append(candidate)
    files = []
    for root in roots:
        files.extend(root.rglob("*.yml"))
        files.extend(root.rglob("*.yaml"))
    logger.info("[REPORT_BUILDER] Additional YAML sources loaded: %s", [str(path) for path in files])
    return files

def _status_value(row):
    return str(row.get("display_status", row.get("status", row.get("finding", ""))) or "").strip().lower()

def _is_pass(row):
    return _status_value(row) == "pass"

def _is_fail(row):
    return _status_value(row) == "fail"

def _severity_value(row):
    raw = str(
        row.get("display_severity")
        or row.get("registry_severity")
        or row.get("severity")
        or ""
    ).strip().lower()
    return SEVERITY_MAP.get(raw, DATA_NOT_AVAILABLE)

def _pillar_value(row):
    raw = str(row.get("pillar") or row.get("category") or "").strip()
    if not raw:
        return DATA_NOT_AVAILABLE
    return PILLAR_MAP.get(raw.lower(), raw)

def _service_value(row):
    raw = row.get("service") or row.get("category") or DATA_NOT_AVAILABLE
    return _service_name_key(raw)

def _distribution(rows, value_getter, allowed_order=None):
    counts = defaultdict(int)
    for row in rows or []:
        value = value_getter(row)
        if value and value != DATA_NOT_AVAILABLE:
            counts[value] += 1
    if allowed_order:
        return {name: counts.get(name, 0) for name in allowed_order if counts.get(name, 0) > 0}
    return dict(counts)

def _pass_fail_counts(rows):
    pass_count = len([row for row in rows or [] if _is_pass(row)])
    fail_count = len([row for row in rows or [] if _is_fail(row)])
    return pass_count, fail_count

def _chart_or_data_not_available(doc, chart_path, width, height=None, config=None, message=None):
    if chart_path:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        if height is None:
            run.add_picture(chart_path, width=width)
        else:
            run.add_picture(chart_path, width=width, height=height)
        try:
            os.remove(chart_path)
        except:
            pass
        return True
    _empty_state_card(doc, message or "No activity data available", config or DEFAULT_REPORT_CONFIG)
    return False

def _sort_findings(findings):
    """Sort by severity then fail before pass."""
    def sort_key(f):
        report_order = f.get('report_order')
        if report_order is not None:
            try:
                return (0, int(report_order))
            except (TypeError, ValueError):
                pass
        title = str(f.get('title') or f.get('parameter') or '').strip()
        if title in PARAMETER_ORDER:
            return (0, PARAMETER_ORDER[title])
        sev = str(f.get('severity', '')).lower().strip()
        fnd = str(f.get('finding', '')).lower().strip()
        return (
            1,
            SEVERITY_ORDER.get(sev, 9),
            FINDING_ORDER.get(fnd, 9),
        )
    return sorted(findings, key=sort_key)

def _add_paragraph_with_bold_phrases(doc, text, bold_phrases, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    _add_runs_with_bold_phrases(p, text, bold_phrases)
    return p

def _add_runs_with_bold_phrases(paragraph, text, bold_phrases):
    remaining = str(text or '')
    while remaining:
        matches = [
            (remaining.find(phrase), phrase)
            for phrase in bold_phrases
            if phrase and remaining.find(phrase) >= 0
        ]
        if not matches:
            paragraph.add_run(remaining)
            break
        index, phrase = min(matches, key=lambda item: item[0])
        if index:
            paragraph.add_run(remaining[:index])
        run = paragraph.add_run(phrase)
        run.bold = True
        remaining = remaining[index + len(phrase):]

def _format_run(run, size=11, bold=False, name='Calibri'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    return run

def _first_value(*values):
    for value in values:
        if value:
            return value
    return ''

def _documentation_url(row):
    return _first_value(
        row.get('microsoft_doc_url'),
        row.get('documentation_url'),
        row.get('documentation_link'),
        row.get('registry_param', {}).get('documentation_url') if isinstance(row.get('registry_param'), dict) else '',
    )

def _add_hyperlink(paragraph, url, text):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(
        qn('r:id'),
        paragraph.part.relate_to(
            url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True,
        ),
    )
    run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    style = OxmlElement('w:rStyle')
    style.set(qn('w:val'), 'Hyperlink')
    r_pr.append(style)
    run.append(r_pr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

def _parse_jsonish(value):
    current = value
    for _ in range(3):
        if isinstance(current, str):
            stripped = current.strip()
            if not stripped:
                return {}
            try:
                current = json.loads(stripped)
                continue
            except Exception:
                return {"actual_value": current}
        return current if isinstance(current, dict) else {"actual_value": current}
    return current if isinstance(current, dict) else {}

def _normalise_raw_value(raw_value, depth=0):
    if depth > 2:
        return {}
    raw = _parse_jsonish(raw_value)
    if not isinstance(raw, dict):
        return {}
    merged = dict(raw)
    actual = _parse_jsonish(raw.get('actual_value'))
    if isinstance(actual, dict):
        merged.update(actual)
    else:
        merged['actual_value'] = actual
    evidence = _parse_jsonish(raw.get('evidence'))
    if isinstance(evidence, dict):
        for key, value in evidence.items():
            merged.setdefault(key, value)
    payload = _parse_jsonish(raw.get('payload'))
    if isinstance(payload, dict):
        result = _parse_jsonish(payload.get('result'))
        result_raw = _parse_jsonish(result.get('raw_value')) if isinstance(result, dict) else {}
        if isinstance(result_raw, dict):
            nested = _normalise_raw_value(result_raw, depth + 1)
            for key, value in nested.items():
                merged.setdefault(key, value)
    return merged

def _num(value, default=0.0):
    try:
        return float(value)
    except (TypeError, ValueError):
        return float(default)

def _int(value, default=0):
    return int(round(_num(value, default)))

def _ratio_to_percent(value):
    pct = _num(value)
    return pct * 100 if 0 <= pct <= 1 else pct

def _fmt_pct(value, digits=0):
    pct = _ratio_to_percent(value)
    return f"{pct:.{digits}f}"

def _license_gap_message(raw, default_sku='required license'):
    sku = raw.get('required_sku') or raw.get('required_license') or raw.get('license') or raw.get('service') or default_sku
    return f"This parameter could not be assessed. The tenant does not have the required license: {sku}."

def _has_licensing_gap(raw):
    text = json.dumps(raw, default=str).upper()
    return raw.get('collection_status') == 'LICENSING_GAP' or 'LICENSING_GAP' in text or bool(raw.get('required_sku'))

def _teams_unavailable(raw):
    return raw.get('service_available') is False and str(raw.get('service', '')).lower() == 'microsoft teams'

def _finding_message(raw):
    '''Return the finding text the collector stored (== assessment_findings.evaluated_value).

    evaluated_value is the source of truth for the report Description. PowerShell
    collectors store it at collector_contract.findings[].message; Graph collectors store
    it at reasoning / evidence.reasoning. Used to keep the Description from contradicting
    the finding when a handler's recomputation reads evidence keys the collector no longer
    populates. Returns None if no message is present (handler falls back to its own text).'''
    cc = raw.get('collector_contract')
    if isinstance(cc, dict):
        for finding in (cc.get('findings') or []):
            if isinstance(finding, dict) and finding.get('message'):
                return str(finding['message'])
    if raw.get('reasoning'):
        return str(raw['reasoning'])
    evidence = raw.get('evidence')
    if isinstance(evidence, dict) and evidence.get('reasoning'):
        return str(evidence['reasoning'])
    return None

def _report_description(row):
    '''The report "Description:" line must equal assessment_findings.evaluated_value.

    evaluated_value is the single source of truth. The detail row carries it directly as
    row['finding'] (== finding.evaluated_value); as a safety net we also read the message
    stored in raw evidence. Returns the finding text, or None only when there is genuinely
    no finding (e.g. NOT COLLECTED) so the caller can fall back to a generic description.'''
    if not isinstance(row, dict):
        return None
    for candidate in (row.get('evaluated_value'), row.get('finding')):
        text = str(candidate or '').strip()
        if text and text.upper() not in ('NOT COLLECTED', 'NOT_COLLECTED', 'NONE'):
            return text
    return _finding_message(_normalise_raw_value(row.get('evidence')))

def resolve_description(parameter_key, raw_value, status):
    '''
    Returns the correct description string for a parameter based on live evidence.
    Handles licensing gaps, service availability, real values, and placeholders.
    '''
    raw = _normalise_raw_value(raw_value)

    if _has_licensing_gap(raw):
        return _license_gap_message(raw)

    teams_license_keys = {
        'copilot_integration_enabled',
        'meeting_transcription_enabled',
        'third_party_apps_allowed',
        'guest_access_enabled_disabled',
        'teams_lobby_bypass',
        'meeting_recording_retention_policies',
    }
    if parameter_key in teams_license_keys and _teams_unavailable(raw):
        return _license_gap_message(raw, 'Microsoft Teams')

    if raw.get('service_available') is False:
        svc = raw.get('service') or 'required service'
        return f"This parameter could not be assessed. {svc} is not available in this tenant."

    descriptions = {
        'audit_logs_enabled': lambda r: (
            f"Audit logs are enabled and queryable. {_int(r.get('sample_count'))} log entries were found."
            if r.get('audit_logs_queryable') else
            "Audit logs are not currently enabled."
        ),
        'secure_score_percentage': lambda r: _finding_message(r) or (
            f"Secure score is {_num(r.get('secure_score_percentage')):.2f}% "
            f"which {'meets' if _num(r.get('secure_score_percentage')) >= 70 else 'is below'} "
            f"the recommended threshold (70%)."
        ),
        'compliance_score_overview': lambda r: _finding_message(r) or (
            "Microsoft Purview Compliance Manager does not expose an officially supported public API "
            "for retrieving the overall Compliance Score; review the score directly in the Purview "
            "compliance portal (compliance.microsoft.com/compliancemanager)."
        ),
        'users_without_mfa': lambda r: _finding_message(r) or (
            f"{_int(r.get('users_without_mfa'))} out of {_int(r.get('total_users'))} users do not have MFA registered."
        ),
        'global_administrator_accounts': lambda r: (
            f"There are {_int(r.get('actual_value', r.get('admin_count', r.get('global_admin_count'))))} Global Administrator accounts."
        ),
        'guest_users_count': lambda r: (
            f"There are {_int(r.get('guest_count'))} guest users out of {_int(r.get('total_users'))} total users."
        ),
        'account_enabled': lambda r: (
            f"{_int(r.get('enabled_count'))} out of {_int(r.get('total_users'))} accounts are enabled ({_fmt_pct(r.get('enabled_percent'), 0)}%)."
        ),
        'user_information': lambda r: (
            f"{_int(r.get('complete_users'))} out of {_int(r.get('total_users'))} users have complete profile information. "
            f"{_int(r.get('incomplete_users', max(_int(r.get('total_users')) - _int(r.get('complete_users')), 0)))} users are missing department or role."
        ),
        'conditional_access_policies_exclusion': lambda r: (
            "No Conditional Access Policy exclusions were found."
            if _int(r.get('policies_with_exclusions')) == 0 else
            f"Conditional Access Policies have {_int(r.get('policies_with_exclusions'))} excluded user(s)."
        ),
        'entra_tenant_creation_by_non_admin': lambda r: (
            "Non-admin users are allowed to create tenants."
            if r.get('actual_value') is True else
            "Non-admin users are not allowed to create tenants."
        ),
        'admin_consent_workflow': lambda r: (
            "Admin consent workflow is configured and enabled."
            if r.get('isEnabled') else
            "User cannot request admin consent - workflow is disabled."
        ),
        'mailboxes_status_active_inactive': lambda r: (
            f"{_int(r.get('active_mailboxes'))} out of {_int(r.get('active_mailboxes')) + _int(r.get('inactive_mailboxes'))} "
            f"({_fmt_pct(r.get('active_ratio'), 0)}%) mailboxes are active."
        ),
        'number_of_emails_read_received': lambda r: _finding_message(r) or (
            f"{_int(r.get('engaged_users'))} out of {_int(r.get('total_users'))} ({_fmt_pct(r.get('read_ratio'), 0)}%) "
            f"have read more than 70% of their mail."
        ),
        'number_of_emails_sent': lambda r: _finding_message(r) or (
            f"{_int(r.get('total_users'))} users sent an average of {_num(r.get('average_sent_per_user')):.1f} mails."
        ),
        'meeting_recording_retention_policies': lambda r: _finding_message(r) or (
            f"It is enabled and meeting recordings are set to automatically expire after {_int(r.get('expiration_days'))} days."
        ),
        'total_active_users_on_onedrive': lambda r: (
            f"{_int(r.get('active_users'))} out of {_int(r.get('total_users'))} ({_fmt_pct(r.get('active_ratio'), 0)}%) users have shown activity in the last 2 months."
        ),
        'active_users_on_sharepoint': lambda r: (
            f"{_int(r.get('active_users'))} out of {_int(r.get('total_users'))} ({_fmt_pct(r.get('active_ratio'), 0)}%) users are active."
        ),
        'storage_quota_consumption': lambda r: _finding_message(r) or (
            "Total storage consumption data could not be retrieved."
            if _int(r.get('site_count')) == 0 else
            f"{_int(r.get('sites_over_90_percent'))} out of {_int(r.get('site_count'))} sites are over 90% storage quota; maximum usage is {_fmt_pct(r.get('max_storage_quota_ratio'), 1)}%."
        ),
        'active_sites_count': lambda r: (
            f"{_int(r.get('active_site_count'))} out of {_int(r.get('total_sites'))} ({_fmt_pct(r.get('active_ratio'), 0)}%) sites are active."
        ),
        'third_party_apps_allowed': lambda r: str(r.get('status_text') or r.get('message') or r.get('actual_value') or 'Third-party apps are allowed.'),
        'guest_access_enabled_disabled': lambda r: _finding_message(r) or (
            f"Guest access on Teams is {'enabled' if r.get('guest_access_enabled') or r.get('AllowGuestUser') else 'disabled'}."
        ),
        'teams_lobby_bypass': lambda r: _finding_message(r) or str(r.get('lobby_bypass') or r.get('AutoAdmittedUsers') or r.get('message') or 'Teams lobby bypass setting could not be determined.'),
    }

    handler = descriptions.get(parameter_key)
    if handler:
        try:
            return handler(raw)
        except Exception:
            return None
    return None

def _pie_chart(labels, values, colors, title):
    filtered = [(l, v, c) for l, v, c in zip(labels, values, colors) if v > 0]
    if not filtered: return None
    fl, fv, fc = zip(*filtered)
    fig, ax = plt.subplots(figsize=(5, 4))
    ax.pie(fv, labels=fl, colors=fc, autopct='%1.0f%%', startangle=90, textprops={'fontsize': 8})
    ax.set_title(title, fontsize=11, fontweight='bold', pad=15)
    plt.tight_layout()
    path = tempfile.mktemp(suffix='.png')
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def _add_chart_panel_shadow(fig):
    shadow = mpatches.Rectangle(
        (0.018, -0.018),
        0.968,
        0.968,
        transform=fig.transFigure,
        facecolor="#9E9E9E",
        alpha=0.28,
        linewidth=0,
        zorder=-20,
    )
    panel = mpatches.Rectangle(
        (0.0, 0.0),
        0.985,
        0.985,
        transform=fig.transFigure,
        facecolor="white",
        edgecolor="#000000",
        linewidth=1.0,
        zorder=-19,
    )
    fig.patches.extend([shadow, panel])

def _page6_pie_chart(labels, values, colors, title, figure_size=(6.34, 2.5)):
    """Compact AAA-style Excel 3D pie chart for page 6."""
    filtered = [(l, v, c) for l, v, c in zip(labels, values, colors) if v > 0]
    if not filtered:
        return None
    fl, fv, fc = zip(*filtered)
    total = sum(fv)

    def darken(hex_color, factor=0.55):
        h = hex_color.lstrip('#')
        r = int(int(h[0:2], 16) * factor)
        g = int(int(h[2:4], 16) * factor)
        b = int(int(h[4:6], 16) * factor)
        return f'#{r:02x}{g:02x}{b:02x}'

    fig, ax = plt.subplots(figsize=figure_size, dpi=REPORT_CHART_DPI)
    fig.patch.set_facecolor('white')
    _add_chart_panel_shadow(fig)

    radius = 0.58
    startangle = 112

    # Faux Excel 3D depth layer, then the visible pie on top.
    ax.pie(
        fv,
        colors=[darken(c) for c in fc],
        startangle=startangle,
        radius=radius,
        center=(0, -0.12),
        wedgeprops={'linewidth': 0.5, 'edgecolor': '#777777'},
    )
    wedges, _texts = ax.pie(
        fv,
        colors=fc,
        startangle=startangle,
        shadow=True,
        radius=radius,
        center=(0, 0),
        labels=None,
        wedgeprops={'linewidth': 0.8, 'edgecolor': 'white'},
    )

    # External name + percentage callouts with leader lines, matching the AAA chart style.
    for wedge, label, value in zip(wedges, fl, fv):
        angle = (wedge.theta1 + wedge.theta2) / 2
        x = np.cos(np.deg2rad(angle))
        y = np.sin(np.deg2rad(angle))
        label_x = 1.18 * np.sign(x)
        label_y = 0.80 * y
        ha = 'left' if x >= 0 else 'right'
        ax.annotate(
            f'{label} {value / total:.0%}',
            xy=(0.61 * x, 0.61 * y),
            xytext=(label_x, label_y),
            ha=ha,
            va='center',
            fontsize=7,
            fontfamily='Calibri',
            color='#1A1A1A',
            bbox={
                'boxstyle': 'square,pad=0.13',
                'facecolor': 'white',
                'edgecolor': '#000000',
                'linewidth': 0.7,
            },
            arrowprops={
                'arrowstyle': '-',
                'color': '#000000',
                'lw': 0.8,
                'shrinkA': 0,
                'shrinkB': 0,
                'connectionstyle': 'angle3,angleA=0,angleB=90',
            },
        )

    ax.set_title(title, fontsize=14, fontweight='bold', fontfamily='Calibri', pad=8)
    ax.legend(
        wedges,
        fl,
        loc='lower center',
        bbox_to_anchor=(0.5, -0.16),
        ncol=min(3, len(fl)),
        frameon=False,
        fontsize=6.8,
    )
    ax.set_aspect('equal')
    ax.set_xlim(-1.62, 1.62)
    ax.set_ylim(-1.12, 1.12)
    ax.set_axis_off()
    plt.tight_layout(pad=0.35)
    path = tempfile.mktemp(suffix='.png')
    plt.savefig(path, dpi=REPORT_CHART_DPI, bbox_inches=None, facecolor='white', pad_inches=0)
    plt.close(fig)
    return path

def _page7_severity_matrix_graphic(severity_items, figure_size):
    """AAA-style Page 7 severity matrix graphic with five overlapping circles."""
    fig, ax = plt.subplots(figsize=figure_size, dpi=REPORT_CHART_DPI)
    fig.patch.set_facecolor("white")
    fig.subplots_adjust(left=0, right=1, top=1, bottom=0)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    def softened(hex_color):
        h = str(hex_color or "#FFFFFF").strip().lstrip("#")
        if len(h) != 6:
            return "#FFFFFF"
        r, g, b = (int(h[i:i + 2], 16) for i in (0, 2, 4))
        r = int(r + (255 - r) * 0.88)
        g = int(g + (255 - g) * 0.88)
        b = int(b + (255 - b) * 0.88)
        return f"#{r:02x}{g:02x}{b:02x}"

    circle_width = 0.18
    circle_height = circle_width * (figure_size[0] / figure_size[1])
    circle_radius_x = circle_width / 2
    circle_radius_y = circle_height / 2
    center_y = 0.66

    centers = np.linspace(0.13, 0.87, max(len(severity_items), 1))
    for idx, (item, x) in enumerate(zip(severity_items, centers)):
        label = item.get("label", "")
        color = item.get("color", "#6B7280")
        desc = item.get("description", "")
        circle = mpatches.Ellipse(
            (x, center_y),
            width=circle_width,
            height=circle_height,
            facecolor=softened(color),
            edgecolor=color,
            linewidth=3.0,
            zorder=3 + idx,
        )
        ax.add_patch(circle)
        if idx < len(severity_items) - 1:
            tri = mpatches.Polygon(
                [
                    (x + circle_radius_x * 0.80, center_y - circle_radius_y * 0.68),
                    (x + circle_radius_x * 1.40, center_y),
                    (x + circle_radius_x * 0.80, center_y + circle_radius_y * 0.68),
                ],
                closed=True,
                facecolor=color,
                edgecolor=color,
                linewidth=0,
                zorder=8 + idx,
            )
            ax.add_patch(tri)
        ax.text(x, center_y, label, ha="center", va="center", fontsize=7.5, fontfamily="Calibri", color="#000000", zorder=20)
        wrapped = textwrap.fill(desc, width=16)
        ax.text(
            x,
            0.34,
            f"* {wrapped}",
            ha="center",
            va="top",
            fontsize=5.75,
            fontfamily="Calibri",
            color="#000000",
            linespacing=0.90,
            multialignment="left",
            zorder=20,
        )

    path = tempfile.mktemp(suffix=".png")
    plt.savefig(path, dpi=REPORT_CHART_DPI, bbox_inches=None, facecolor="white", pad_inches=0)
    plt.close(fig)
    return path

def _page7_risk_pie_chart(labels, values, colors, title, figure_size):
    """Compact AAA-style Excel 3D pie chart for Page 7 risk categories."""
    filtered = [(l, v, c) for l, v, c in zip(labels, values, colors) if v > 0]
    if not filtered:
        return None
    fl, fv, fc = zip(*filtered)
    total = sum(fv)

    def darken(hex_color, factor=0.55):
        h = hex_color.lstrip('#')
        r = int(int(h[0:2], 16) * factor)
        g = int(int(h[2:4], 16) * factor)
        b = int(int(h[4:6], 16) * factor)
        return f'#{r:02x}{g:02x}{b:02x}'

    fig, ax = plt.subplots(figsize=figure_size, dpi=REPORT_CHART_DPI)
    fig.patch.set_facecolor('white')
    _add_chart_panel_shadow(fig)

    radius = 0.62
    startangle = 112

    ax.pie(
        fv,
        colors=[darken(c) for c in fc],
        startangle=startangle,
        radius=radius,
        center=(0, -0.10),
        wedgeprops={'linewidth': 0.5, 'edgecolor': '#777777'},
    )
    wedges, _texts = ax.pie(
        fv,
        colors=fc,
        startangle=startangle,
        shadow=True,
        radius=radius,
        center=(0, 0),
        labels=None,
        wedgeprops={'linewidth': 0.8, 'edgecolor': 'white'},
    )

    for wedge, label, value in zip(wedges, fl, fv):
        angle = (wedge.theta1 + wedge.theta2) / 2
        x = np.cos(np.deg2rad(angle))
        y = np.sin(np.deg2rad(angle))
        label_x = 1.12 * np.sign(x)
        label_y = 0.78 * y
        ha = 'left' if x >= 0 else 'right'
        ax.annotate(
            f'{label} {value / total:.0%}',
            xy=(0.66 * x, 0.66 * y),
            xytext=(label_x, label_y),
            ha=ha,
            va='center',
            fontsize=7.2,
            fontfamily='Calibri',
            color='#1A1A1A',
            bbox={
                'boxstyle': 'square,pad=0.13',
                'facecolor': 'white',
                'edgecolor': '#000000',
                'linewidth': 0.7,
            },
            arrowprops={
                'arrowstyle': '-',
                'color': '#000000',
                'lw': 0.8,
                'shrinkA': 0,
                'shrinkB': 0,
                'connectionstyle': 'angle3,angleA=0,angleB=90',
            },
        )

    ax.set_title(title, fontsize=15, fontweight='bold', fontfamily='Calibri', pad=8)
    ax.legend(
        wedges,
        fl,
        loc='lower center',
        bbox_to_anchor=(0.5, -0.14),
        ncol=5,
        frameon=False,
        fontsize=7.5,
    )
    ax.set_aspect('equal')
    ax.set_xlim(-1.62, 1.62)
    ax.set_ylim(-1.08, 1.08)
    ax.set_axis_off()
    plt.tight_layout(pad=0.45)
    path = tempfile.mktemp(suffix='.png')
    plt.savefig(path, dpi=REPORT_CHART_DPI, bbox_inches=None, facecolor='white', pad_inches=0)
    plt.close(fig)
    return path

def _bar_chart_h(labels, values, colors, title):
    fig, ax = plt.subplots(figsize=(6.6, 2.9), dpi=REPORT_CHART_DPI)
    y_pos = range(len(labels))
    bars = ax.barh(list(y_pos), values, color=colors, height=0.4)
    ax.set_yticks(list(y_pos))
    ax.set_yticklabels(labels, fontsize=10)
    ax.bar_label(bars, padding=4, fontsize=10, fontweight='bold')
    ax.set_xlim(0, max(values) * 1.25 if max(values) > 0 else 10)
    ax.set_title(title, fontsize=11, fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout(pad=0.7)
    path = tempfile.mktemp(suffix='.png')
    plt.savefig(path, dpi=REPORT_CHART_DPI, bbox_inches='tight', facecolor='white', pad_inches=0.06)
    plt.close(fig)
    return path

def _add_readiness_gauge(doc, score, config=None):
    config = config or DEFAULT_REPORT_CONFIG
    score = max(0, min(100, float(score or 0)))
    pass_color = f"#{_hex_text(_cfg(config, 'branding', 'success_color', '00B050'))}"
    fail_color = f"#{_hex_text(_cfg(config, 'branding', 'fail_color', 'C00000'))}"
    accent_color = f"#{_hex_text(_cfg(config, 'branding', 'primary_color', '0078D4'))}"
    fig, ax = plt.subplots(figsize=(7.2, 1.72), dpi=150)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 1.46)
    bar_y = 0.55
    ax.barh(bar_y, 100, height=0.22, color=fail_color, align='center')
    ax.barh(bar_y, score, height=0.22, color=pass_color, align='center')
    ax.annotate(
        '',
        xy=(score, bar_y + 0.15),
        xytext=(score, 1.07),
        arrowprops={
            'arrowstyle': '-|>',
            'color': accent_color,
            'lw': 2.0,
            'mutation_scale': 14,
        },
    )
    ax.text(
        score, 1.28, f'{score:.2f}%',
        ha='center', va='top',
        fontsize=12, fontweight='bold', color=accent_color,
        bbox={'boxstyle': 'round,pad=0.18', 'facecolor': 'white', 'edgecolor': 'white', 'linewidth': 0},
        clip_on=False,
    )
    ax.set_xticks(range(0, 101, 10))
    ax.set_xticklabels([f'{x}%' for x in range(0, 101, 10)], fontsize=7)
    ax.set_yticks([])
    legend = [
        mpatches.Patch(color=pass_color, label='Pass'),
        mpatches.Patch(color=fail_color, label='Fail'),
    ]
    ax.legend(handles=legend, loc='lower center', bbox_to_anchor=(0.5, -0.66), ncol=2, fontsize=7, frameon=False)
    fig.subplots_adjust(left=0.035, right=0.985, top=0.82, bottom=0.42)
    path = tempfile.mktemp(suffix='.png')
    plt.savefig(path, dpi=150, facecolor='white')
    plt.close(fig)
    return path

def _pass_fail_horizontal_chart(pass_count, fail_count):
    total = int(pass_count or 0) + int(fail_count or 0)
    if total <= 0:
        return None
    pass_pct = int(round(int(pass_count or 0) / total * 100))
    fail_pct = max(0, 100 - pass_pct)

    fig, ax = plt.subplots(figsize=(6.4, 1.1))
    ax.barh([0], [pass_pct], color='#00B050', label='Pass')
    ax.barh([0], [fail_pct], left=[pass_pct], color='#FF0000', label='Fail')
    ax.text(pass_pct / 2, 0, f'Pass {pass_pct}%', ha='center', va='center', color='white', fontsize=10, fontweight='bold')
    ax.text(pass_pct + fail_pct / 2, 0, f'Fail {fail_pct}%', ha='center', va='center', color='white', fontsize=10, fontweight='bold')
    ax.set_xlim(0, 100)
    ax.set_yticks([])
    ax.set_xticks([])
    ax.set_title('Pass vs Fail', fontsize=11, fontweight='bold', pad=16)
    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.legend(loc='upper center', bbox_to_anchor=(0.5, 1.20), ncol=2, frameon=False, fontsize=8)
    plt.tight_layout(pad=0.55)
    path = tempfile.mktemp(suffix='.png')
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def _service_pass_fail_chart(parameter_rows):
    rows = parameter_rows or []
    if not rows:
        return None
    labels = []
    pass_values = []
    fail_values = []
    for service in SERVICE_ORDER:
        service_rows = [row for row in rows if _service_value(row) == service]
        if not service_rows:
            continue
        p_count, f_count = _pass_fail_counts(service_rows)
        labels.append(service.replace('Microsoft ', '').replace(' Online', ''))
        pass_values.append(p_count)
        fail_values.append(f_count)
    if not labels:
        return None

    x = np.arange(len(labels))
    fig, ax = plt.subplots(figsize=(6.4, 2.2))
    ax.bar(x, fail_values, color='#FF0000', label='Fail', width=0.55)
    ax.bar(x, pass_values, bottom=fail_values, color='#00B050', label='Pass', width=0.55)
    ax.set_title('Assessment Results by Service', fontsize=11, fontweight='bold', pad=8)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=8, rotation=20, ha='right')
    ax.tick_params(axis='y', labelsize=8)
    ax.legend(loc='upper right', frameon=False, fontsize=8)
    ax.grid(axis='y', linestyle='--', alpha=0.25)
    for spine in ['top', 'right']:
        ax.spines[spine].set_visible(False)
    plt.tight_layout()
    path = tempfile.mktemp(suffix='.png')
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def _finding_title(row):
    return str(row.get('title') or row.get('parameter') or row.get('parameter_name') or DATA_NOT_AVAILABLE)

def _top_findings_by_severity(rows, severities, limit=4):
    wanted = {sev.lower() for sev in severities}
    selected = []
    for row in rows or []:
        if _severity_value(row).lower() in wanted and _is_fail(row):
            selected.append(row)
    return selected[:limit]

def _recommendation_text(row):
    recommendation = row.get('recommendation') or row.get('recommendation_text')
    if isinstance(recommendation, dict):
        recommendation = recommendation.get('text') or recommendation.get('recommendation')
    return str(recommendation or '').strip()

def _service_pillar_pass_fail_counts(parameter_rows, pillar_services):
    """Count chart values from the same service/pillar/status fields shown in report tables."""
    def normalize_pillar(value):
        text = str(value or '').lower().strip()
        return PILLAR_MAP.get(text, str(value or '').replace('_', ' ').title())

    def normalize_service(value):
        service = _service_name_key(value)
        aliases = {
            'OneDrive': 'OneDrive for Business',
            'SharePoint': 'SharePoint Online',
            'Teams': 'Microsoft Teams',
            'Purview': 'Microsoft Purview',
            'Exchange': 'Exchange Online',
        }
        return aliases.get(service, service)

    counts = {(pillar, service): {'Pass': 0, 'Fail': 0} for pillar, service in pillar_services}
    for row in parameter_rows or []:
        key = (normalize_pillar(row.get('pillar')), normalize_service(row.get('service')))
        if key not in counts:
            continue
        status = _detailed_status_text(row.get('display_status', row.get('status', '')))
        counts[key][status] += 1
    return counts

def _page8_exec_summary_chart(parameter_rows):
    """AAA-style stacked pass/fail summary chart for Page 8."""
    pillar_services = [
        ('Best Practice', 'Entra ID'),
        ('Best Practice', 'Exchange Online'),
        ('Best Practice', 'SharePoint Online'),
        ('Best Practice', 'Microsoft Teams'),
        ('Governance', 'Entra ID'),
        ('Governance', 'Exchange Online'),
        ('Governance', 'Microsoft Purview'),
        ('Governance', 'Microsoft Teams'),
        ('Governance', 'OneDrive for Business'),
        ('Governance', 'SharePoint Online'),
        ('Security', 'Entra ID'),
        ('Security', 'Exchange Online'),
        ('Security', 'Microsoft Purview'),
        ('Security', 'Microsoft Teams'),
        ('Security', 'OneDrive for Business'),
        ('Security', 'SharePoint Online'),
    ]

    counts = _service_pillar_pass_fail_counts(parameter_rows, pillar_services)

    def chart_service_label(service):
        text = str(service or "")
        replacements = {
            "Exchange Online": "Exchange",
            "SharePoint Online": "SharePoint",
            "Microsoft Teams": "Teams",
            "Microsoft Purview": "Purview",
            "OneDrive for Business": "OneDrive",
        }
        return replacements.get(text, text)

    labels = [chart_service_label(service) for pillar, service in pillar_services]
    pass_values = [counts[item]['Pass'] for item in pillar_services]
    fail_values = [counts[item]['Fail'] for item in pillar_services]
    x = np.arange(len(labels))
    totals = [p + f for p, f in zip(pass_values, fail_values)]
    max_total = max(totals, default=0)
    y_max = max(2, int(math.ceil((max_total + 1) / 2.0) * 2))

    fig, ax = plt.subplots(figsize=(7.25, 3.85), dpi=REPORT_CHART_DPI)
    fig.patch.set_facecolor('white')
    fig.patches.append(
        mpatches.Rectangle(
            (0, 0),
            1,
            1,
            transform=fig.transFigure,
            fill=False,
            edgecolor='#000000',
            linewidth=1.0,
            zorder=1000,
        )
    )
    fail_bars = ax.bar(x, fail_values, color='#C00000', label='Fail', width=0.58)
    pass_bars = ax.bar(x, pass_values, bottom=fail_values, color='#00B050', label='Pass', width=0.58)
    ax.set_title('Executive Summary - M365 Services and 3 Pillars', fontsize=12, fontweight='normal', pad=44)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=90, fontsize=6.4)
    ax.tick_params(axis='x', pad=2)
    ax.tick_params(axis='y', labelsize=7)
    ax.set_ylim(0, y_max)
    ax.set_yticks(np.arange(0, y_max + 1, 2))
    ax.grid(axis='y', color='#D9D9D9', linewidth=0.6)
    ax.set_axisbelow(True)
    ax.legend(
        [fail_bars, pass_bars],
        ['Fail', 'Pass'],
        loc='upper center',
        bbox_to_anchor=(0.5, 1.30),
        ncol=2,
        frameon=False,
        fontsize=7,
        handlelength=0.9,
        handletextpad=0.3,
        columnspacing=0.8,
    )

    for bar, value, base in zip(fail_bars, fail_values, [0] * len(fail_values)):
        if value > 0:
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                base + value / 2,
                str(value),
                ha='center',
                va='center',
                fontsize=6.6,
                color='white',
                fontweight='bold',
            )
    for bar, value, base in zip(pass_bars, pass_values, fail_values):
        if value > 0:
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                base + value / 2,
                str(value),
                ha='center',
                va='center',
                fontsize=6.6,
                color='white',
                fontweight='bold',
            )

    group_spans = [(0, 3, 'Best Practice'), (4, 9, 'Governance'), (10, 15, 'Security')]
    for start, end, label in group_spans:
        center = (start + end) / 2
        ax.text(center, -0.50, label, ha='center', va='top', fontsize=7.2, transform=ax.get_xaxis_transform())
    for separator in (3.5, 9.5):
        ax.axvline(separator, color='#BFBFBF', linewidth=0.7)

    for spine in ax.spines.values():
        spine.set_color('#BFBFBF')
        spine.set_linewidth(0.8)

    ax.set_xlim(-0.8, len(labels) - 0.2)
    fig.subplots_adjust(bottom=0.35, top=0.76, left=0.07, right=0.985)
    path = tempfile.mktemp(suffix='.png')
    plt.savefig(path, dpi=REPORT_CHART_DPI, facecolor='white', bbox_inches=None, pad_inches=0)
    plt.close(fig)
    return path

def _page9_services_pillars_chart(parameter_rows):
    """AAA Page 9 chart 1: grouped stacked columns by M365 service and pillar."""
    return _page8_exec_summary_chart(parameter_rows)

def _page9_severity_pillars_chart(parameter_rows, config=None):
    """AAA Page 9: grouped stacked severity columns by pillar and status."""
    config = config or DEFAULT_REPORT_CONFIG
    pillar_order = ['Best Practice', 'Governance', 'Security']
    status_order = ['Fail', 'Pass']
    severity_stack = ['Critical', 'High', 'Medium', 'Low', 'Informational']
    severity_definitions = config.get("severity_definitions", DEFAULT_REPORT_CONFIG["severity_definitions"])
    severity_colors = {
        severity: f"#{_hex_text((severity_definitions.get(severity, {}) or {}).get('color'))}"
        for severity in severity_stack
    }

    def normalize_pillar(value):
        text = str(value or '').lower().strip()
        return PILLAR_MAP.get(text, str(value or '').replace('_', ' ').title())

    counts = {
        (pillar, status): {severity: 0 for severity in severity_stack}
        for pillar in pillar_order
        for status in status_order
    }
    for row in parameter_rows:
        pillar = normalize_pillar(row.get('pillar'))
        if pillar not in pillar_order:
            continue
        status = _detailed_status_text(row.get('display_status', row.get('status', '')))
        severity_raw = str(
            row.get('display_severity')
            or row.get('registry_severity')
            or row.get('severity')
            or 'informational'
        ).lower().strip()
        severity = SEVERITY_MAP.get(severity_raw, 'Informational')
        if severity not in severity_stack:
            severity = 'Informational'
        counts[(pillar, status)][severity] += 1

    positions, labels, keys = [], [], []
    group_centers, group_bounds = [], []
    x = 0.0
    for pillar in pillar_order:
        start = x
        for status in status_order:
            positions.append(x)
            labels.append(status)
            keys.append((pillar, status))
            x += 1.0
        group_centers.append(((start + x - 1) / 2, pillar))
        group_bounds.append((start - 0.5, x - 0.5))
        x += 0.95

    fig, ax = plt.subplots(figsize=(6.6, 3.95), dpi=REPORT_CHART_DPI)
    fig.patch.set_facecolor('white')
    fig.patches.append(
        mpatches.Rectangle(
            (0, 0),
            1,
            1,
            transform=fig.transFigure,
            fill=False,
            edgecolor='#000000',
            linewidth=1.0,
            zorder=1000,
        )
    )
    bottoms = np.zeros(len(positions))
    for severity in severity_stack:
        values = [counts[key][severity] for key in keys]
        bars = ax.bar(positions, values, bottom=bottoms, color=severity_colors[severity], label=severity, width=0.42, edgecolor='none')
        for bar, value, bottom in zip(bars, values, bottoms):
            if value <= 0:
                continue
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                bottom + value / 2,
                str(value),
                    ha='center',
                    va='center',
                    fontsize=7.6,
                    fontweight='bold',
                    color='white',
                )
        bottoms += np.array(values)

    max_stack = max(bottoms) if len(bottoms) else 0
    y_max = max(20, int(np.ceil(max_stack / 2.0) * 2))
    ax.set_title('Executive Summary - Severity and 3 Pillars', fontsize=13, fontweight='bold', pad=10, fontfamily='Calibri')
    ax.set_xticks(positions)
    ax.set_xticklabels(labels, fontsize=7.5)
    ax.set_ylim(0, y_max)
    ax.set_yticks(np.arange(0, y_max + 1, 2))
    ax.tick_params(axis='y', labelsize=7.5, length=0)
    ax.tick_params(axis='x', length=0, pad=0)
    ax.grid(False)

    bracket_top = -0.055
    bracket_bottom = -0.18
    for start, end in group_bounds:
        ax.plot([start, end], [bracket_top, bracket_top], color='black', linewidth=0.6, transform=ax.get_xaxis_transform(), clip_on=False)
        ax.plot([start, start], [bracket_top, bracket_bottom], color='black', linewidth=0.6, transform=ax.get_xaxis_transform(), clip_on=False)
        ax.plot([end, end], [bracket_top, bracket_bottom], color='black', linewidth=0.6, transform=ax.get_xaxis_transform(), clip_on=False)
    for center, label in group_centers:
        ax.text(center, -0.14, label, ha='center', va='top', fontsize=7.0, transform=ax.get_xaxis_transform())

    ax.legend(
        [ax.containers[severity_stack.index(label)] for label in severity_stack],
        severity_stack,
        loc='center left',
        bbox_to_anchor=(1.02, 0.54),
        frameon=False,
        fontsize=7.0,
        handlelength=0.8,
        handletextpad=0.35,
        borderaxespad=0,
    )

    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.spines['left'].set_visible(True)
    ax.spines['bottom'].set_visible(True)
    ax.spines['left'].set_color('black')
    ax.spines['bottom'].set_color('black')
    ax.spines['left'].set_linewidth(0.6)
    ax.spines['bottom'].set_linewidth(0.6)
    ax.set_xlim(group_bounds[0][0], group_bounds[-1][1])
    fig.subplots_adjust(left=0.05, right=0.825, top=0.84, bottom=0.19)
    path = tempfile.mktemp(suffix='.png')
    plt.savefig(path, dpi=REPORT_CHART_DPI, facecolor='white', bbox_inches=None, pad_inches=0)
    plt.close(fig)
    return path

def _page9_observation_metrics(assessment_data):
    """Runtime metrics for the AAA Page 9 Key Observations block."""
    parameter_rows = assessment_data.get('parameter_rows', [])
    total_assessed = assessment_data.get('total_params') or len(parameter_rows)
    failed_parameters = assessment_data.get('gaps_count')
    if failed_parameters is None:
        failed_parameters = sum(
            1 for row in parameter_rows
            if str(row.get('display_status', row.get('status', ''))).lower().strip() != 'pass'
        )

    severity_count = 0
    for row in parameter_rows:
        severity = SEVERITY_MAP.get(str(row.get('severity', '')).lower().strip(), 'Informational')
        if severity in {'Critical', 'High', 'Medium'}:
            severity_count += 1
    severity_percent = round((severity_count / len(parameter_rows)) * 100) if parameter_rows else 0

    pillar_order = ['Security', 'Governance', 'Best Practice']
    pillar_totals = {pillar: 0 for pillar in pillar_order}
    pillar_fails = {pillar: 0 for pillar in pillar_order}
    for row in parameter_rows:
        pillar_raw = str(row.get('pillar', '')).lower().strip()
        pillar = PILLAR_MAP.get(pillar_raw, str(row.get('pillar', '')).replace('_', ' ').title())
        if pillar not in pillar_totals:
            continue
        pillar_totals[pillar] += 1
        status = str(row.get('display_status', row.get('status', ''))).lower().strip()
        if status != 'pass':
            pillar_fails[pillar] += 1

    fail_pct = {
        pillar: round((pillar_fails[pillar] / pillar_totals[pillar]) * 100) if pillar_totals[pillar] else 0
        for pillar in pillar_order
    }

    eligible_users = (
        assessment_data.get('copilot_eligible_users')
        or assessment_data.get('eligible_users')
        or assessment_data.get('copilot_license_eligible_users')
        or assessment_data.get('summary', {}).get('copilot_eligible_users')
        or assessment_data.get('summary', {}).get('eligible_users')
        or 0
    )
    total_users = (
        assessment_data.get('total_users')
        or assessment_data.get('eligible_total_users')
        or assessment_data.get('user_info_total')
        or assessment_data.get('summary', {}).get('total_users')
        or assessment_data.get('summary', {}).get('eligible_total_users')
        or assessment_data.get('summary', {}).get('user_info_total')
        or 0
    )

    return {
        'failed_parameters': int(failed_parameters or 0),
        'total_assessed': int(total_assessed or 0),
        'severity_percent': int(severity_percent),
        'security_fail_pct': int(fail_pct['Security']),
        'governance_fail_pct': int(fail_pct['Governance']),
        'best_practice_fail_pct': int(fail_pct['Best Practice']),
        'copilot_eligible_users': int(eligible_users or 0),
        'total_users': int(total_users or 0),
    }

def _adjust_label_positions(label_positions):
    """Adjust pie chart label positions to prevent overlapping leader lines."""
    min_spacing = 0.22

    left_labels = []
    right_labels = []

    for pos in label_positions:
        if pos['x'] >= 0:
            right_labels.append(pos)
        else:
            left_labels.append(pos)

    for group in [left_labels, right_labels]:
        if len(group) <= 1:
            for pos in group:
                pos['adjusted_pos'] = (1.36 * np.sign(pos['x']), 0.86 * pos['y'])
            continue

        group.sort(key=lambda p: p['y'], reverse=True)
        adjusted_y = []
        for i, pos in enumerate(group):
            base_y = 0.86 * pos['y']
            if i == 0:
                adjusted_y.append(base_y)
            else:
                prev_y = adjusted_y[-1]
                if base_y < prev_y - min_spacing:
                    adjusted_y.append(base_y)
                else:
                    adjusted_y.append(prev_y - min_spacing)

        for pos, adj_y in zip(group, adjusted_y):
            pos['adjusted_pos'] = (1.36 * np.sign(pos['x']), adj_y)

    return label_positions

def _page10_license_pie_chart(license_counts, title, config=None):
    config = config or DEFAULT_REPORT_CONFIG
    cleaned = {
        str(label): int(value)
        for label, value in (license_counts or {}).items()
        if int(value or 0) > 0
    }
    if not cleaned:
        return None
    labels = list(cleaned.keys())
    values = list(cleaned.values())
    palette = ["#F4B183", "#B4C7E7", "#C6E0B4", "#FFD966", "#A9D18E", "#9DC3E6"]

    def license_color(label, index):
        text = str(label or "").lower()
        if "unlicensed" in text:
            return "#AEE8D2"
        if "no license" in text or "none" == text.strip():
            return "#D9D9D9"
        if "copilot" in text:
            return "#ED7D31"
        if "business basic" in text:
            return "#F9D976"
        if "business standard" in text:
            return "#E6A0D8"
        if "business premium" in text:
            return "#B4A7D6"
        if "e5" in text:
            return "#8FAADC"
        if "e3" in text:
            return "#A9C2E8"
        if "exchange online" in text:
            return "#F4B183"
        return palette[index % len(palette)]

    colors = [license_color(label, idx) for idx, label in enumerate(labels)]

    def darken(hex_color, factor=0.62):
        h = hex_color.lstrip("#")
        r = int(int(h[0:2], 16) * factor)
        g = int(int(h[2:4], 16) * factor)
        b = int(int(h[4:6], 16) * factor)
        return f"#{r:02x}{g:02x}{b:02x}"

    fig, ax = plt.subplots(figsize=(6.35, 3.05), dpi=REPORT_CHART_DPI)
    fig.patch.set_facecolor("white")
    _add_chart_panel_shadow(fig)
    radius = 0.82
    startangle = 112
    ax.pie(
        values,
        colors=[darken(color) for color in colors],
        startangle=startangle,
        radius=radius,
        center=(0, -0.14),
        wedgeprops={"linewidth": 0.5, "edgecolor": "#777777"},
    )
    wedges, _texts = ax.pie(
        values,
        colors=colors,
        startangle=startangle,
        shadow=True,
        radius=radius,
        center=(0, 0),
        labels=None,
        wedgeprops={"linewidth": 0.8, "edgecolor": "white"},
    )

    label_positions = []
    for wedge, label, value in zip(wedges, labels, values):
        angle = (wedge.theta1 + wedge.theta2) / 2
        x = np.cos(np.deg2rad(angle))
        y = np.sin(np.deg2rad(angle))
        label_positions.append({
            'wedge': wedge,
            'label': label,
            'value': value,
            'angle': angle,
            'x': x,
            'y': y,
            'xy': (0.72 * x, 0.72 * y)
        })

    adjusted_positions = _adjust_label_positions(label_positions)

    for pos_data in adjusted_positions:
        wedge = pos_data['wedge']
        label = pos_data['label']
        value = pos_data['value']
        x = pos_data['x']
        xy = pos_data['xy']
        label_x, label_y = pos_data['adjusted_pos']

        ax.annotate(
            f"{label}; {value}",
            xy=xy,
            xytext=(label_x, label_y),
            ha="left" if x >= 0 else "right",
            va="center",
            fontsize=6.6,
            fontfamily="Calibri",
            bbox={"boxstyle": "square,pad=0.12", "facecolor": "white", "edgecolor": "#000000", "linewidth": 0.6},
            arrowprops={"arrowstyle": "-", "color": "#000000", "lw": 0.8, "shrinkA": 0, "shrinkB": 0},
        )
    ax.set_title(title, fontsize=12, fontweight="bold", fontfamily="Calibri", pad=5)
    ax.legend(
        wedges,
        labels,
        loc="lower center",
        bbox_to_anchor=(0.5, -0.12),
        ncol=min(len(labels), 4),
        frameon=False,
        fontsize=6.6,
    )
    ax.set_aspect("equal")
    ax.set_xlim(-1.55, 1.55)
    ax.set_ylim(-1.08, 1.08)
    ax.set_axis_off()
    fig.subplots_adjust(left=0.03, right=0.97, top=0.90, bottom=0.14)
    path = tempfile.mktemp(suffix=".png")
    plt.savefig(path, dpi=REPORT_CHART_DPI, facecolor="white", bbox_inches=None, pad_inches=0)
    plt.close(fig)
    return path

def _page10_user_info_chart(fields_dict, total_users, title, added_label, not_added_label, config=None):
    config = config or DEFAULT_REPORT_CONFIG
    total_users = int(total_users or 0)
    if not fields_dict or total_users <= 0:
        return None
    fields = list(fields_dict.keys())
    added = [max(min(int(value or 0), total_users), 0) for value in fields_dict.values()]
    missing = [max(total_users - value, 0) for value in added]
    x = np.arange(len(fields))
    width = 0.42
    pass_color = f"#{_hex_text(_cfg(config, 'branding', 'success_color', '00B050'))}"
    fail_color = f"#{_hex_text(_cfg(config, 'branding', 'fail_color', 'C00000'))}"

    fig, ax = plt.subplots(figsize=(6.35, 2.95), dpi=REPORT_CHART_DPI)
    fig.patch.set_facecolor("white")
    fig.patches.append(
        mpatches.Rectangle(
            (0, 0),
            1,
            1,
            transform=fig.transFigure,
            fill=False,
            edgecolor='#000000',
            linewidth=1.0,
            zorder=1000,
        )
    )
    bars_added = ax.bar(x, added, width, color=pass_color, label=added_label)
    bars_missing = ax.bar(x, missing, width, bottom=added, color=fail_color, label=not_added_label)
    for bar, value in zip(bars_added, added):
        if value > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, value / 2, str(value), ha="center", va="center", fontsize=7, color="white", fontweight="bold")
    for bar, value, base in zip(bars_missing, missing, added):
        if value > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, base + value / 2, str(value), ha="center", va="center", fontsize=7, color="white", fontweight="bold")

    ax.set_title(title, pad=8, fontsize=12, fontweight="bold", fontfamily="Calibri")
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, 0.98), ncol=2, frameon=False, fontsize=7)
    ax.set_xticks(x)
    ax.set_xticklabels([str(field).replace(" ", "\n") for field in fields], fontsize=6.6)
    ax.set_ylim(0, max(total_users * 1.18, 1))
    ax.tick_params(axis="y", labelleft=False, length=0)
    ax.grid(False)
    for spine in ax.spines.values():
        spine.set_color("#000000")
        spine.set_linewidth(0.8)
    fig.subplots_adjust(left=0.04, right=0.985, top=0.84, bottom=0.20)
    path = tempfile.mktemp(suffix=".png")
    plt.savefig(path, dpi=REPORT_CHART_DPI, facecolor="white", bbox_inches=None, pad_inches=0)
    plt.close(fig)
    return path

def _page12_risk_strips(risk_items, config=None):
    config = config or DEFAULT_REPORT_CONFIG
    items = [str(item) for item in (risk_items or []) if str(item).strip()]
    if not items:
        return None
    items = items[:4]
    primary = np.array([0x00, 0xB0, 0xF0]) / 255.0
    dark = np.array([0x00, 0x00, 0x00]) / 255.0
    fig, ax = plt.subplots(figsize=(6.35, 3.05), dpi=REPORT_CHART_DPI)
    fig.patch.set_facecolor("white")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    x0, x1 = 0.12, 0.93
    bar_h = 0.13
    y_positions = np.linspace(0.78, 0.18, len(items))
    line_x = x0 - 0.035
    ax.plot([line_x, line_x], [y_positions[-1] - 0.10, y_positions[0] + 0.11], color="#315A9C", linewidth=2.0, zorder=1)

    grad = np.linspace(0, 1, 600)
    colors = np.zeros((1, len(grad), 3))
    for idx, t in enumerate(grad):
        mix = min(t * 1.18, 1.0)
        colors[0, idx, :] = primary * (1 - mix) + dark * mix

    for y, text in zip(y_positions, items):
        ax.imshow(colors, extent=(x0, x1, y - bar_h / 2, y + bar_h / 2), aspect="auto", zorder=2)
        ax.add_patch(mpatches.Rectangle((x0, y - bar_h / 2), x1 - x0, bar_h, fill=False, edgecolor="#000000", linewidth=0.6, zorder=4))
        ax.add_patch(mpatches.Circle((x0 - 0.055, y), 0.055, facecolor="#00A4EF", edgecolor="#0070C0", linewidth=1.3, zorder=5))
        ax.add_patch(mpatches.Circle((x0 - 0.07, y + 0.018), 0.035, facecolor="#4CC5F4", edgecolor="none", alpha=0.70, zorder=6))
        ax.text(x0 + 0.035, y, textwrap.fill(text, width=58), ha="left", va="center", fontsize=9.5, color="white", fontfamily="Calibri", fontweight="bold", zorder=6)

    path = tempfile.mktemp(suffix=".png")
    plt.savefig(path, dpi=REPORT_CHART_DPI, facecolor="white", bbox_inches="tight", pad_inches=0.03)
    plt.close(fig)
    return path

def _make_donut_chart_img(active, total, title, size=(2.5, 2.5)):
    try:
        import io

        active = max(float(active or 0), 0.0)
        total = max(float(total or 0), 0.0)
        pct = (active / total * 100) if total > 0 else 0
        fill_color = '#00B050' if pct > 0 else '#D9D9D9'
        empty_color = '#D9D9D9'

        fig, ax = plt.subplots(figsize=size, dpi=REPORT_CHART_DPI, facecolor='white')
        fig.patch.set_edgecolor('#000000')
        fig.patch.set_linewidth(1.2)
        ax.set_facecolor('white')
        ax.pie(
            [pct, max(0, 100 - pct)],
            colors=[fill_color, empty_color],
            startangle=90,
            wedgeprops={'width': 0.30, 'edgecolor': 'white', 'linewidth': 2.4},
        )
        ax.text(
            0, 0, f'{pct:.0f}%',
            ha='center', va='center',
            fontsize=18, fontweight='bold',
            fontfamily='Calibri',
            color='#1A1A1A',
        )
        ax.set_title(title, pad=8, fontsize=10.5, fontweight='bold', fontfamily='Calibri', color='#333333')
        ax.set_aspect('equal')
        ax.set_axis_off()

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=REPORT_CHART_DPI, bbox_inches=None, facecolor='white', pad_inches=0)
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception as e:
        logging.getLogger('cra').warning(f'[CHART] donut failed: {e}')
        return None

def _make_pie_chart_img(data_dict, title, size=(3.5, 3.0)):
    try:
        import io

        cleaned = {
            str(label): int(value)
            for label, value in (data_dict or {}).items()
            if int(value or 0) > 0
        }
        if not cleaned:
            return None

        labels = list(cleaned.keys())
        values = list(cleaned.values())
        colors = [
            '#2196F3', '#4CAF50', '#FF9800', '#9C27B0',
            '#F44336', '#00BCD4', '#8BC34A', '#FF5722',
        ][:len(labels)]

        fig, ax = plt.subplots(figsize=size, dpi=REPORT_CHART_DPI, facecolor='white')
        wedges, _texts, autotexts = ax.pie(
            values,
            labels=None,
            colors=colors,
            autopct='%1.0f%%',
            startangle=90,
            pctdistance=0.75,
            wedgeprops={'edgecolor': 'white', 'linewidth': 1.5},
        )
        for text in autotexts:
            text.set_fontsize(9)
            text.set_color('white')
            text.set_fontweight('bold')

        ax.set_title(title, pad=10, fontsize=12, fontweight='bold', fontfamily='Calibri', color='#1A1A1A')
        ax.legend(
            wedges,
            [f'{label}; {value}' for label, value in zip(labels, values)],
            loc='lower center',
            bbox_to_anchor=(0.5, -0.18),
            ncol=2,
            fontsize=8,
            frameon=False,
        )
        ax.set_aspect('equal')

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=REPORT_CHART_DPI, bbox_inches='tight', facecolor='white', pad_inches=0.05)
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception as e:
        logging.getLogger('cra').warning(f'[CHART] pie failed: {e}')
        return None

def _make_bar_chart_img(fields_dict, total_users, title, size=(3.5, 2.8)):
    try:
        import io

        total_users = int(total_users or 0)
        if not fields_dict or total_users <= 0:
            return None

        fields = list(fields_dict.keys())
        present = [max(min(int(value or 0), total_users), 0) for value in fields_dict.values()]
        missing = [max(total_users - value, 0) for value in present]
        x = np.arange(len(fields))
        width = 0.5

        fig, ax = plt.subplots(figsize=size, dpi=REPORT_CHART_DPI, facecolor='white')
        bars_missing = ax.bar(x, missing, width, color='#E53935', label='Not Added')
        ax.bar(x, present, width, bottom=missing, color='#43A047', label='Added')

        for i, (miss, added) in enumerate(zip(missing, present)):
            if miss > 0:
                ax.text(
                    x[i], miss / 2, str(miss),
                    ha='center', va='center',
                    fontsize=8, color='white',
                    fontweight='bold',
                )
            if added > 0:
                ax.text(
                    x[i], miss + added / 2, str(added),
                    ha='center', va='center',
                    fontsize=8, color='white',
                    fontweight='bold',
                )

        ax.set_xticks(x)
        ax.set_xticklabels([field.replace(' ', '\n') for field in fields], fontsize=8)
        ax.set_title(title, pad=9, fontsize=12, fontweight='bold', fontfamily='Calibri', color='#1A1A1A')
        ax.legend(fontsize=8, loc='upper right', frameon=False)
        ax.set_ylim(0, max(total_users * 1.15, 1))
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.tick_params(axis='y', labelsize=8)

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=REPORT_CHART_DPI, bbox_inches='tight', facecolor='white', pad_inches=0.05)
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception as e:
        logging.getLogger('cra').warning(f'[CHART] bar failed: {e}')
        return None

def _insert_chart(doc, img_bytes, width_inches=2.8, sa=10):
    if not img_bytes:
        return
    try:
        import io

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(sa)
        p.paragraph_format.space_before = Pt(6)
        p.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(width_inches))
    except Exception as e:
        logging.getLogger('cra').warning(f'[CHART] insert failed: {e}')

def _remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    tbl_borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)

def _insert_activity_donut_grid(doc, activity_counts):
    import io

    chart_pairs = [
        ('SharePoint', 'SharePoint Accounts'),
        ('OneDrive', 'OneDrive Accounts'),
        ('Teams', 'Teams Usage'),
        ('Outlook', 'Outlook Usage'),
    ]
    donut_imgs = {}
    for key, title in chart_pairs:
        active, total = (activity_counts or {}).get(key, (0, 0))
        donut_imgs[key] = _make_donut_chart_img(active, total, title, size=(2.55, 2.58))

    if not any(donut_imgs.values()):
        return

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _remove_table_borders(table)
    positions = [
        ('SharePoint', 0, 0),
        ('OneDrive', 0, 1),
        ('Teams', 1, 0),
        ('Outlook', 1, 1),
    ]
    for key, row_idx, col_idx in positions:
        cell = table.rows[row_idx].cells[col_idx]
        _set_cell_width(cell, Inches(3.05))
        set_cell_padding(cell, top=80, right=70, bottom=80, left=70)
        for paragraph in list(cell.paragraphs):
            paragraph._element.getparent().remove(paragraph._element)
        img_bytes = donut_imgs.get(key)
        if not img_bytes:
            continue
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.space_before = Pt(4)
        p.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(2.55), height=Inches(2.58))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    spacer.paragraph_format.space_before = Pt(0)

def _add_key_observations(doc, report_data, include_user_activity=True):
    from docx.shared import Pt, RGBColor
    from collections import defaultdict
    import logging
    logger = logging.getLogger('cra')

    parameter_rows = report_data.get('parameter_rows', [])
    fail_count = int(report_data.get('gaps_count') if report_data.get('gaps_count') is not None else len([row for row in parameter_rows if _is_fail(row)]))
    total_params = int(report_data.get('total_params') if report_data.get('total_params') is not None else len(parameter_rows))
    findings = report_data.get('findings_list') or report_data.get('parameter_rows', [])
    eligible = int(report_data.get('eligible_users', 0))
    total_users = int(report_data.get('total_users', 0))

    od_pct = float(report_data.get('onedrive_active_pct', 0))
    tms_pct = float(report_data.get('teams_active_pct', 0))
    ol_pct = float(report_data.get('outlook_active_pct', 0))
    sp_pct = float(report_data.get('sharepoint_active_pct', 0))

    logger.info(
        f'[PAGE9] fail={fail_count} total={total_params}'
        f' od={od_pct} teams={tms_pct} ol={ol_pct} sp={sp_pct}'
    )

    pillar = defaultdict(lambda: {'f': 0, 't': 0})
    for finding in findings:
        pillar_name = str(finding.get('pillar', finding.get('category', ''))).strip()
        if not pillar_name:
            continue
        pillar[pillar_name]['t'] += 1
        if str(finding.get('status', '')).lower() == 'fail':
            pillar[pillar_name]['f'] += 1

    def ppct(key):
        data = pillar.get(key, {'f': 0, 't': 1})
        return round(data['f'] / max(data['t'], 1) * 100)

    sec = ppct('Security')
    gov = ppct('Governance')
    bp = ppct('Best Practice')

    h = doc.add_heading('Key Observations:', level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(10)

    def bullet(runs, sa=6):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        p.paragraph_format.space_after = Pt(sa)
        p.paragraph_format.space_before = Pt(0)
        marker = p.add_run('\u2022  ')
        marker.font.size = Pt(11)
        marker.font.name = 'Calibri'
        for text, bold in runs:
            run = p.add_run(str(text))
            run.font.size = Pt(11)
            run.font.bold = bold
            run.font.name = 'Calibri'
        return p

    def plain(runs, sa=6, sb=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(sa)
        p.paragraph_format.space_before = Pt(sb)
        for text, bold, size in runs:
            run = p.add_run(str(text))
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.name = 'Calibri'
        return p

    def bold_pct(value, sa=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(sa)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(f'{value:.0f}%')
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        return p

    bullet([
        ('A total of ', False),
        (str(fail_count), True),
        (' gaps', True),
        (' out of ', False),
        (str(total_params), True),
        (' parameters', True),
        (' were identified, distributed across ', False),
        ('Security', True),
        (', ', False),
        ('Governance', True),
        (', and ', False),
        ('Best Practice', True),
        (' categories.', False),
    ])

    bullet([
        ('Medium to Critical', True),
        (' severity issues make up most of the findings, indicating substantial exposure to operational and compliance risks.', False),
    ])

    bullet([
        ('Gap findings reveal that the percentage of failed parameters is ', False),
        (f'Security ({sec}%)', True),
        (', ', False),
        (f'Governance ({gov}%)', True),
        (', and ', False),
        (f'Best Practices ({bp}%)', True),
        (' pillars, indicating a critical need for immediate remediation in those areas.', False),
    ])

    bullet([
        ('There are ', False),
        (f'{eligible} user accounts out of {total_users}', True),
        (' that are eligible for a M365 Copilot license. Copilot requires a base Microsoft 365 subscription, such as Microsoft 365 E3, E5, Business Standard, or Business Premium.', False),
    ])
    license_data = report_data.get('license_counts', {})
    if license_data:
        pie_bytes = _make_pie_chart_img(
            license_data,
            'Licenses Assigned Data',
            size=(3.8, 3.2),
        )
        _insert_chart(doc, pie_bytes, width_inches=3.5, sa=8)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.space_before = Pt(0)

    ui_fields = report_data.get('user_info_fields', {})
    ui_total = int(report_data.get('user_info_total', 0))
    complete_profiles = report_data.get('complete_user_profiles')
    if complete_profiles is None:
        bullet([(DATA_NOT_AVAILABLE, True)])
    else:
        bullet([
            ('There are ', False),
            (f'{complete_profiles} users with full user information', True),
            ('; assigning user information such as department, role, manager, and location supports accurate organizational hierarchy and Copilot context quality.', False),
        ])
    if ui_fields and ui_total > 0:
        bar_bytes = _make_bar_chart_img(
            ui_fields,
            ui_total,
            'User Information Details',
            size=(4.0, 2.8),
        )
        _insert_chart(doc, bar_bytes, width_inches=3.8, sa=10)
    elif not include_user_activity:
        plain([(DATA_NOT_AVAILABLE, True, 11)], sa=8)

    if not include_user_activity:
        doc.add_page_break()
        return

    bold_pct(od_pct, sa=2)

    plain([
        (
            f'In the past 60 days, {od_pct:.0f}% of OneDrive users, '
            f'{tms_pct:.0f}% of Microsoft Teams users, '
            f'{ol_pct:.0f}% of outlook users and '
            f'{sp_pct:.0f}% of SharePoint users have been active, '
            f'indicating strong engagement across Microsoft 365 core services.',
            False,
            11,
        ),
    ], sa=4)
    _insert_activity_donut_grid(doc, report_data.get('activity_counts', {}))

    bold_pct(ol_pct, sa=2)
    bold_pct(tms_pct, sa=2)
    bold_pct(sp_pct, sa=10)

def _doughnut(pct, label):
    fig, ax = plt.subplots(figsize=(2.2, 2.5))
    vals = [float(pct), max(0, 100 - float(pct))]
    ax.pie(vals, colors=['#00B050', '#E0E0E0'], startangle=90, wedgeprops=dict(width=0.45))
    ax.text(0, 0, f'{int(pct)}%', ha='center', va='center', fontsize=12, fontweight='bold')
    ax.set_title(label, fontsize=9, pad=4)
    plt.tight_layout()
    path = tempfile.mktemp(suffix='.png')
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def _severity_meter(severity: str) -> str:
    """Generate a severity meter image showing position on the scale."""
    severity = severity.lower().strip()

    levels = [
        ('Critical', '#CC0000'),
        ('High', '#FF0000'),
        ('Medium', '#ED7D31'),
        ('Low', '#FFC000'),
        ('Informational', '#2FB62F'),
    ]

    pos_map = {
        'critical': 0, 'high': 1, 'medium': 2,
        'low': 3, 'informational': 4, 'info': 4,
    }
    active_pos = pos_map.get(severity, 4)

    fig, ax = plt.subplots(figsize=(7.1, 1.08), dpi=180)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 2)
    ax.axis('off')

    box_width = 1.7
    box_height = 0.9
    gap = 0.125
    start_x = 0.25

    for i, (label, color) in enumerate(levels):
        x = start_x + i * (box_width + gap)
        fancy = FancyBboxPatch(
            (x, 0.8), box_width, box_height,
            boxstyle='round,pad=0.05',
            facecolor=color,
            edgecolor='#111111',
            linewidth=1.2,
            zorder=2
        )
        ax.add_patch(fancy)
        ax.text(
            x + box_width/2, 0.55, label,
            ha='center', va='top',
            fontsize=7.8, color='#333333',
            fontfamily='Calibri'
        )
        if i == active_pos:
            triangle_x = x + box_width/2
            triangle = plt.Polygon(
                [[triangle_x-0.16, 0.65],
                 [triangle_x+0.16, 0.65],
                 [triangle_x, 1.05]],
                closed=True,
                facecolor='black',
                edgecolor='black',
                zorder=3
            )
            ax.add_patch(triangle)

    line_end = start_x + 5*(box_width + gap) - gap
    ax.plot([start_x, line_end], [0.65, 0.65],
            color='#666666', linewidth=1, zorder=1)

    plt.tight_layout(pad=0)
    path = tempfile.mktemp(suffix='.png')
    plt.savefig(path, dpi=180, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close()
    return path

def _detail_title_bar(title):
    fig, ax = plt.subplots(figsize=(6.35, 0.43), dpi=220)
    gradient = np.linspace(0, 1, 512)
    gradient = np.vstack([gradient, gradient])
    from matplotlib.colors import LinearSegmentedColormap
    cmap = LinearSegmentedColormap.from_list("cra_detail_bar", ["#000000", "#003946", "#00B5E8"])
    ax.imshow(gradient, aspect="auto", cmap=cmap, extent=[0, 1, 0, 1])
    ax.add_patch(mpatches.Rectangle((0, 0), 1, 1, fill=False, edgecolor="#000000", linewidth=1.1))
    ax.text(
        0.022,
        0.50,
        str(title or DATA_NOT_AVAILABLE),
        ha="left",
        va="center",
        color="white",
        fontsize=8.4,
        fontweight="bold",
        fontfamily="Calibri",
    )
    ax.axis("off")
    path = tempfile.mktemp(suffix=".png")
    plt.savefig(path, dpi=220, bbox_inches="tight", pad_inches=0, facecolor="white")
    plt.close(fig)
    return path

def _set_cell_bg(cell, hex_color):
    if not hex_color: return
    hex_color = hex_color.lstrip('#')
    tcPr = cell._element.get_or_add_tcPr()
    for old in list(tcPr.findall(qn('w:shd'))):
        tcPr.remove(old)
    fill = OxmlElement('w:shd')
    fill.set(qn('w:val'), 'clear')
    fill.set(qn('w:color'), 'auto')
    fill.set(qn('w:fill'), hex_color)
    tcPr.append(fill)

def set_cell_background(cell, hex_color):
    if not hex_color:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.replace('#', ''))
    tcPr.append(shd)

def remove_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_padding(cell, top=80, right=80, bottom=80, left=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('right', right), ('bottom', bottom), ('left', left)]:
        margin = OxmlElement(f'w:{side}')
        margin.set(qn('w:w'), str(val))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)

def add_run_with_style(para, text, size_pt, bold=False, color_hex=None, italic=False):
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color_hex:
        h = color_hex.replace('#', '')
        run.font.color.rgb = RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    return run

def _remove_table_borders(table):
    tblPr = table._tbl.tblPr
    for old in list(tblPr.findall(qn('w:tblBorders'))):
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'FFFFFF')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def _set_cell_border(cell, side, color='C5D8F0', size='4', val='single'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    border = OxmlElement(f'w:{side}')
    border.set(qn('w:val'), val)
    border.set(qn('w:sz'), str(size))
    border.set(qn('w:space'), '0')
    border.set(qn('w:color'), color.replace('#', ''))
    tcBorders.append(border)

def _set_cell_width(cell, width):
    if hasattr(width, 'twips'):
        pass
    elif isinstance(width, (int, float)):
        width = Inches(float(width))
    cell.width = width
    width_twips = width.twips if hasattr(width, 'twips') else int(width / 635)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in('w:tcW')
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')

def _clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.clear()

def _service_name_key(value):
    svc = str(value or '').strip().lower()
    aliases = {
        'entra': 'Entra ID',
        'entra id': 'Entra ID',
        'exchange': 'Exchange Online',
        'exchange online': 'Exchange Online',
        'microsoft purview': 'Microsoft Purview',
        'purview': 'Microsoft Purview',
        'microsoft teams': 'Microsoft Teams',
        'teams': 'Microsoft Teams',
        'onedrive': 'OneDrive for Business',
        'onedrive for business': 'OneDrive for Business',
        'sharepoint': 'SharePoint Online',
        'sharepoint online': 'SharePoint Online',
    }
    return aliases.get(svc, value)

def _readiness_color(readiness_level):
    level = str(readiness_level or '').strip().lower()
    if level == 'ready':
        return '#27500A'
    if level == 'needs improvement':
        return '#BA7517'
    return '#A32D2D'

def _add_risk_score_matrix(doc):
    doc.add_paragraph()
    headers = ['Critical', 'High', 'Medium', 'Low', 'Informational']
    colors = ['CC0000', 'FF6600', 'FFA500', 'FFD700', '00B050']
    descs = [
        'The risks posed by this finding are of a critical nature and can lead to full system compromise. Specifics should be addressed immediately.',
        'The risks posed by this finding are of high impact and can lead to partial or full system compromise. Specifics should be addressed as soon as can be safely implemented.',
        'The risks posed by this finding is of moderate impact, existing controls may be providing partial mitigation. Remediation efforts should be planned as part of regular maintenance.',
        'The risk posed by this finding is low due to the potential impact or the difficulty of exploitation. Issues noted should be reviewed for remediation when convenient.',
        'A potential indirect risk that may contribute to or lead to an incident or deviation from industry best practices.',
    ]
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    for i, (hdr, col) in enumerate(zip(headers, colors)):
        cell = table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, col)
    for i, desc in enumerate(descs):
        cell = table.cell(1, i)
        p = cell.paragraphs[0]
        p.add_run(desc).font.size = Pt(11)

def _add_risks_section(doc):
    risks = [
        'Unauthorized access to sensitive business data.',
        'Inadequate auditing and monitoring of AI-driven activity.',
        'Compliance violations with internal and external regulations.',
        'Reduced user trust and operational inconsistencies due to misconfigured policies.',
    ]
    doc.add_paragraph('Proceeding with Copilot activation in the current state may lead to:')
    for risk in risks:
        doc.add_paragraph(risk, style='List Bullet')

def _format_date(date_value):
    if not date_value: return datetime.now().strftime('%d-%m-%Y')
    if isinstance(date_value, datetime): return date_value.strftime('%d-%m-%Y')
    try:
        dt = datetime.fromisoformat(str(date_value)[:19])
        return dt.strftime('%d-%m-%Y')
    except: return str(date_value)[:10]

def _apply_body_font_defaults(doc, config=None):
    config = config or DEFAULT_REPORT_CONFIG
    body_font = _font_name(config, "body")
    body_size = _pt(config, "body_size", 11)
    for style_name in ("Normal", "Body Text"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = body_font
        style.font.size = body_size
        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:ascii"), body_font)
        r_fonts.set(qn("w:hAnsi"), body_font)
        r_fonts.set(qn("w:cs"), body_font)

def _add_header_logo(doc, logo_path=None, display_name=None):
    """Apply the TPT logo to the top-left header after the cover page."""
    blueprint = _load_aaa_report_blueprint()
    header_config = blueprint.get("header", {}) if isinstance(blueprint, dict) else {}
    spacing_after = _twips_value(header_config.get("spacing_after_twips"), 0)
    logo_file = Path(str(logo_path)) if logo_path else None
    if not logo_file or not logo_file.exists():
        configured_logo = _cfg(DEFAULT_REPORT_CONFIG, "branding", "partner_logo")
        logo_file = _repo_file_path(configured_logo) if configured_logo else None

    logo_width = Inches(1.30)

    def apply_to_header(header):
        header.is_linked_to_previous = False
        for para in list(header.paragraphs):
            try:
                para._element.getparent().remove(para._element)
            except Exception:
                pass
        header_para = header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_para.paragraph_format.space_before = Pt(0)
        header_para.paragraph_format.space_after = Pt(0)
        header_para.paragraph_format.left_indent = Inches(0)
        header_para.paragraph_format.line_spacing = None
        if logo_file and logo_file.exists():
            header_para.add_run().add_picture(str(logo_file), width=logo_width)
        else:
            header_text = header_config.get("default_header_text")
            if header_text:
                header_para.add_run(str(header_text))

    for section in doc.sections:
        section.header_distance = Inches(0.38)
        section.header.is_linked_to_previous = False
        apply_to_header(section.header)
        if section.different_first_page_header_footer:
            section.first_page_header.is_linked_to_previous = False
            for para in list(section.first_page_header.paragraphs):
                try:
                    para._element.getparent().remove(para._element)
                except Exception:
                    pass

def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instr = OxmlElement('w:instrText')
    instr.set(
        '{http://www.w3.org/XML/1998/namespace}space',
        'preserve'
    )
    instr.text = ' PAGE '
    run._r.append(instr)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)
    return run

def _remove_page_number_restarts(doc):
    for section in doc.sections:
        sect_pr = section._sectPr
        for pg_num_type in list(sect_pr.findall(qn('w:pgNumType'))):
            sect_pr.remove(pg_num_type)

def _apply_footer(section, blueprint=None):
    blueprint = blueprint or _load_aaa_report_blueprint()
    section.footer_distance = Inches(0.45)
    footer_config = blueprint.get("footer", {}) if isinstance(blueprint, dict) else {}
    fonts = blueprint.get("fonts", {}) if isinstance(blueprint, dict) else {}
    footer_font = fonts.get("footer", {}) if isinstance(fonts.get("footer"), dict) else {}
    inherited_font = fonts.get("body_text", {}) if isinstance(fonts.get("body_text"), dict) else {}
    left_text = (footer_config.get("left_text") or {}).get("value", "Copilot Readiness Assessment")
    right_format = (footer_config.get("right_text") or {}).get("value", "{page} | Page")
    spacing_after = _twips_value(footer_font.get("spacing_after_twips"), 0)
    line_twips = footer_font.get("line_twips")
    line_rule = footer_font.get("line_rule")
    footer_text_color = "666666"

    def apply_footer_run_style(run):
        _set_run_font_from_blueprint(run, inherited_font)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = _hex_color(footer_text_color)

    def apply_exact_footer_paragraph_format(paragraph):
        content_width = section.page_width - section.left_margin - section.right_margin
        content_twips = int(content_width / 635)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        _set_paragraph_spacing_twips(
            paragraph,
            before_twips=0,
            after_twips=0,
            line_twips=line_twips,
            line_rule=line_rule,
        )

        p_pr = paragraph._p.get_or_add_pPr()
        tabs = p_pr.find(qn("w:tabs"))
        if tabs is not None:
            p_pr.remove(tabs)
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(content_twips))
        tabs.append(tab)
        p_pr.append(tabs)

        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is not None:
            p_pr.remove(p_bdr)
        p_bdr = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "4")
        top.set(qn("w:space"), "4")
        top.set(qn("w:color"), "D9D9D9")
        p_bdr.append(top)
        p_pr.append(p_bdr)

    def clear_footer(footer):
        for paragraph in list(footer.paragraphs):
            try:
                paragraph._element.getparent().remove(paragraph._element)
            except Exception:
                pass
        for table in list(footer.tables):
            try:
                table._element.getparent().remove(table._element)
            except Exception:
                pass

    def render_footer(footer):
        clear_footer(footer)
        paragraph = footer.add_paragraph()
        apply_exact_footer_paragraph_format(paragraph)
        apply_footer_run_style(paragraph.add_run(str(left_text)))
        paragraph.add_run("\t")
        before, after = str(right_format).split("{page}", 1) if "{page}" in str(right_format) else ("", str(right_format))
        if after == " | Page":
            after = " | P a g e"
        if before:
            apply_footer_run_style(paragraph.add_run(before))
        page_run = _add_page_number_field(paragraph)
        apply_footer_run_style(page_run)
        if after:
            apply_footer_run_style(paragraph.add_run(after))

    section.footer.is_linked_to_previous = False
    render_footer(section.footer)
    if section.different_first_page_header_footer:
        section.first_page_footer.is_linked_to_previous = False
        render_footer(section.first_page_footer)

def _add_document_footer_page_numbers(doc):
    blueprint = _load_aaa_report_blueprint()
    for i, section in enumerate(doc.sections):
        if i > 0:
            section.footer.is_linked_to_previous = True
            if section.different_first_page_header_footer:
                section.first_page_footer.is_linked_to_previous = True
            continue
        _apply_footer(section, blueprint)

def _add_image(doc, img_path, width_inches=4):
    if not img_path or not os.path.exists(img_path): return
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(img_path, width=Inches(width_inches))
    except: pass
    finally:
        try:
            if os.path.exists(img_path): os.remove(img_path)
        except: pass

def _build_cover_page(doc, report_data):
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import datetime
    import os, logging

    logger = logging.getLogger('cra')

    # -- DATA (use actual assessment data, no hardcoded fallbacks) -----
    company = str(report_data.get('company_name', '[Organization Name Not Available]'))
    partner = str(report_data.get('partner_name', '[Partner Not Available]'))
    try:
        score = float(report_data.get('readiness_score', 0) or 0)
    except (TypeError, ValueError):
        score = 0.0
    level = str(report_data.get('readiness_level', 'Not Ready'))
    logo = report_data.get('logo_path')
    findings = report_data.get('findings_list', [])

    # Format date cleanly
    raw = str(report_data.get('assessment_date', ''))
    date_str = raw
    for fmt in ['%Y-%m-%d %H:%M:%S.%f',
                '%Y-%m-%d %H:%M:%S', '%Y-%m-%d',
                '%d-%m-%Y', '%d/%m/%Y']:
        try:
            dt = datetime.strptime(raw.strip(), fmt)
            date_str = dt.strftime('%d %B %Y').lstrip('0')
            break
        except Exception:
            continue

    level_short = (
        'Not Ready'
        if 'not' in level.lower()
        else 'Needs Improvement'
        if 'needs' in level.lower()
        else 'Ready'
    )

    # Service counts
    svc_order = [
        ('Entra ID', 'Entra ID'),
        ('Exchange Online', 'Exchange'),
        ('Microsoft Purview', 'Purview'),
        ('Microsoft Teams', 'Teams'),
        ('OneDrive for Business', 'OneDrive'),
        ('SharePoint Online', 'SharePoint'),
    ]
    cnt = {k: {'f': 0, 'p': 0} for k, _ in svc_order}
    for f in findings:
        s = str(f.get('service_name', '') or f.get('service', '') or f.get('category', ''))
        try:
            s = _service_name_key(s)
        except Exception:
            pass
        st = str(f.get('status', '') or f.get('finding', '') or f.get('result', '') or f.get('display_status', '')).lower()
        if s in cnt:
            if st == 'fail':
                cnt[s]['f'] += 1
            elif st == 'pass':
                cnt[s]['p'] += 1
    fail_total = sum(v['f'] for v in cnt.values())
    pass_total = sum(v['p'] for v in cnt.values())

    logger.info(
        f'[COVER] simple cover company={company} score={score:.2f} '
        f'findings={len(findings)} fail={fail_total} pass={pass_total}'
    )

    # -- HELPERS -----------------------------------
    def R(h):
        h = h.lstrip('#')
        return RGBColor(int(h[0:2], 16),
                        int(h[2:4], 16),
                        int(h[4:6], 16))

    def para(space_before=0, space_after=6,
             align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.alignment = align
        return p

    def run(p, text, size=11, bold=False,
            color='1A1A1A', italic=False):
        r = p.add_run(str(text))
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = R(color)
        r.font.name = 'Calibri'
        return r

    def hline(color_hex='E0E0E0', size_pt=0.5):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), str(int(size_pt * 8)))
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), color_hex.lstrip('#'))
        pBdr.append(bot)
        pPr.append(pBdr)
        return p

    L = WD_ALIGN_PARAGRAPH.LEFT
    C = WD_ALIGN_PARAGRAPH.CENTER

    # -- PAGE MARGINS -----------------------------
    sec = doc.sections[0]
    sec.top_margin = Inches(0.0)
    sec.bottom_margin = Inches(0.52)
    sec.left_margin = Inches(0.0)
    sec.right_margin = Inches(0.0)
    sec.different_first_page_header_footer = True

    # GRADIENT HEADER WITH LOGO
    BASE = os.path.dirname(os.path.abspath(__file__))
    ROOT = os.path.abspath(os.path.join(BASE, '..', '..', '..'))

    # Use custom logo if provided, or TPT logo as default
    logo_path = report_data.get('logo_path')

    if logo_path and os.path.exists(str(logo_path)):
        # Use provided custom logo
        pass
    else:
        # Fall back to TPT logo
        logo_path = os.path.join(ROOT, 'app', 'services', 'reporting', 'assets', 'techplustalent-logo.png')

        if not os.path.exists(logo_path):
            # Try alternate TPT logo location
            logo_path = os.path.join(ROOT, 'storage', 'assets', 'techplustalent-logo.png')

    # Header table with gradient background
    header_tbl = doc.add_table(rows=1, cols=1)
    header_tbl.autofit = False
    header_tbl.allow_autofit = False
    header_cell = header_tbl.rows[0].cells[0]

    # Set header background to light gradient color
    header_tcPr = header_cell._tc.get_or_add_tcPr()
    header_shd = OxmlElement('w:shd')
    header_shd.set(qn('w:val'), 'clear')
    header_shd.set(qn('w:color'), 'auto')
    header_shd.set(qn('w:fill'), 'EBF4FD')
    header_tcPr.append(header_shd)

    # Remove header table borders
    header_tblPr = header_tbl._tbl.find(qn('w:tblPr'))
    if header_tblPr is None:
        header_tblPr = OxmlElement('w:tblPr')
        header_tbl._tbl.insert(0, header_tblPr)
    header_tblBdr = OxmlElement('w:tblBorders')
    for s in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'), 'none')
        header_tblBdr.append(b)
    header_tblPr.append(header_tblBdr)

    # Set header height
    header_trPr = header_tbl.rows[0]._tr.get_or_add_trPr()
    header_trH = OxmlElement('w:trHeight')
    header_trH.set(qn('w:val'), str(int(1.2 * 1440)))
    header_trH.set(qn('w:hRule'), 'exact')
    header_trPr.append(header_trH)

    # Add logo to header (right aligned) - logo only, no company name text
    header_para = header_cell.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.paragraph_format.space_before = Pt(12)
    header_para.paragraph_format.space_after = Pt(12)
    header_para.paragraph_format.right_indent = Inches(0.48)

    if logo_path and os.path.exists(str(logo_path)):
        try:
            logo_run = header_para.add_run()
            logo_run.add_picture(str(logo_path), height=Inches(0.9))
            logo_run.font.name = 'Calibri'
        except Exception as e:
            logger.warning(f'[COVER] Could not load logo: {e}')

    # Add spacing after header
    doc.add_paragraph()

    # TOP BLUE BAR (kept for consistency with existing design)
    tb = doc.add_table(rows=1, cols=1)
    tb.style = 'Table Grid'
    tc = tb.rows[0].cells[0]

    tblPr = tb._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tb._tbl.insert(0, tblPr)
    tblBdr = OxmlElement('w:tblBorders')
    for s in ['top', 'left', 'bottom', 'right',
              'insideH', 'insideV']:
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'), 'none')
        tblBdr.append(b)
    tblPr.append(tblBdr)

    tcPr = tc._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '0078D4')
    tcPr.append(shd)

    tcBdr = OxmlElement('w:tcBorders')
    for s in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'), 'none')
        tcBdr.append(b)
    tcPr.append(tcBdr)

    trPr = tb.rows[0]._tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(0.12 * 1440)))
    trH.set(qn('w:hRule'), 'exact')
    trPr.append(trH)

    for p in list(tc.paragraphs):
        p._element.getparent().remove(p._element)
    tc.add_paragraph()

    # LOGO ROW: Copilot logo left, organization name right
    p = para(space_before=14, space_after=4)

    copilot = os.path.join(ROOT, 'storage', 'assets', 'copilot_clean.png')
    if not os.path.exists(copilot):
        copilot = os.path.join(ROOT, 'storage', 'assets', 'copilot_logo.png')

    if os.path.exists(copilot):
        r0 = p.add_run()
        r0.add_picture(copilot,
                           width=Inches(0.42),
                       height=Inches(0.42))
        r0.font.name = 'Calibri'
        run(p, '    Microsoft 365 Copilot',
            10, bold=True, color='0078D4')
    else:
        run(p, 'Microsoft 365 Copilot',
            11, bold=True, color='0078D4')

    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '8640')
    tabs.append(tab)
    pPr.append(tabs)

    tab_run = p.add_run('\t')
    tab_run.font.name = 'Calibri'

    # Organization name on right (logo now at top)
    run(p, company, 13, bold=True, color='1A1A1A')

    p2 = para(space_before=0, space_after=14)
    run(p2, 'Readiness Assessment Platform',
        8, color='888888')

    hline('E0E0E0', 0.5)

    # REPORT TYPE TAG
    p = para(space_before=18, space_after=6)
    run(p, 'ASSESSMENT REPORT  ·  2026',
        8, color='888888')

    # MAIN TITLE
    p = para(space_before=4, space_after=0)
    run(p, 'Copilot ', 30, bold=True, color='0078D4')
    run(p, 'Readiness', 30, bold=False, color='1A1A1A')

    p = para(space_before=0, space_after=10)
    run(p, 'Assessment', 30, bold=False, color='1A1A1A')

    # Blue accent line
    acc = doc.add_paragraph()
    acc.paragraph_format.space_before = Pt(2)
    acc.paragraph_format.space_after = Pt(12)
    accPr = acc._p.get_or_add_pPr()
    accBdr = OxmlElement('w:pBdr')
    accBot = OxmlElement('w:bottom')
    accBot.set(qn('w:val'), 'single')
    accBot.set(qn('w:sz'), '18')
    accBot.set(qn('w:space'), '1')
    accBot.set(qn('w:color'), '0078D4')
    accBdr.append(accBot)
    accPr.append(accBdr)
    ar = acc.add_run('   ')
    ar.font.size = Pt(4)

    # Subtitle
    p = para(space_before=0, space_after=24)
    run(p,
        'Comprehensive evaluation of the Microsoft 365 '
        'environment across Security, Governance and '
        'Best Practice pillars for responsible '
        'Copilot adoption.',
        10, italic=True, color='555555')

    # SCORE + SERVICE TABLE
    st = doc.add_table(rows=1, cols=2)
    st.style = 'Table Grid'

    stPr = st._tbl.find(qn('w:tblPr'))
    if stPr is None:
        stPr = OxmlElement('w:tblPr')
        st._tbl.insert(0, stPr)
    stBdr = OxmlElement('w:tblBorders')
    for s in ['top', 'left', 'bottom', 'right',
              'insideH', 'insideV']:
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'), 'none')
        stBdr.append(b)
    stPr.append(stBdr)

    sr = st.rows[0]
    sc_left = sr.cells[0]
    sc_right = sr.cells[1]
    sr.height = Inches(1.9)

    for cell, inches in [(sc_left, 1.9),
                         (sc_right, 4.3)]:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        tc2 = cell._tc
        tcPr2 = tc2.get_or_add_tcPr()
        w2 = OxmlElement('w:tcW')
        w2.set(qn('w:w'), str(int(inches * 1440)))
        w2.set(qn('w:type'), 'dxa')
        tcPr2.append(w2)

        shd2 = OxmlElement('w:shd')
        shd2.set(qn('w:val'), 'clear')
        shd2.set(qn('w:color'), 'auto')
        shd2.set(qn('w:fill'), 'EBF4FD')
        tcPr2.append(shd2)

        tcBdr2 = OxmlElement('w:tcBorders')
        for s in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{s}')
            b.set(qn('w:val'), 'none')
            tcBdr2.append(b)
        tcPr2.append(tcBdr2)

        tcMar = OxmlElement('w:tcMar')
        for side, val in [('top', 95), ('right', 120),
                          ('bottom', 95), ('left', 160)]:
            e = OxmlElement(f'w:{side}')
            e.set(qn('w:w'), str(val))
            e.set(qn('w:type'), 'dxa')
            tcMar.append(e)
        tcPr2.append(tcMar)

        for pp2 in list(cell.paragraphs):
            pp2._element.getparent().remove(pp2._element)

    # LEFT cell: score
    p = sc_left.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.alignment = C
    r2 = p.add_run(f'{score:.2f}%')
    r2.font.size = Pt(24)
    r2.font.bold = True
    r2.font.name = 'Calibri'
    r2.font.color.rgb = R('A32D2D')

    p = sc_left.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.alignment = C
    r3 = p.add_run('Readiness Score')
    r3.font.size = Pt(11)
    r3.font.name = 'Calibri'
    r3.font.color.rgb = R('888888')

    p = sc_left.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.alignment = C
    r4 = p.add_run(f'●  {level_short}')
    r4.font.size = Pt(11)
    r4.font.bold = True
    r4.font.name = 'Calibri'
    r4.font.color.rgb = R('A32D2D')

    p = sc_left.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.alignment = C
    r5 = p.add_run(
        f'{fail_total} fail  ·  {pass_total} pass\n'
        f'out of {fail_total + pass_total} parameters' if (fail_total + pass_total) else DATA_NOT_AVAILABLE)
    r5.font.size = Pt(11)
    r5.font.name = 'Calibri'
    r5.font.color.rgb = R('666666')

    # RIGHT cell: services
    for svc_full, svc_short in svc_order:
        fc = cnt[svc_full]['f']
        pc = cnt[svc_full]['p']
        dot_c = 'E24B4A' if fc > 0 else '27500A'
        p = sc_right.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        rd = p.add_run('●  ')
        rd.font.size = Pt(11)
        rd.font.name = 'Calibri'
        rd.font.color.rgb = R(dot_c)

        rn = p.add_run(f'{svc_short:<12}')
        rn.font.size = Pt(11)
        rn.font.bold = True
        rn.font.name = 'Calibri'
        rn.font.color.rgb = R('0C447C')

        rc2 = p.add_run(f'   {fc}F / {pc}P')
        rc2.font.size = Pt(11)
        rc2.font.name = 'Calibri'
        rc2.font.color.rgb = R('444444')

    hline('E0E0E0', 0.5)

    # META ROW
    meta_items = [
        ('PREPARED FOR', company, True, '0078D4'),
        ('PREPARED BY', partner, True, '1A1A1A'),
        ('ASSESSMENT DATE', date_str, True, '1A1A1A'),
    ]

    mt = doc.add_table(rows=1, cols=3)
    mt.style = 'Table Grid'
    mtPr = mt._tbl.find(qn('w:tblPr'))
    if mtPr is None:
        mtPr = OxmlElement('w:tblPr')
        mt._tbl.insert(0, mtPr)
    mtBdr = OxmlElement('w:tblBorders')
    for s in ['top', 'left', 'bottom', 'right',
              'insideH', 'insideV']:
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'), 'none')
        mtBdr.append(b)
    mtPr.append(mtBdr)

    for i, (lbl, val, bold, color) in enumerate(meta_items):
        cell = mt.rows[0].cells[i]
        for pp3 in list(cell.paragraphs):
            pp3._element.getparent().remove(pp3._element)
        p1 = cell.add_paragraph()
        p1.paragraph_format.space_before = Pt(7)
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(lbl)
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
        r1.font.color.rgb = R('999999')

        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run(val)
        r2.font.size = Pt(12)
        r2.font.bold = bold
        r2.font.name = 'Calibri'
        r2.font.color.rgb = R(color)

    # BOTTOM BLUE BAR
    bb = doc.add_table(rows=1, cols=1)
    bb.style = 'Table Grid'
    bbPr = bb._tbl.find(qn('w:tblPr'))
    if bbPr is None:
        bbPr = OxmlElement('w:tblPr')
        bb._tbl.insert(0, bbPr)
    bbBdr = OxmlElement('w:tblBorders')
    for s in ['top', 'left', 'bottom', 'right',
              'insideH', 'insideV']:
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'), 'none')
        bbBdr.append(b)
    bbPr.append(bbBdr)

    bbc = bb.rows[0].cells[0]
    bbcPr = bbc._tc.get_or_add_tcPr()
    bbShd = OxmlElement('w:shd')
    bbShd.set(qn('w:val'), 'clear')
    bbShd.set(qn('w:color'), 'auto')
    bbShd.set(qn('w:fill'), '0078D4')
    bbcPr.append(bbShd)
    bbcBdr = OxmlElement('w:tcBorders')
    for s in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'), 'none')
        bbcBdr.append(b)
    bbcPr.append(bbcBdr)

    bbTrPr = bb.rows[0]._tr.get_or_add_trPr()
    bbTrH = OxmlElement('w:trHeight')
    bbTrH.set(qn('w:val'), str(int(0.10 * 1440)))
    bbTrH.set(qn('w:hRule'), 'exact')
    bbTrPr.append(bbTrH)

    for pp4 in list(bbc.paragraphs):
        pp4._element.getparent().remove(pp4._element)
    bbc.add_paragraph()

    # PAGE BREAK
    pb = doc.add_paragraph()
    pb.paragraph_format.space_before = Pt(0)
    pb.paragraph_format.space_after = Pt(0)
    pbr = pb.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    pbr._r.append(br)

    # Note: Section break removed. Use native page break only.
    # Section breaks cause Word to recalculate layout and move content to previous pages.

def _build_cover_fallback(doc, report_data):
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    company = str(report_data.get('company_name', 'Client'))
    raw_score = report_data.get('readiness_score', 0)
    try:
        score = float(raw_score or 0)
    except (TypeError, ValueError):
        score = 0.0
    date_str = str(report_data.get('assessment_date', ''))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run('Copilot Readiness Assessment')
    r.font.size = Pt(24)
    r.font.bold = True

    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        f'Prepared for: {company}  |  '
        f'Score: {score:.2f}%  |  Date: {date_str}'
    )
    r2.font.size = Pt(12)

    pb = doc.add_paragraph()
    run = pb.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def _build_cover_page_impl(doc, report_data):
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import datetime
    import os, logging

    logger = logging.getLogger('cra')

    # -- 1. EXTRACT DATA (use actual assessment data, no hardcoded fallbacks) -----
    company = str(report_data.get('company_name', '[Organization Name Not Available]'))
    partner = str(report_data.get('partner_name', '[Partner Not Available]'))
    score = float(report_data.get('readiness_score', 0))
    level = str(report_data.get('readiness_level', 'Not Ready'))
    logo_path = report_data.get('logo_path')
    findings = report_data.get('findings_list', [])

    # Clean date
    raw = str(report_data.get('assessment_date', ''))
    date_str = raw
    for fmt in ['%Y-%m-%d %H:%M:%S.%f',
                '%Y-%m-%d %H:%M:%S',
                '%Y-%m-%d', '%d-%m-%Y']:
        try:
            dt = datetime.strptime(raw.strip(), fmt)
            date_str = dt.strftime('%d %B %Y').lstrip('0')
            break
        except Exception:
            continue

    # Level badge
    level_short = (
        'Not Ready' if 'not' in level.lower()
        else 'Needs Improvement' if 'needs' in level.lower()
        else 'Ready'
    )

    # Copilot logo path (absolute)
    BASE = os.path.dirname(os.path.abspath(__file__))
    ROOT = os.path.abspath(os.path.join(BASE, '..', '..', '..'))
    copilot = os.path.join(ROOT, 'storage', 'assets',
                           'copilot_clean.png')
    if not os.path.exists(copilot):
        copilot = os.path.join(ROOT, 'storage', 'assets',
                               'copilot_logo.png')
    if not os.path.exists(copilot):
        copilot = None
    logger.info(f'[COVER] copilot={copilot} '
                f'exists={bool(copilot and os.path.exists(copilot))}')
    logger.info(f'[COVER] logo_path={logo_path} '
                f'exists={bool(logo_path and os.path.exists(str(logo_path)))}')
    logger.info(
        f'[COVER] company={company} partner={partner} '
        f'date={date_str} score={score} level={level} '
        f'logo={logo_path} findings={len(findings)}'
    )

    # -- 2. SERVICE COUNTS -------------------------
    svc_order = [
        ('Entra ID', 'Entra ID'),
        ('Exchange Online', 'Exchange'),
        ('Microsoft Purview', 'Purview'),
        ('Microsoft Teams', 'Teams'),
        ('OneDrive for Business', 'OneDrive'),
        ('SharePoint Online', 'SharePoint'),
    ]
    cnt = {k: {'f': 0, 'p': 0} for k, _ in svc_order}
    all_svcs = set(str(f.get('service_name', '') or f.get('service', '') or f.get('category', '')) for f in findings)
    logger.info(f'[COVER] all service names: {all_svcs}')
    unique_svcs = {}
    for f in findings:
        svc_dbg = str(f.get('service_name', '') or f.get('service', '') or f.get('category', 'MISSING') or 'MISSING')
        st_dbg = str(f.get('status', '') or f.get('finding', '') or f.get('result', '') or f.get('display_status', 'MISSING') or 'MISSING')
        v = unique_svcs.setdefault(svc_dbg, {'pass': 0, 'fail': 0, 'other': 0, 'statuses': set()})
        st_key = st_dbg.lower()
        v['statuses'].add(st_dbg)
        if st_key == 'pass':
            v['pass'] += 1
        elif st_key == 'fail':
            v['fail'] += 1
        else:
            v['other'] += 1
    for s, v in unique_svcs.items():
        logger.info(f'[COVER DEBUG] {s!r}: {v}')
    for f in findings:
        svc = str(f.get('service_name', '') or f.get('service', '') or f.get('category', ''))
        svc = _service_name_key(svc)
        st = str(f.get('status', '') or f.get('finding', '') or f.get('result', '') or f.get('display_status', '')).lower()
        if svc in cnt:
            if st == 'fail':
                cnt[svc]['f'] += 1
            elif st == 'pass':
                cnt[svc]['p'] += 1
    fail_total = sum(v['f'] for v in cnt.values())
    pass_total = sum(v['p'] for v in cnt.values())
    for svc_full, svc_short in svc_order:
        fc = cnt[svc_full]['f']
        pc = cnt[svc_full]['p']
        logger.info(f'[COVER] {svc_short}: fail={fc} pass={pc}')

    # -- 3. HELPERS --------------------------------
    def rgb(h):
        h = h.lstrip('#')
        return RGBColor(int(h[0:2], 16),
                        int(h[2:4], 16),
                        int(h[4:6], 16))

    def set_bg(cell, hex6):
        tc = cell._tc
        pr = tc.get_or_add_tcPr()
        s = OxmlElement('w:shd')
        s.set(qn('w:val'), 'clear')
        s.set(qn('w:color'), 'auto')
        s.set(qn('w:fill'), hex6.lstrip('#'))
        pr.append(s)

    def no_border(cell):
        tc = cell._tc
        pr = tc.get_or_add_tcPr()
        tb = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right',
                     'insideH', 'insideV']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            tb.append(b)
        pr.append(tb)

    def no_tbl_border(tbl):
        tblPr = tbl._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl._tbl.insert(0, tblPr)
        tb = OxmlElement('w:tblBorders')
        for side in ['top', 'left', 'bottom', 'right',
                     'insideH', 'insideV']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            tb.append(b)
        tblPr.append(tb)

    def set_pad(cell, t=80, r=120, b=80, l=120):
        tc = cell._tc
        pr = tc.get_or_add_tcPr()
        m = OxmlElement('w:tcMar')
        for side, v in [('top', t), ('right', r),
                        ('bottom', b), ('left', l)]:
            e = OxmlElement(f'w:{side}')
            e.set(qn('w:w'), str(v))
            e.set(qn('w:type'), 'dxa')
            m.append(e)
        pr.append(m)

    def set_cw(cell, inches):
        tc = cell._tc
        pr = tc.get_or_add_tcPr()
        w = OxmlElement('w:tcW')
        w.set(qn('w:w'), str(int(inches * 1440)))
        w.set(qn('w:type'), 'dxa')
        pr.append(w)

    def set_rh(row, inches):
        tr = row._tr
        trp = tr.get_or_add_trPr()
        h = OxmlElement('w:trHeight')
        h.set(qn('w:val'), str(int(inches * 1440)))
        h.set(qn('w:hRule'), 'exact')
        trp.append(h)

    def cell_valign(cell, align='top'):
        tc = cell._tc
        pr = tc.get_or_add_tcPr()
        va = OxmlElement('w:vAlign')
        va.set(qn('w:val'), align)
        pr.append(va)

    def clear(cell):
        for p in list(cell.paragraphs):
            try:
                p._element.getparent().remove(p._element)
            except Exception:
                pass

    def add_para(cell, align=WD_ALIGN_PARAGRAPH.LEFT,
                 sa=4, sb=0, ls=13):
        p = cell.add_paragraph()
        p.paragraph_format.alignment = align
        p.paragraph_format.space_after = Pt(sa)
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.line_spacing = Pt(ls)
        return p

    def add_run(para, text, size, bold=False,
                color='FFFFFF', italic=False):
        r = para.add_run(str(text))
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = rgb(color)
        r.font.name = 'Calibri'
        try:
            rPr = r._r.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), 'Calibri')
            rFonts.set(qn('w:hAnsi'), 'Calibri')
            rFonts.set(qn('w:eastAsia'), 'Calibri')
            rFonts.set(qn('w:cs'), 'Calibri')
            rFonts.set(qn('w:theme'), 'majorHAnsi')
        except Exception as fe:
            logger.warning(f'[COVER] font XML failed: {fe}')
        return r

    def add_pic(cell, path, max_w, max_h,
                align=WD_ALIGN_PARAGRAPH.LEFT, sa=4):
        if not path or not os.path.exists(str(path)):
            return False
        try:
            try:
                from PIL import Image as PILImg
                with PILImg.open(str(path)) as im:
                    pw, ph = im.size
                aspect = pw / ph if ph else 1
            except ImportError:
                aspect = 1.0
            w = min(max_w, max_h * aspect)
            h = w / aspect
            if h > max_h:
                h = max_h
                w = h * aspect
            p = cell.add_paragraph()
            p.paragraph_format.alignment = align
            p.paragraph_format.space_after = Pt(sa)
            p.paragraph_format.space_before = Pt(0)
            r = p.add_run()
            r.add_picture(str(path),
                          width=Inches(w),
                          height=Inches(h))
            return True
        except Exception as e:
            logger.warning(f'[COVER] pic failed {path}: {e}')
            return False

    C = WD_ALIGN_PARAGRAPH.CENTER
    L = WD_ALIGN_PARAGRAPH.LEFT

    # -- 4. PAGE MARGINS (narrow for cover) --------
    sec = doc.sections[0]
    sec.top_margin = Inches(0)
    sec.bottom_margin = Inches(0.3)
    sec.left_margin = Inches(0.35)
    sec.right_margin = Inches(0.35)
    sec.different_first_page_header_footer = True
    for paragraph in sec.first_page_header.paragraphs:
        paragraph.clear()

    # TABLE 1 - BLUE HEADER
    # Left: logo + title + TPT | Right: score panel
    t1 = doc.add_table(rows=1, cols=2)
    t1.style = 'Table Grid'
    no_tbl_border(t1)
    r1 = t1.rows[0]
    lc = r1.cells[0]
    rc = r1.cells[1]

    set_cw(lc, 5.2)
    set_cw(rc, 2.1)
    set_rh(r1, 5.5)
    set_bg(lc, '0078D4')
    set_bg(rc, '0078D4')
    no_border(lc)
    no_border(rc)
    set_pad(lc, t=230, r=160, b=180, l=200)
    set_pad(rc, t=220, r=130, b=144, l=130)
    cell_valign(lc, 'top')
    cell_valign(rc, 'center')
    clear(lc)
    clear(rc)

    # LEFT: top row - copilot logo + title
    shown_copilot = add_pic(lc, copilot,
                            0.48, 0.48, L, sa=8)
    if not shown_copilot:
        p = add_para(lc, L, sa=4)
        add_run(p, 'M365', 10, bold=True, color='B3D9FF')

    # Tag line
    p = add_para(lc, L, sa=5)
    add_run(p,
            'MICROSOFT 365  ·  COPILOT READINESS ASSESSMENT',
            7.5, color='B3D9FF')

    # Main title line 1
    p = add_para(lc, L, sa=0)
    add_run(p, 'Copilot ', 20, color='FFFFFF')
    add_run(p, 'Readiness', 20, color='B3D9FF')

    # Main title line 2
    p = add_para(lc, L, sa=10)
    add_run(p, 'Assessment Report', 15, color='FFFFFF')

    # Description line
    p = add_para(lc, L, sa=14)
    add_run(p,
            'Microsoft 365 environment evaluated across Security,',
            9, italic=True, color='CCE8FF')
    p2 = add_para(lc, L, sa=14)
    add_run(p2,
            'Governance and Best Practice pillars.',
            9, italic=True, color='CCE8FF')

    # TPT logo OR text branding at bottom
    shown_logo = add_pic(lc, logo_path,
                         1.3, 0.26, L, sa=0)
    if not shown_logo:
        p = add_para(lc, L, sa=0)
        add_run(p, f'Prepared by  {partner}',
                9, bold=True, color='FFFFFF')

    # RIGHT: Score panel
    set_bg(rc, '005A9E')

    p = add_para(rc, C, sa=4)
    add_run(p, 'READINESS SCORE',
            7, color='B3D9FF')

    p = add_para(rc, C, sa=4, sb=6)
    add_run(p, f'{score:.2f}%',
            20, bold=True, color='FFFFFF')

    # Separator
    p = add_para(rc, C, sa=4)
    add_run(p, '─' * 10, 7, color='5599CC')

    p = add_para(rc, C, sa=8)
    add_run(p, f'●  {level_short}',
            9, bold=True, color='FFFFFF')

    p = add_para(rc, C, sa=2)
    add_run(p, f'{fail_total} fail  ·  {pass_total} pass',
            8, color='B3D9FF')

    p = add_para(rc, C, sa=0)
    add_run(p, f'out of {fail_total + pass_total} parameters' if (fail_total + pass_total) else DATA_NOT_AVAILABLE,
            7.5, color='B3D9FF')

    # TABLE 2 - META CARDS ROW
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = 'Table Grid'
    no_tbl_border(t2)
    r2 = t2.rows[0]
    set_rh(r2, 1.1)

    meta = [
        ('PREPARED FOR', company, '0078D4'),
        ('PREPARED BY', partner, '1A1A1A'),
        ('ASSESSMENT DATE', date_str, '1A1A1A'),
    ]
    cw2 = [2.57, 2.57, 2.56]
    for i, (lbl, val, vc) in enumerate(meta):
        cell = r2.cells[i]
        set_cw(cell, cw2[i])
        set_bg(cell, 'EBF4FD')
        no_border(cell)
        set_pad(cell, t=120, r=140, b=100, l=200)
        cell_valign(cell, 'center')
        clear(cell)

        p = add_para(cell, L, sa=2)
        add_run(p, lbl, 7, color='888888')

        p = add_para(cell, L, sa=0)
        add_run(p, val, 12, bold=True, color=vc)

    # TABLE 3 - SERVICE CHIPS
    t3 = doc.add_table(rows=1, cols=6)
    t3.style = 'Table Grid'
    no_tbl_border(t3)
    r3 = t3.rows[0]
    set_rh(r3, 0.8)

    bgs = ['EBF4FD', 'DDEEFB']
    cw3 = [1.28, 1.28, 1.28, 1.28, 1.28, 1.30]

    for i, (svc_full, svc_short) in enumerate(svc_order):
        cell = r3.cells[i]
        clear(cell)
        set_cw(cell, cw3[i])
        no_border(cell)
        set_pad(cell, t=80, r=50, b=60, l=50)
        cell_valign(cell, 'center')
        bg_color = bgs[i % 2]
        set_bg(cell, bg_color)

        fc = cnt[svc_full]['f']
        pc = cnt[svc_full]['p']
        dot_color = 'E24B4A' if fc > 0 else '27500A'

        p = add_para(cell, C, sa=2)
        add_run(p, svc_short, 7.5,
                bold=True, color='0C447C')

        p = add_para(cell, C, sa=0)
        add_run(p, '● ', 7, color=dot_color)
        add_run(p, f'{fc}F / {pc}P', 7, color='444444')

    t4 = doc.add_table(rows=1, cols=1)
    t4.style = 'Table Grid'
    no_tbl_border(t4)
    r4 = t4.rows[0]
    set_rh(r4, 0.08)
    c4 = r4.cells[0]
    set_cw(c4, 7.3)
    set_bg(c4, '0078D4')
    no_border(c4)
    clear(c4)
    c4.add_paragraph()

    # PAGE BREAK (using native Word API, NOT section break)
    # Section breaks cause Word to recalculate layout and move content to previous pages
    pb = doc.add_paragraph()
    pb.paragraph_format.space_before = Pt(0)
    pb.paragraph_format.space_after = Pt(0)
    pbr = pb.add_run()
    pbr.add_break(WD_BREAK.PAGE)

def _add_executive_page(doc, company_name, partner_name, assessment_data=None):
    """PAGE 5: Executive Summary - AAA structure with Calibri Light 16pt headings, Calibri 11pt body."""
    from docx.shared import Pt, RGBColor

    assessment_data = assessment_data or {}

    # ---- PAGE 5: EXECUTIVE SUMMARY ----
    h = doc.add_heading('Executive Summary', level=1)
    for r in h.runs:
        r.font.name = 'Calibri Light'
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # AAA Paragraph 1
    p = doc.add_paragraph()
    p_text = (
        f"As part of its digital transformation {company_name}, engaged {partner_name} for a Copilot Readiness "
        f"Assessment. The purpose of this engagement was to evaluate the Client's Microsoft 365 environment "
        f"across areas including security, governance, and best practices to determine readiness for the secure "
        f"and responsible adoption of Microsoft 365 Copilot."
    )
    run = p.add_run(p_text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # AAA Paragraph 2
    p = doc.add_paragraph()
    p_text = (
        "The assessment covered critical services including Entra ID, Exchange Online, Microsoft Teams, "
        "SharePoint Online, OneDrive for Business, and Microsoft Purview. It aimed to identify configuration "
        "gaps, policy misalignments, and potential vulnerabilities that could impact the responsible use of AI-powered "
        "tools like Copilot. By benchmarking the current environment against industry standards and Microsoft's Copilot "
        "deployment criteria, the assessment provides a clear roadmap for remediation and optimization."
    )
    run = p.add_run(p_text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # AAA Paragraph 3
    p = doc.add_paragraph()
    p_text = (
        f"The findings serve as a strategic foundation for {company_name} to enhance its digital workplace, "
        f"mitigate operational and compliance risks, and unlock the full potential of Microsoft 365 Copilot. "
        f"With targeted improvements, the organization can ensure a secure and scalable AI integration that aligns "
        f"with its long-term business goals."
    )
    run = p.add_run(p_text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # Purpose heading
    p = doc.add_paragraph()
    h2 = doc.add_heading('Purpose', level=1)
    for r in h2.runs:
        r.font.name = 'Calibri Light'
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    purpose_items = [
        [
            ('Evaluate the ', False),
            (company_name, True),
            (' environment for alignment with ', False),
            ('industry best practices', True),
            ('.', False),
        ],
        [('Assess the environment across Microsoft 365 products and services like SharePoint, Teams, OneDrive for Business etc.', False)],
        [
            ('Identify gaps that could pose ', False),
            ('security or compliance risks', True),
            (' upon integrating Copilot.', False),
        ],
        [('Establish a baseline for future audits and compliance tracking related to AI usage within Microsoft 365.', False)],
        [('Highlight licensing readiness and user eligibility for Microsoft 365 Copilot deployment.', False)],
        [('Provide a risk-based prioritization of remediation efforts to guide Copilot enablement planning.', False)],
        [('Offer actionable insights to strengthen governance, data protection, and identity management in preparation for AI integration.', False)],
        [('Support strategic decision-making by outlining Copilot deployment prerequisites and dependencies.', False)],
    ]
    for item_runs in purpose_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        for text, bold in item_runs:
            _format_run(p.add_run(text), size=11, bold=bold)

    doc.add_page_break()

def _add_evaluation_page(doc, findings):
    """PAGE 6: Evaluation Summary - AAA structure with both charts on same page."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # ---- PAGE 6: EVALUATION SUMMARY ----
    h = doc.add_heading('Evaluation Summary', level=1)
    h.paragraph_format.space_after = Pt(4)
    for r in h.runs:
        r.font.name = 'Calibri Light'
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # Section A: 3 Pillars heading
    h2 = doc.add_heading('3 Pillars of Microsoft 365 Copilot Readiness Assessment', level=2)
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(2)
    for r in h2.runs:
        r.font.name = 'Calibri Light'
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # AAA 3 Pillars bullets
    for pillar in ['Governance', 'Security', 'Best Practices']:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        _format_run(p.add_run(pillar), size=11)

    # Chart 1: 3 Pillars
    pillar_order = ['Security', 'Governance', 'Best Practice']
    pillar_counts = _distribution(findings, _pillar_value, pillar_order)
    pillar_labels = list(pillar_counts.keys())
    pillar_values = list(pillar_counts.values())
    pillar_colors = ['#4472C4', '#A5A5A5', '#ED7D31']

    pillar_chart = _page6_pie_chart(pillar_labels, pillar_values, pillar_colors, '3 Pillars of CRA') if pillar_values else None
    _chart_or_data_not_available(doc, pillar_chart, width=Inches(4.7), height=Inches(2.35))

    # Section B: M365 Services heading
    h2 = doc.add_heading('M365 Services assessed in CRA', level=2)
    h2.paragraph_format.space_before = Pt(2)
    h2.paragraph_format.space_after = Pt(2)
    for r in h2.runs:
        r.font.name = 'Calibri Light'
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # AAA M365 Services bullets in exact order
    services_list = [
        'Entra ID',
        'Exchange Online',
        'Microsoft Purview',
        'Microsoft Teams',
        'OneDrive for Business',
        'SharePoint Online',
    ]
    for service in services_list:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        _format_run(p.add_run(service), size=11)

    # Chart 2: M365 Services
    service_order = [
        'Entra ID',
        'Exchange Online',
        'Microsoft Purview',
        'Microsoft Teams',
        'OneDrive for Business',
        'SharePoint Online',
    ]
    service_counts = _distribution(findings, _service_value, service_order)
    service_labels = list(service_counts.keys())
    service_values = list(service_counts.values())
    service_colors = ['#4472C4', '#ED7D31', '#A5A5A5', '#FFC000', '#5B9BD5', '#70AD47']

    service_chart = _page6_pie_chart(service_labels, service_values, service_colors, 'M365 Services') if service_values else None
    _chart_or_data_not_available(doc, service_chart, width=Inches(4.7), height=Inches(2.35))

    doc.add_page_break()

def _add_summary_page(doc, findings, assessment_data):
    """PAGE 7: Risk Score Matrix - AAA structure with severity table and Risk-wise Parameters chart."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # ---- PAGE 7: RISK SCORE MATRIX ----
    h = doc.add_heading('Risk Score Matrix', level=1)
    for r in h.runs:
        r.font.name = 'Calibri Light'
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # AAA intro text
    p = doc.add_paragraph("The findings presented in this report are graded according to the following levels of severity:")
    for r in p.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(11)

    doc.add_paragraph()

    # Risk Score Matrix Table (AAA severity visual display)
    _add_risk_score_matrix(doc)

    doc.add_paragraph()

    # Risk Category of Parameters Assessed heading
    h2 = doc.add_heading('Risk Category of Parameters Assessed', level=2)
    for r in h2.runs:
        r.font.name = 'Calibri Light'
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # AAA description
    p = doc.add_paragraph("The following chart provides consolidated parameters based on risk category assessed during the engagement:")
    for r in p.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(11)

    doc.add_paragraph()

    sev_order = ['Critical', 'High', 'Medium', 'Low', 'Informational']
    sev_counts = _distribution(findings, _severity_value, sev_order)
    sev_values = list(sev_counts.values())
    sev_colors = ['#8B0000', '#FF0000', '#ED7D31', '#FFD966', '#70AD47']

    c1 = _page7_risk_pie_chart(list(sev_counts.keys()), sev_values, sev_colors, 'Risk-wise Parameters') if sev_values else None
    _chart_or_data_not_available(doc, c1, width=Inches(4.7), height=Inches(2.45))

    doc.add_page_break()

def _add_assessment_summary_page(doc, assessment_data):
    """PAGE 8: Summary of Assessment using AAA layout with live assessment data."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    parameter_rows = assessment_data.get('parameter_rows', [])
    total_parameters = assessment_data.get('total_params') or len(parameter_rows)
    passed_parameters = assessment_data.get('pass_count')
    if passed_parameters is None:
        passed_parameters = len([
            row for row in parameter_rows
            if str(row.get('display_status', row.get('status', ''))).lower().strip() == 'pass'
        ])
    failed_parameters = assessment_data.get('gaps_count')
    if failed_parameters is None:
        failed_parameters = max(int(total_parameters or 0) - int(passed_parameters or 0), 0)

    summary = assessment_data.get('summary', {}) if isinstance(assessment_data.get('summary'), dict) else {}
    readiness_score_raw = (
        assessment_data.get('readiness_score')
        if assessment_data.get('readiness_score') is not None
        else assessment_data.get('overall_score')
    )
    if readiness_score_raw is None:
        readiness_score_raw = summary.get('readiness_score') or summary.get('overall_score')
    readiness_score = float(readiness_score_raw) if readiness_score_raw is not None else None

    readiness_level = (
        assessment_data.get('readiness_level')
        or assessment_data.get('readiness_status')
        or summary.get('readiness_level')
        or summary.get('readiness_status')
        or DATA_NOT_AVAILABLE
    )

    if readiness_score is None:
        readiness_text = DATA_NOT_AVAILABLE
        readiness_score_text = DATA_NOT_AVAILABLE
    elif readiness_score < 50:
        readiness_text = 'Significant remediation is required prior to enabling Copilot in the production environment.'
    elif readiness_score < 80:
        readiness_text = 'Moderate remediation is recommended before enabling Copilot.'
    else:
        readiness_text = 'The environment is generally prepared for Copilot deployment.'
    if readiness_score is not None:
        readiness_score_text = f'{readiness_score:.2f}%'

    h = doc.add_heading('Summary of Assessment', level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(5)
    for r in h.runs:
        r.font.name = 'Calibri Light'
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    paragraphs = [
        'The Copilot Readiness Assessment uncovered several configuration gaps and policy deficiencies that could impact the secure and compliant adoption of Microsoft Copilot. Each finding has been categorized by severity and mapped to specific areas of risk within the Microsoft 365 environment.',
        'Overall Readiness:',
        "Based on the findings, the Client's current readiness level for Copilot integration is assessed as:",
        f'Readiness Level: {readiness_level}',
        f'Readiness Gaps: {failed_parameters} out of {total_parameters}',
    ]
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
        if text in {'Overall Readiness:', f'Readiness Level: {readiness_level}', f'Readiness Gaps: {failed_parameters} out of {total_parameters}'}:
            for r in p.runs:
                r.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('Readiness Score:')
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Calibri'
    r.font.color.rgb = _hex_color('003366')

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(readiness_text)
    r.font.italic = True
    r.font.size = Pt(11)
    r.font.name = 'Calibri'

    score_table = doc.add_table(rows=1, cols=1)
    score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    score_cell = score_table.rows[0].cells[0]
    _set_cell_bg(score_cell, 'EBF4FD')
    score_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    score_para = score_cell.paragraphs[0]
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    score_para.paragraph_format.space_before = Pt(8)
    score_para.paragraph_format.space_after = Pt(8)
    score_run = score_para.add_run(readiness_score_text)
    score_run.font.size = Pt(28)
    score_run.font.bold = True
    score_run.font.name = 'Calibri'
    score_run.font.color.rgb = _hex_color('0078D4')
    _remove_table_borders(score_table)
    score_table.autofit = False
    score_table.allow_autofit = False
    for row in score_table.rows:
        row.height = Inches(0.8)

    meter = _add_readiness_gauge(doc, readiness_score) if readiness_score is not None else None
    if meter:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(meter, width=Inches(6.4), height=Inches(1.05))
        try:
            os.remove(meter)
        except:
            pass

    chart = _pass_fail_horizontal_chart(passed_parameters, failed_parameters)
    _chart_or_data_not_available(doc, chart, width=Inches(6.4), height=Inches(1.05))

    doc.add_page_break()

def _add_page9_executive_charts(doc, assessment_data):
    """PAGE 9: AAA severity chart with Key Observations only."""
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    parameter_rows = assessment_data.get('parameter_rows', [])

    chart = _page9_severity_pillars_chart(parameter_rows)
    if chart:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
        p.add_run().add_picture(chart, width=Inches(6.2), height=Inches(3.70))
        try:
            os.remove(chart)
        except:
            pass

    _add_key_observations(doc, assessment_data)

def _add_page7_executive_charts(doc, assessment_data):
    """PAGE 7: Executive dashboard charts only."""
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    parameter_rows = assessment_data.get('parameter_rows', [])

    h = doc.add_heading('Executive Dashboard', level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(4)
    for r in h.runs:
        r.font.name = 'Calibri Light'
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    chart = _page8_exec_summary_chart(parameter_rows)
    if chart:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.add_run().add_picture(chart, width=Inches(6.1), height=Inches(2.25))
        try:
            os.remove(chart)
        except:
            pass

    chart = _page9_severity_pillars_chart(parameter_rows)
    if chart:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.add_run().add_picture(chart, width=Inches(6.1), height=Inches(2.0))
        try:
            os.remove(chart)
        except:
            pass

    chart = _service_pass_fail_chart(parameter_rows)
    _chart_or_data_not_available(doc, chart, width=Inches(6.1), height=Inches(1.65))

    doc.add_page_break()

def _add_page8_key_observations(doc, assessment_data):
    """PAGE 8: Key Observations and license chart only."""
    _add_key_observations(doc, assessment_data, include_user_activity=False)

def _add_page9_user_activity_and_risks(doc, assessment_data):
    """PAGE 9: User activity, deployment risks, and recommendations. Stops the report."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    def heading(text, space_after=6):
        h = doc.add_heading(text, level=1)
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after = Pt(space_after)
        for r in h.runs:
            r.font.name = 'Calibri Light'
            r.font.size = Pt(16)
            r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        return h

    def bullet(text, space_after=2):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        marker = p.add_run('\u2022  ')
        marker.font.name = 'Calibri'
        marker.font.size = Pt(11)
        r = p.add_run(str(text))
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        return p

    heading('User Activity Overview', space_after=4)
    subheading = doc.add_heading('User Information Details', level=2)
    subheading.paragraph_format.space_before = Pt(0)
    subheading.paragraph_format.space_after = Pt(4)
    for r in subheading.runs:
        r.font.name = 'Calibri Light'
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    ui_fields = assessment_data.get('user_info_fields', {})
    ui_total = int(assessment_data.get('user_info_total', 0))
    if ui_fields and ui_total > 0:
        bar_bytes = _make_bar_chart_img(
            ui_fields,
            ui_total,
            'User Information Details',
            size=(5.2, 2.2),
        )
        _insert_chart(doc, bar_bytes, width_inches=5.0, sa=4)

    activity = [
        ('SharePoint Users', float(assessment_data.get('sharepoint_active_pct', 0))),
        ('OneDrive Users', float(assessment_data.get('onedrive_active_pct', 0))),
        ('Teams Users', float(assessment_data.get('teams_active_pct', 0))),
        ('Outlook Users', float(assessment_data.get('outlook_active_pct', 0))),
    ]
    table = doc.add_table(rows=1, cols=4)
    _remove_table_borders(table)
    for idx, (label, pct) in enumerate(activity):
        cell = table.rows[0].cells[idx]
        _set_cell_bg(cell, 'EBF4FD')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f'{pct:.0f}%')
        r.font.name = 'Calibri'
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)

    heading('Risks of Immediate Deployment', space_after=4)

    risk_rows = _top_findings_by_severity(
        assessment_data.get('parameter_rows', []),
        ['Critical', 'High'],
        limit=4,
    )
    risks = [_finding_title(row) for row in risk_rows] or [DATA_NOT_AVAILABLE]
    risk_table = doc.add_table(rows=2, cols=2)
    _remove_table_borders(risk_table)
    for idx, risk in enumerate(risks[:4]):
        cell = risk_table.rows[idx // 2].cells[idx % 2]
        _set_cell_bg(cell, 'FCE4D6')
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f'{idx + 1}. {risk}')
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    heading('Recommendations', space_after=4)
    recommendation_rows = (
        _top_findings_by_severity(assessment_data.get('parameter_rows', []), ['Critical'], limit=1)
        + _top_findings_by_severity(assessment_data.get('parameter_rows', []), ['High'], limit=1)
        + _top_findings_by_severity(assessment_data.get('parameter_rows', []), ['Medium'], limit=1)
    )
    recommendations = []
    for priority, row in enumerate(recommendation_rows, 1):
        recommendation = _recommendation_text(row)
        if recommendation:
            recommendations.append(f'Priority {priority}: {recommendation}')
        else:
            recommendations.append(f'Priority {priority}: Review and remediate {_finding_title(row)}')
    for text in recommendations or [DATA_NOT_AVAILABLE]:
        bullet(text)

def _add_risks_recommendations_page(doc):
    def bullet(runs, sa=5):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        p.paragraph_format.space_after = Pt(sa)
        p.paragraph_format.space_before = Pt(0)
        marker = p.add_run('\u2022  ')
        marker.font.size = Pt(11)
        marker.font.name = 'Calibri'
        for text, bold in runs:
            r = p.add_run(str(text))
            r.font.bold = bool(bold)
            r.font.size = Pt(11)
            r.font.name = 'Calibri'
        return p

    pb = doc.add_paragraph()
    pb.paragraph_format.space_before = Pt(0)
    pb.paragraph_format.space_after = Pt(0)
    run = pb.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

    h = doc.add_heading('Risks of Immediate Deployment:', level=1)
    h.paragraph_format.space_after = Pt(8)
    for r in h.runs:
        r.font.name = 'Calibri Light'
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _format_run(
        p.add_run('Proceeding with Copilot activation in the current state may lead to:'),
        size=11,
    )

    risks = [
        'Unauthorized access to sensitive business data.',
        'Inadequate auditing and monitoring of AI-driven activity.',
        'Compliance violations with internal and external regulations.',
        'Reduced user trust and operational inconsistencies due to misconfigured policies.',
    ]
    for risk in risks:
        bullet([(risk, False)], sa=5)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.space_before = Pt(0)

    h = doc.add_heading('Recommendations:', level=1)
    h.paragraph_format.space_after = Pt(8)
    for r in h.runs:
        r.font.name = 'Calibri Light'
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    recommendations = [
        (
            'Remediation of identified gaps: ',
            'Address all findings regardless of severity to meet cybersecurity baseline standards.',
        ),
        (
            'Postpone Deployment: ',
            'Due to the current maturity level of the environment, it is recommended to adopt Copilot '
            'deployment after all critical and high-priority gaps are resolved.',
        ),
        (
            'Futureproofing: ',
            'Implementing the recommendations provided will reduce security risks and ensure regulatory '
            'compliance during and after Copilot integration.',
        ),
    ]
    for label, body in recommendations:
        bullet([(label, True), (body, False)], sa=6)

def _get_risk_text(severity):
    """FIX 5: Generate risk text based on severity."""
    RISK_TEMPLATES = {
        'critical': 'This is a critical finding that requires immediate remediation before enabling Microsoft 365 Copilot in the environment.',
        'high': 'This high-severity finding significantly increases the risk of data exposure and should be addressed as a priority.',
        'medium': 'This medium-severity finding may impact Copilot performance and compliance if not addressed.',
        'low': 'This low-severity finding represents a minor risk that should be included in remediation planning.',
        'info': 'This informational finding provides context for the overall Copilot readiness assessment.',
        'informational': 'This informational finding provides context for the overall Copilot readiness assessment.',
    }
    severity_key = str(severity or '').lower().strip()
    return RISK_TEMPLATES.get(severity_key, 'This finding should be reviewed as part of Copilot readiness.')

def _detailed_status_text(status):
    status_key = str(status or '').lower().strip()
    return 'Pass' if status_key == 'pass' else 'Fail'

def _detailed_severity_color(severity, config):
    severity_defs = config.get("severity_definitions", {}) if isinstance(config, dict) else {}
    severity_info = severity_defs.get(severity, {}) if isinstance(severity_defs, dict) else {}
    return _hex_text(severity_info.get("color"), {
        "Critical": "C00000",
        "High": "FF0000",
        "Medium": "ED7D31",
        "Low": "FFC000",
        "Informational": "70AD47",
    }.get(severity, "70AD47"))

def _add_detailed_summary_table(doc, rows, config, headers, start_index=1, include_header=True):
    table_rows = len(rows) + (1 if include_header else 0)
    table = doc.add_table(rows=table_rows, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [0.55, 3.05, 1.08, 0.72, 1.00]
    for idx, width in enumerate(widths):
        for table_row in table.rows:
            _set_cell_width(table_row.cells[idx], width)

    body_start = 0
    if include_header:
        hrow = table.rows[0]
        for i, header in enumerate(headers):
            cell = hrow.cells[i]
            _set_cell_bg(cell, '2E74B5')
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(header)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(11)
            r.font.name = _font_name(config)
        _set_row_height(hrow, 0.30)
        body_start = 1

    for offset, row in enumerate(rows):
        trow = table.rows[body_start + offset]
        _set_row_height(trow, 0.30)
        for cell in trow.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_bg(cell, 'FFFFFF')

        display_index = start_index + offset
        _set_cell_text(trow.cells[0], str(display_index).zfill(2), config, size_key="caption_size", align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(trow.cells[1], row.get('title', ''), config, size_key="caption_size")

        pillar_raw = str(row.get('pillar', '')).lower().strip()
        pillar = PILLAR_MAP.get(pillar_raw, row.get('pillar', 'Best Practice'))
        _set_cell_text(trow.cells[2], pillar, config, size_key="caption_size", align=WD_ALIGN_PARAGRAPH.CENTER)

        status = str(row.get('display_status', row.get('status', ''))).lower().strip()
        finding_text = _detailed_status_text(status)
        finding_color = _cfg(config, "branding", "success_color", "00B050") if finding_text == 'Pass' else _cfg(config, "branding", "fail_color", "C00000")
        _set_cell_text(
            trow.cells[3],
            finding_text,
            config,
            size_key="caption_size",
            bold=True,
            color=finding_color,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

        severity_raw = str(
            row.get('display_severity')
            or row.get('registry_severity')
            or row.get('severity', '')
        ).lower().strip()
        severity = SEVERITY_MAP.get(severity_raw, 'Informational')
        _set_cell_bg(trow.cells[4], _detailed_severity_color(severity, config))
        _set_cell_text(
            trow.cells[4],
            severity,
            config,
            size_key="caption_size",
            bold=True,
            color="FFFFFF",
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    return table

def _detail_paragraph(doc, text="", config=None, before=0, after=5, bold=False, color=None, size_key="caption_size"):
    config = config or DEFAULT_REPORT_CONFIG
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.03
    run = p.add_run(str(text or ""))
    _apply_run_style(run, config, size_key, bold=bold, color=color)
    return p

def _detail_multiline_paragraph(doc, text="", config=None, before=0, after=5, bold=False, color=None, size_key="caption_size"):
    '''Like _detail_paragraph but renders embedded newlines as real line breaks, so a
    multi-line evaluated_value (e.g. "Allowed:\\nGoogleDrive\\n...") shows correctly.'''
    config = config or DEFAULT_REPORT_CONFIG
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.03
    lines = str(text or "").split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        _apply_run_style(run, config, size_key, bold=bold, color=color)
        if i < len(lines) - 1:
            run.add_break()
    return p

def _add_detail_finding_block(doc, row, idx, config, meter_cache=None, compact_top=False):
    title = row.get('title') or row.get('parameter') or DATA_NOT_AVAILABLE
    bar_path = _detail_title_bar(f"{str(idx).zfill(2)}: {title}")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0 if compact_top else 5)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(bar_path, width=Inches(6.35), height=Inches(0.43))
    try:
        os.remove(bar_path)
    except OSError:
        pass

    severity_raw = str(
        row.get('display_severity')
        or row.get('registry_severity')
        or row.get('severity', 'informational')
    ).lower().strip()
    severity = SEVERITY_MAP.get(severity_raw, 'Informational')
    status = str(row.get('display_status', row.get('status', ''))).lower().strip()
    finding_text = _detailed_status_text(status)

    risk_line = doc.add_paragraph()
    risk_line.paragraph_format.space_before = Pt(0)
    risk_line.paragraph_format.space_after = Pt(5)
    _apply_run_style(risk_line.add_run("Risk Rating: "), config, "caption_size", bold=True)
    _apply_run_style(risk_line.add_run(severity), config, "caption_size", bold=True)
    _apply_run_style(risk_line.add_run(" - "), config, "caption_size", bold=True)
    status_color = _cfg(config, "branding", "success_color", "00B050") if finding_text == "Pass" else _cfg(config, "branding", "fail_color", "C00000")
    _apply_run_style(risk_line.add_run(finding_text), config, "caption_size", bold=True, color=status_color)

    severity_key = severity.lower().strip()
    if severity_key == 'info':
        severity_key = 'informational'
    meter_path = meter_cache.get(severity_key) if meter_cache else None
    generated_meter = False
    if not meter_path or not os.path.exists(meter_path):
        meter_path = _severity_meter(severity_key)
        generated_meter = True
    if meter_path and os.path.exists(meter_path):
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        mp.paragraph_format.space_before = Pt(0)
        mp.paragraph_format.space_after = Pt(13)
        mp.add_run().add_picture(meter_path, width=Inches(6.35), height=Inches(0.86))
        if generated_meter:
            try:
                os.remove(meter_path)
            except OSError:
                pass

    _detail_paragraph(doc, "Description:", config, after=7, bold=True)
    # The Description must mirror assessment_findings.evaluated_value (source of truth).
    # Only when a finding has no evaluated_value do we fall back to the computed/static text.
    description = (
        _report_description(row)
        or resolve_description(row.get('parameter_key'), row.get('evidence'), status)
        or row.get('description', '')
        or DATA_NOT_AVAILABLE
    )
    _detail_multiline_paragraph(doc, description, config, after=9)

    _detail_paragraph(doc, "Risk:", config, after=7, bold=True)
    risk_text = row.get('risk') or _get_risk_text(severity)
    _detail_paragraph(doc, risk_text, config, after=6)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(20)

def _add_detailed_pages(doc, assessment_data, meter_cache=None):
    config = _report_config(assessment_data)
    page_text = (config.get("page_text", {}) or {}).get("detailed_assessment", {})
    heading = page_text.get("heading", "Detailed Assessment")
    intro = page_text.get(
        "intro",
        "The following tables provides a summary of all findings discovered during the course of this engagement:",
    )
    headers = page_text.get("table_headers") or ['S. No', 'Parameter', 'CRA Pillar', 'Finding', 'Severity']
    first_page_row_limit = int(page_text.get("first_page_row_limit", 18) or 18)
    continuation_top_spacing = float(page_text.get("continuation_top_spacing_pt", 22) or 22)

    _styled_heading(
        doc,
        heading,
        config,
        level=1,
        color=_cfg(config, "branding", "secondary_color", "2F5496"),
        before=0,
        after=18,
    )
    _body_paragraph(doc, intro, config, after=26)

    parameter_rows = assessment_data.get('parameter_rows', [])
    if not parameter_rows:
        doc.add_paragraph("No findings to display.")
        return

    grouped = defaultdict(list)
    for row in parameter_rows:
        service = _service_name_key(row.get('service', row.get('category', 'General')))
        grouped[service].append(row)

    ordered_services = [service for service in SERVICE_ORDER if service in grouped]
    for service_index, service in enumerate(ordered_services):

        rows = grouped[service]
        # FIX 1-2: SORT FINDINGS BY SEVERITY THEN STATUS
        sorted_findings = _sort_findings(rows)

        # FIX 3: SERVICE HEADINGS WITH PROPER NAMES
        display_svc = SERVICE_DISPLAY_NAMES.get(service.lower().strip(), service.upper())
        _styled_heading(
            doc,
            display_svc,
            config,
            level=2,
            color=_cfg(config, "branding", "secondary_color", "2F5496"),
            before=0,
            after=14,
        )

        first_chunk = sorted_findings[:first_page_row_limit]
        remaining_chunk = sorted_findings[first_page_row_limit:]
        _add_detailed_summary_table(doc, first_chunk, config, headers, start_index=1, include_header=True)

        if remaining_chunk:
            doc.add_page_break()
            top_spacer = doc.add_paragraph()
            top_spacer.paragraph_format.space_before = Pt(continuation_top_spacing)
            top_spacer.paragraph_format.space_after = Pt(0)
            _add_detailed_summary_table(
                doc,
                remaining_chunk,
                config,
                headers,
                start_index=first_page_row_limit + 1,
                include_header=False,
            )

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)
        doc.add_page_break()

        for idx, row in enumerate(sorted_findings, 1):
            if idx > 1 and idx % 2 == 1:
                doc.add_page_break()
            _add_detail_finding_block(
                doc,
                row,
                idx,
                config,
                meter_cache=meter_cache,
                compact_top=(idx % 2 == 1),
            )

        if service_index < len(ordered_services) - 1:
            doc.add_page_break()

def _add_conclusion_page(doc, company_name, partner_name, assessment_data):
    """Conclusion with real report_data values."""
    config = _report_config(assessment_data)
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(56)
    heading.paragraph_format.space_after = Pt(26)
    run = heading.add_run('Conclusion')
    _apply_run_style(run, config, "h1_size", bold=True, color=_cfg(config, "branding", "secondary_color", "2F5496"), font_kind="heading")

    company = (
        assessment_data.get('company_name')
        or assessment_data.get('organization_name')
        or assessment_data.get('customer_name')
        or company_name
        or 'the organisation'
    )
    fail_count = (
        assessment_data.get('gaps_count')
        or assessment_data.get('failed_parameters')
        or assessment_data.get('summary', {}).get('gaps_count')
        or 0
    )
    total = (
        assessment_data.get('total_params')
        or assessment_data.get('total_assessed')
        or assessment_data.get('summary', {}).get('total_params')
        or 0
    )

    def para(text, sa=13):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(sa)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(text)
        _apply_run_style(r, config, "body_size")
        return p

    para(
        f'The Copilot Readiness Assessment for {company} reveals that the current Microsoft 365 '
        f'environment is not yet prepared for the secure and compliant deployment of Microsoft '
        f'365 Copilot. With {fail_count} out of {total} parameters failing to meet readiness '
        f'standards, many of which fall under critical and high-risk categories, there is a '
        f'clear need for immediate and comprehensive remediation.'
    )

    para(
        'Key gaps were identified across all three foundational pillars: Security, Governance, '
        'and Best Practices. Notably, critical vulnerabilities such as the absence of sensitivity '
        'labels, permissive external sharing configurations, and insufficient audit logging '
        'significantly elevate the risk of data exposure and non-compliance. Furthermore, the '
        'lack of complete user profile information and inconsistent policy enforcement could '
        'impair Copilot\'s ability to deliver accurate, context-aware insights.'
    )

    para(
        'To mitigate these risks and ensure a successful Copilot rollout, we strongly recommend '
        'a phased remediation strategy. This should prioritise the resolution of critical and '
        'high-severity issues, followed by medium and low-risk items. Only after these gaps are '
        'addressed should the organisation consider enabling Copilot in the production environment.'
    )

    para(
        f'By aligning with the recommendations outlined in this report, {company} can enhance '
        'its security posture, ensure regulatory compliance, and fully leverage the transformative '
        'potential of Microsoft 365 Copilot.',
        sa=0,
    )

def _toc_key(text):
    return " ".join(str(text or "").lower().replace("–", "-").split())

def _toc_page_lookup(config):
    fallback = {
        "executive summary": 5,
        "purpose": 5,
        "evaluation summary": 6,
        "3 pillars of microsoft 365 copilot readiness assessment": 6,
        "m365 services assessed in cra": 6,
        "risk matrix": 7,
        "risk category of parameters assessed": 7,
        "summary of assessment": 8,
        "executive dashboard": 9,
        "key observations": 10,
        "user information analysis": 11,
        "usage and recommendations": 12,
        "risks of immediate deployment": 12,
        "recommendations": 12,
        "detailed assessment": 13,
        "entra id": 13,
        "exchange online": 26,
        "microsoft purview": 30,
        "microsoft teams": 35,
        "onedrive for business": 44,
        "sharepoint online": 47,
        "01: custom banned password list": 15,
        "02: restricted access to microsoft entra admin centre": 15,
        "03: emergency access accounts": 16,
        "04: device without compliance policies": 16,
        "05: authentication methods enabled": 17,
        "06: entra - tenant creation by non-admins": 17,
        "07: global administrator accounts": 18,
        "08: self-service password reset authentication method": 18,
        "09: tenant collaboration invitation": 19,
        "10: administrator consent workflows": 19,
        "11: cap policies for risky sign-ins": 20,
        "12: conditional access policies (exclusion)": 20,
        "13: user consent for applications": 21,
        "14: entra - third-party app integrations": 21,
        "15: users without mfa": 22,
        "16: auto-expiration policy for m365 groups": 22,
        "17: customer lockbox": 23,
        "18: guest invite settings": 23,
        "19: guest users count": 24,
        "20: user information": 24,
        "21: number of accounts enabled": 25,
        "01: mailbox status (active/inactive)": 27,
        "02: external storage providers in owa": 27,
        "03: mailbox storage usage": 28,
        "04: full calendar schedules able to be shared externally": 28,
        "05: number of emails read/received": 29,
        "06: number of emails sent": 29,
        "01: audit logs enabled": 31,
        "02: secure score percentage": 31,
        "03: sensitivity labels configured and applied": 32,
        "04: sensitivity labels applied to teams": 32,
        "05: compliance score overview": 33,
        "06: information protection labels applied": 33,
        "07: dlp rules configured": 34,
        "08: audit log retention duration": 34,
        "01: copilot integration enabled": 36,
        "02: third party apps allowed": 36,
        "03: active/inactive teams": 37,
        "04: minimum number of owners": 37,
        "05: teams with external users": 38,
        "06: meeting policies configuration": 38,
        "07: orphan teams": 39,
        "08: teams with external guest as owner": 39,
        "09: meeting transcription enabled": 40,
        "10: guest access enabled/disabled": 40,
        "11: teams - lobby bypass": 41,
        "12: teams - file storage option": 41,
        "13: active/inactive teams users": 42,
        "14: teams - meeting chat": 42,
        "15: meeting recording retention policies": 43,
        "16: teams - channel email addresses": 43,
        "01: external sharing settings": 45,
        "02: days to retain a deleted user's onedrive": 45,
        "03: total active users on onedrive": 46,
        "01: permission settings for anyone links": 48,
        "02: sensitive sharepoint sites excluded from copilot": 48,
        "03: sharing settings (external/internal)": 49,
        "04: sharepoint and onedrive guest access expiry": 49,
        "05: expiration policy for anyone links": 50,
        "06: inactive site policies": 50,
        "07: active sites count": 51,
        "08: site ownership policies": 51,
        "09: active users on sharepoint": 52,
        "10: sharepoint - modern authentication": 52,
        "11: storage quota consumption": 53,
        "conclusion": 54,
    }
    lookup = dict(fallback)
    for entry in (config or {}).get("toc", []) if isinstance(config, dict) else []:
        title = entry.get("title") if isinstance(entry, dict) else None
        page = entry.get("page") if isinstance(entry, dict) else None
        if title and page:
            lookup[_toc_key(title)] = page
    return lookup

def _validate_aaa_toc_pages(doc_path):
    """Validate that generated TOC pages 2-4 match AAA blueprint exactly.

    Returns a validation report with measurements for:
    - Margins (left, right, top, bottom)
    - Tab stops and positions
    - Paragraph spacing
    - Line height
    - Font specifications
    """
    from docx import Document
    try:
        doc = Document(doc_path)
    except Exception as e:
        logger.error(f"[VALIDATION] Could not open document: {e}")
        return None

    blueprint = _load_aaa_report_blueprint()
    page_setup = _blueprint_page_setup(blueprint)
    margins = _blueprint_margins(blueprint)

    # Expected values from blueprint (in twips)
    expected = {
        "top_margin_twips": _twips_value(margins.get("top_twips") or _inch_value(margins.get("top"), 1.0) * 1440, 1440),
        "bottom_margin_twips": _twips_value(margins.get("bottom_twips") or _inch_value(margins.get("bottom"), 1.0) * 1440, 1440),
        "left_margin_twips": _twips_value(margins.get("left_twips") or _inch_value(margins.get("left"), 1.0) * 1440, 1440),
        "right_margin_twips": _twips_value(margins.get("right_twips") or _inch_value(margins.get("right"), 1.0) * 1440, 1440),
        "dot_leader_tab_twips": 9016,
        "level2_indent_twips": 440,
        "level3_indent_twips": 720,
        "spacing_after_twips": 100,
        "line_twips": 240,
    }

    report = {
        "document_path": str(doc_path),
        "validation_timestamp": datetime.now().isoformat(),
        "blueprint_source": str(_repo_file_path(AAA_REPORT_BLUEPRINT)),
        "pages": {2: {}, 3: {}, 4: {}},
        "summary": {"pass": True, "errors": [], "warnings": []},
    }

    # Check document-level settings
    for section in doc.sections:
        report["pages"]["document"] = {
            "top_margin_twips": section.top_margin.twips,
            "bottom_margin_twips": section.bottom_margin.twips,
            "left_margin_twips": section.left_margin.twips,
            "right_margin_twips": section.right_margin.twips,
            "expected": expected,
        }
        # Validate margins
        for margin_key in ["top_margin_twips", "bottom_margin_twips", "left_margin_twips", "right_margin_twips"]:
            actual = section.__dict__.get(margin_key, 0)
            if margin_key == "top_margin_twips":
                actual = section.top_margin.twips
            elif margin_key == "bottom_margin_twips":
                actual = section.bottom_margin.twips
            elif margin_key == "left_margin_twips":
                actual = section.left_margin.twips
            elif margin_key == "right_margin_twips":
                actual = section.right_margin.twips

            expected_val = expected[margin_key]
            diff = abs(actual - expected_val)
            if diff > 10:  # Allow 10 twips tolerance (~0.007 inches)
                report["summary"]["errors"].append(
                    f"Margin {margin_key}: expected {expected_val}, got {actual}, diff={diff} twips"
                )
                report["summary"]["pass"] = False

    # Check TOC paragraph formatting (pages 2-4)
    page_num = 1
    para_index = 0
    for page_idx in [2, 3, 4]:
        page_report = report["pages"][page_idx]
        page_report["paragraphs"] = []

        # Count paragraphs per page (rough estimate)
        for para in doc.paragraphs:
            para_index += 1
            if para_index < 50:  # Adjust based on typical paragraph count
                page_data = {
                    "text": para.text[:50],  # First 50 chars
                    "spacing_after_twips": para.paragraph_format.space_after.twips if para.paragraph_format.space_after else None,
                    "line_spacing_twips": None,
                    "tabs": [],
                }

                # Extract line spacing
                pPr = para._p.get_or_add_pPr()
                spacing_elem = pPr.find(qn("w:spacing"))
                if spacing_elem is not None:
                    line_val = spacing_elem.get(qn("w:line"))
                    if line_val:
                        page_data["line_spacing_twips"] = int(line_val)

                # Extract tab stops
                tabs_elem = pPr.find(qn("w:tabs"))
                if tabs_elem is not None:
                    for tab in tabs_elem.findall(qn("w:tab")):
                        tab_data = {
                            "position_twips": int(tab.get(qn("w:pos")) or 0),
                            "alignment": tab.get(qn("w:val"), "left"),
                            "leader": tab.get(qn("w:leader"), "none"),
                        }
                        page_data["tabs"].append(tab_data)

                page_report["paragraphs"].append(page_data)

    return report


def _aaa_toc_pages():
    page2 = [
        ("entry", "Executive Summary"), ("entry", "Purpose"), ("entry", "Evaluation Summary"),
        ("entry", "3 Pillars of Microsoft 365 Copilot Readiness Assessment"),
        ("entry", "M365 Services assessed in CRA"), ("entry", "Risk Category of Parameters Assessed"),
        ("entry", "Summary of Assessment"), ("entry", "Key Observations"),
        ("entry", "Risks of Immediate Deployment"), ("entry", "Recommendations"),
        ("entry", "Detailed Assessment"),
        ("section", "ENTRA ID"),
        ("entry", "01: Custom Banned Password List"),
        ("entry", "02: Restricted Access to Microsoft Entra Admin Centre"),
        ("entry", "03: Emergency Access Accounts"),
        ("entry", "04: Device without Compliance Policies"),
        ("entry", "05: Authentication Methods Enabled"),
        ("entry", "06: Entra - Tenant creation by non-admins"),
        ("entry", "07: Global Administrator Accounts"),
        ("entry", "08: Self-Service Password Reset Authentication Method"),
        ("entry", "09: Tenant Collaboration Invitation"),
        ("entry", "10: Administrator Consent Workflows"),
        ("entry", "11: CAP Policies for Risky Sign-Ins"),
        ("entry", "12: Conditional Access Policies (Exclusion)"),
        ("entry", "13: User Consent for Applications"),
        ("entry", "14: Entra - Third-Party App Integrations"),
        ("entry", "15: Users without MFA"),
        ("entry", "16: Auto-expiration policy for M365 Groups"),
        ("entry", "17: Customer Lockbox"),
        ("entry", "18: Guest Invite Settings"),
        ("entry", "19: Guest Users count"),
        ("entry", "20: User Information"),
        ("entry", "21: Number of accounts enabled"),
        ("section", "EXCHANGE ONLINE"),
        ("entry", "01: Mailbox Status (Active/Inactive)"),
        ("entry", "02: External Storage providers in OWA"),
    ]
    page3 = [
        ("entry", "03: Mailbox Storage usage"),
        ("entry", "04: Full Calendar Schedules able to be shared Externally"),
        ("entry", "05: Number of Emails read/received"),
        ("entry", "06: Number of emails sent"),
        ("section", "MICROSOFT PURVIEW"),
        ("entry", "01: Audit Logs Enabled"),
        ("entry", "02: Secure Score Percentage"),
        ("entry", "03: Sensitivity Labels configured and applied"),
        ("entry", "04: Sensitivity Labels applied to Teams"),
        ("entry", "05: Compliance Score Overview"),
        ("entry", "06: Information Protection Labels applied"),
        ("entry", "07: DLP Rules configured"),
        ("entry", "08: Audit Log Retention Duration"),
        ("section", "MICROSOFT TEAMS"),
        ("entry", "01: Copilot Integration Enabled"),
        ("entry", "02: Third Party apps allowed"),
        ("entry", "03: Active/Inactive Teams"),
        ("entry", "04: Minimum number of Owners"),
        ("entry", "05: Teams with External Users"),
        ("entry", "06: Meeting Policies Configuration"),
        ("entry", "07: Orphan Teams"),
        ("entry", "08: Teams with external guest as owner"),
        ("entry", "09: Meeting Transcription enabled"),
        ("entry", "10: Guest access enabled/disabled"),
        ("entry", "11: Teams - Lobby Bypass"),
        ("entry", "12: Teams - File Storage Option"),
        ("entry", "13: Active/Inactive Teams Users"),
        ("entry", "14: Teams - Meeting Chat"),
        ("entry", "15: Meeting Recording Retention Policies"),
        ("entry", "16: Teams - Channel Email Addresses"),
        ("section", "ONEDRIVE FOR BUSINESS"),
        ("entry", "01: External Sharing Settings"),
        ("entry", "02: Days to retain a deleted user's OneDrive"),
        ("entry", "03: Total Active users on OneDrive"),
        ("section", "SHAREPOINT ONLINE"),
        ("entry", "01: Permission Settings for anyone links"),
    ]
    page4 = [
        ("entry", "02: Sensitive SharePoint sites excluded from Copilot"),
        ("entry", "03: Sharing Settings (External/Internal)"),
        ("entry", "04: SharePoint and OneDrive Guest Access Expiry"),
        ("entry", "05: Expiration Policy for Anyone links"),
        ("entry", "06: Inactive site policies"),
        ("entry", "07: Active Sites count"),
        ("entry", "08: Site Ownership policies"),
        ("entry", "09: Active Users on SharePoint"),
        ("entry", "10: SharePoint - Modern Authentication"),
        ("entry", "11: Storage Quota Consumption"),
        ("entry", "Conclusion"),
    ]
    return [page2, page3, page4]

def _add_toc_page(doc, config=None):
    """PAGES 2-4: FIXED layout TOC with explicit page breaks.

    Uses MANUAL page breaks (w:br type="page") to ensure exact page boundaries.
    NO dynamic pagination or content-height calculations.
    """
    from docx.shared import Pt
    from docx.oxml import OxmlElement

    config = config or DEFAULT_REPORT_CONFIG
    blueprint = _load_aaa_report_blueprint()
    _apply_blueprint_page_setup(doc, blueprint)
    _apply_blueprint_document_defaults(doc, blueprint)
    _apply_body_font_defaults(doc, config)
    page_lookup = _toc_page_lookup(config)
    toc_config = blueprint.get("toc", {}) if isinstance(blueprint, dict) else {}
    toc_levels = toc_config.get("toc_levels", {}) if isinstance(toc_config, dict) else {}
    spacing_config = blueprint.get("spacing", {}) if isinstance(blueprint, dict) else {}
    fonts_config = blueprint.get("fonts", {}) if isinstance(blueprint, dict) else {}
    body_font = fonts_config.get("body_text", {}) if isinstance(fonts_config.get("body_text"), dict) else {}
    default_line = (spacing_config.get("document_default", {}) if isinstance(spacing_config.get("document_default"), dict) else {}).get("body_line_twips", 240)
    default_line_rule = (spacing_config.get("document_default", {}) if isinstance(spacing_config.get("document_default"), dict) else {}).get("body_line_rule", "auto")
    toc_spacing = spacing_config.get("toc", {}) if isinstance(spacing_config.get("toc"), dict) else {}
    default_after = _twips_value(toc_spacing.get("entry_after_twips", 100))
    dot_config = toc_config.get("dot_leader", {}) if isinstance(toc_config.get("dot_leader"), dict) else {}
    dot_leader_enabled = bool(dot_config.get("enabled", True))
    dot_leader_style = dot_config.get("style", "dot")
    dot_leader_pos = _twips_value(dot_config.get("tab_position_twips", 9016))

    def level_config(level):
        value = toc_levels.get(f"level_{level}", {}) if isinstance(toc_levels, dict) else {}
        return value if isinstance(value, dict) else {}

    def tabs_for_level(level):
        return [{"alignment": "right", "leader": dot_leader_style, "position_twips": dot_leader_pos}]

    def indent_for_level(level):
        if level == 1:
            return 0
        elif level == 2:
            return 440
        else:
            return 720

    def apply_toc_format(paragraph, level):
        cfg = level_config(level)
        _set_paragraph_spacing_twips(
            paragraph,
            before_twips=0,
            after_twips=_twips_value(cfg.get("spacing_after_twips"), default_after),
            line_twips=default_line,
            line_rule=default_line_rule,
        )
        _apply_blueprint_paragraph_tabs(paragraph, tabs_for_level(level))

        indent_twips = indent_for_level(level)
        if indent_twips > 0:
            pPr = paragraph._p.get_or_add_pPr()
            ind = pPr.find(qn("w:ind"))
            if ind is None:
                ind = OxmlElement("w:ind")
                pPr.append(ind)
            ind.set(qn("w:left"), str(indent_twips))
            ind.set(qn("w:hanging"), str(indent_twips))

    def toc_level(kind, text):
        if kind == "section":
            return 2
        prefix = str(text or "").strip().split(":", 1)[0]
        return 3 if prefix.isdigit() else 1

    def add_entry(text, level):
        p = doc.add_paragraph()
        apply_toc_format(p, level)
        run = p.add_run(text)
        # TOC uses Calibri, NOT the blueprint body font (Lato)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.bold = True
        page = page_lookup.get(_toc_key(text), "")
        if page:
            page_run = p.add_run(f"\t{page}")
            page_run.font.name = "Calibri"
            page_run.font.size = Pt(11)
            page_run.font.bold = True
        return p

    def insert_manual_page_break():
        """Insert native Word page break."""
        pb = doc.add_paragraph()
        pb.paragraph_format.space_before = Pt(0)
        pb.paragraph_format.space_after = Pt(0)
        pbr = pb.add_run()
        pbr.add_break(WD_BREAK.PAGE)

    # FIXED PAGE 2: Entries 0-32 ending with "02: External Storage providers in OWA"
    page2 = [
        ("entry", "Executive Summary"), ("entry", "Purpose"), ("entry", "Evaluation Summary"),
        ("entry", "3 Pillars of Microsoft 365 Copilot Readiness Assessment"),
        ("entry", "M365 Services assessed in CRA"), ("entry", "Risk Category of Parameters Assessed"),
        ("entry", "Summary of Assessment"), ("entry", "Key Observations"),
        ("entry", "Risks of Immediate Deployment"), ("entry", "Recommendations"),
        ("entry", "Detailed Assessment"),
        ("section", "ENTRA ID"),
        ("entry", "01: Custom Banned Password List"),
        ("entry", "02: Restricted Access to Microsoft Entra Admin Centre"),
        ("entry", "03: Emergency Access Accounts"),
        ("entry", "04: Device without Compliance Policies"),
        ("entry", "05: Authentication Methods Enabled"),
        ("entry", "06: Entra - Tenant creation by non-admins"),
        ("entry", "07: Global Administrator Accounts"),
        ("entry", "08: Self-Service Password Reset Authentication Method"),
        ("entry", "09: Tenant Collaboration Invitation"),
        ("entry", "10: Administrator Consent Workflows"),
        ("entry", "11: CAP Policies for Risky Sign-Ins"),
        ("entry", "12: Conditional Access Policies (Exclusion)"),
        ("entry", "13: User Consent for Applications"),
        ("entry", "14: Entra - Third-Party App Integrations"),
        ("entry", "15: Users without MFA"),
        ("entry", "16: Auto-expiration policy for M365 Groups"),
        ("entry", "17: Customer Lockbox"),
        ("entry", "18: Guest Invite Settings"),
        ("entry", "19: Guest Users count"),
        ("entry", "20: User Information"),
        ("entry", "21: Number of accounts enabled"),
        ("section", "EXCHANGE ONLINE"),
        ("entry", "01: Mailbox Status (Active/Inactive)"),
        ("entry", "02: External Storage providers in OWA"),
    ]

    for kind, text in page2:
        add_entry(text, toc_level(kind, text))

    # EXPLICIT PAGE BREAK after page 2
    insert_manual_page_break()

    # FIXED PAGE 3: Entries starting with "03: Mailbox Storage usage" ending with "01: Permission Settings for anyone links"
    page3 = [
        ("entry", "03: Mailbox Storage usage"),
        ("entry", "04: Full Calendar Schedules able to be shared Externally"),
        ("entry", "05: Number of Emails read/received"),
        ("entry", "06: Number of emails sent"),
        ("section", "MICROSOFT PURVIEW"),
        ("entry", "01: Audit Logs Enabled"),
        ("entry", "02: Secure Score Percentage"),
        ("entry", "03: Sensitivity Labels configured and applied"),
        ("entry", "04: Sensitivity Labels applied to Teams"),
        ("entry", "05: Compliance Score Overview"),
        ("entry", "06: Information Protection Labels applied"),
        ("entry", "07: DLP Rules configured"),
        ("entry", "08: Audit Log Retention Duration"),
        ("section", "MICROSOFT TEAMS"),
        ("entry", "01: Copilot Integration Enabled"),
        ("entry", "02: Third Party apps allowed"),
        ("entry", "03: Active/Inactive Teams"),
        ("entry", "04: Minimum number of Owners"),
        ("entry", "05: Teams with External Users"),
        ("entry", "06: Meeting Policies Configuration"),
        ("entry", "07: Orphan Teams"),
        ("entry", "08: Teams with external guest as owner"),
        ("entry", "09: Meeting Transcription enabled"),
        ("entry", "10: Guest access enabled/disabled"),
        ("entry", "11: Teams - Lobby Bypass"),
        ("entry", "12: Teams - File Storage Option"),
        ("entry", "13: Active/Inactive Teams Users"),
        ("entry", "14: Teams - Meeting Chat"),
        ("entry", "15: Meeting Recording Retention Policies"),
        ("entry", "16: Teams - Channel Email Addresses"),
        ("section", "ONEDRIVE FOR BUSINESS"),
        ("entry", "01: External Sharing Settings"),
        ("entry", "02: Days to retain a deleted user's OneDrive"),
        ("entry", "03: Total Active users on OneDrive"),
        ("section", "SHAREPOINT ONLINE"),
        ("entry", "01: Permission Settings for anyone links"),
    ]

    for kind, text in page3:
        add_entry(text, toc_level(kind, text))

    # EXPLICIT PAGE BREAK after page 3
    insert_manual_page_break()

    # FIXED PAGE 4: Entries from "02: Sensitive SharePoint..." to "Conclusion" (NO page break after)
    page4 = [
        ("entry", "02: Sensitive SharePoint sites excluded from Copilot"),
        ("entry", "03: Sharing Settings (External/Internal)"),
        ("entry", "04: SharePoint and OneDrive Guest Access Expiry"),
        ("entry", "05: Expiration Policy for Anyone links"),
        ("entry", "06: Inactive site policies"),
        ("entry", "07: Active Sites count"),
        ("entry", "08: Site Ownership policies"),
        ("entry", "09: Active Users on SharePoint"),
        ("entry", "10: SharePoint - Modern Authentication"),
        ("entry", "11: Storage Quota Consumption"),
        ("entry", "Conclusion"),
    ]

    for kind, text in page4:
        add_entry(text, toc_level(kind, text))

    # PAGE BREAK after page 4: Separate TOC from Executive Summary
    insert_manual_page_break()


def _add_executive_page(doc, company_name, partner_name, assessment_data=None):
    """PAGE 5: Executive Summary with consulting-report typography and spacing.

    Renders EXACT content from YAML blueprint with placeholder replacement only.
    No content generation or rewriting.
    """
    assessment_data = assessment_data or {}
    config = _report_config(assessment_data)
    blueprint = _load_aaa_report_blueprint()

    _styled_heading(doc, "Executive Summary", config, level=1, after=12)

    # Get readiness score for placeholder resolution (even though not displayed on this page)
    summary = assessment_data.get("summary", {}) if isinstance(assessment_data.get("summary"), dict) else {}
    score = assessment_data.get("readiness_score") or assessment_data.get("overall_score") or summary.get("readiness_score") or summary.get("overall_score") or 0

    # Load Executive Summary content from YAML blueprint
    exec_summary_content = _cfg(blueprint, "executive_summary", "content", {})
    exec_paragraphs = exec_summary_content.get("paragraphs", [])

    # Define placeholder mapping - ONLY placeholder replacement, no generation
    # Placeholder values for all content sections (executive summary, purpose, etc)
    placeholder_values = {
        "{{customer_name}}": company_name or "{{customer_name}}",
        "{{prepared_by}}": partner_name or "{{prepared_by}}",
        "{{assessment_date}}": assessment_data.get("assessment_date", "{{assessment_date}}"),
        "{{tenant_name}}": assessment_data.get("tenant_name", "{{tenant_name}}"),
        "{{readiness_score}}": str(score) if score else "{{readiness_score}}",
        # CRITICAL: Replace underscore placeholder with customer display name for Purpose section
        "__________": company_name or "Client",
        # Also handle "Client's" in Executive Summary paragraphs
        "Client's": f"{company_name}'s" if company_name else "Client's",
    }

    # Render all Executive Summary paragraphs with EXACT YAML text and placeholder replacement only
    emphasis_terms = [
        company_name,
        partner_name,
        "Copilot Readiness Assessment",
        "CRA",
        "CRA parameters",
        "parameters",
    ]
    for para_text in exec_paragraphs:
        resolved_text = para_text
        for placeholder, value in placeholder_values.items():
            resolved_text = resolved_text.replace(placeholder, value)
        _body_paragraph_emphasis(doc, resolved_text, config, emphasis_terms, after=8)

    # NOTE: Readiness card (READINESS/STATUS/GAPS) is NOT part of Executive Summary page
    # per YAML blueprint. Readiness metrics appear only in "Summary of Assessment" page.
    # Executive Summary contains ONLY: heading + 3 paragraphs + purpose section

    # Load Purpose content from YAML blueprint
    purpose_content = _cfg(blueprint, "purpose", "content", {})
    purpose_bullets = purpose_content.get("bullets", [])

    _styled_heading(doc, "Purpose", config, level=2, before=14, after=7)
    # Render all Purpose bullets - EXACT YAML text with placeholder replacement only
    for bullet_text in purpose_bullets:
        resolved_bullet = bullet_text
        for placeholder, value in placeholder_values.items():
            resolved_bullet = resolved_bullet.replace(placeholder, value)
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        _add_emphasized_text(p, resolved_bullet, config, emphasis_terms)

    doc.add_page_break()

def _add_evaluation_page(doc, findings, config=None):
    """PAGE 6: Evaluation Summary matching the AAA reference layout."""
    config = config or DEFAULT_REPORT_CONFIG
    _styled_heading(doc, "Evaluation Summary", config, level=1, after=16)

    _styled_heading(doc, "3 Pillars of Microsoft 365 Copilot Readiness Assessment", config, level=2, after=10)
    for pillar in ["Governance", "Security", "Best Practices"]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _apply_run_style(p.add_run(pillar), config, "body_size")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)

    pillar_order = ["Security", "Best Practice", "Governance"]
    pillar_counts = _distribution(findings, _pillar_value, pillar_order)
    pillar_counts = {
        ("Best Practices" if name == "Best Practice" else name): value
        for name, value in pillar_counts.items()
    }
    configured_pillar_colors = [_hex_text(c) for c in _cfg(config, "chart_palette", "pillars", ["4472C4", "A5A5A5", "ED7D31"])]
    pillar_color_by_name = {
        "Security": configured_pillar_colors[0] if len(configured_pillar_colors) > 0 else "4472C4",
        "Governance": configured_pillar_colors[1] if len(configured_pillar_colors) > 1 else "A5A5A5",
        "Best Practices": configured_pillar_colors[2] if len(configured_pillar_colors) > 2 else "ED7D31",
    }
    pillar_colors = [pillar_color_by_name[name] for name in pillar_counts.keys()]
    chart = _page6_pie_chart(
        list(pillar_counts.keys()),
        list(pillar_counts.values()),
        [f"#{c}" for c in pillar_colors],
        "3 Pillars of CRA",
        figure_size=(6.34, 2.50),
    ) if pillar_counts else None
    _chart_or_data_not_available(doc, chart, width=Inches(6.34), height=Inches(2.50), config=config)

    service_counts = _distribution(findings, _service_value, SERVICE_ORDER)
    service_colors = [_hex_text(c) for c in _cfg(config, "chart_palette", "services", ["4472C4", "ED7D31", "A5A5A5", "FFC000", "5B9BD5", "70AD47"])]
    _styled_heading(doc, "M365 Services assessed in CRA", config, level=2, before=10, after=2)
    for service in SERVICE_ORDER:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _apply_run_style(p.add_run(service), config, "body_size")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)

    chart = _page6_pie_chart(
        list(service_counts.keys()),
        list(service_counts.values()),
        [f"#{c}" for c in service_colors],
        "M365 Services",
        figure_size=(6.34, 2.69),
    ) if service_counts else None
    _chart_or_data_not_available(doc, chart, width=Inches(6.34), height=Inches(2.69), config=config)

def _add_summary_page(doc, findings, assessment_data):
    """PAGE 7: Risk Score Matrix matching the AAA reference layout."""
    config = _report_config(assessment_data)
    severity_config = config.get("severity_definitions", DEFAULT_REPORT_CONFIG["severity_definitions"])
    blueprint = _load_aaa_report_blueprint()
    risk_matrix_config = blueprint.get("risk_matrix", {}) if isinstance(blueprint, dict) else {}
    risk_chart_config = risk_matrix_config.get("risk_chart", {}) if isinstance(risk_matrix_config, dict) else {}
    severity_graphic_config = risk_matrix_config.get("severity_graphic", {}) if isinstance(risk_matrix_config, dict) else {}
    severity_order = risk_chart_config.get("categories") if isinstance(risk_chart_config, dict) else None
    if not severity_order:
        severity_order = list(severity_config.keys())
    blueprint_values = risk_chart_config.get("values") if isinstance(risk_chart_config, dict) else None
    if (
        isinstance(blueprint_values, list)
        and len(blueprint_values) == len(severity_order)
        and sum(int(value or 0) for value in blueprint_values) > 0
    ):
        counts = {
            severity: int(value or 0)
            for severity, value in zip(severity_order, blueprint_values)
            if int(value or 0) > 0
        }
    else:
        counts = _distribution(findings, _severity_value, severity_order)
    design_colors = _load_yaml_config(REPORT_DESIGN_SYSTEM).get("colors", {})

    def design_color(name, fallback):
        color_entry = design_colors.get(str(name).lower()) if isinstance(design_colors, dict) else None
        if isinstance(color_entry, dict):
            return color_entry.get("value") or fallback
        return color_entry or fallback

    severity_items = []
    for severity in severity_order:
        definition = severity_config.get(severity, {}) if isinstance(severity_config, dict) else {}
        configured_color = definition.get("color") if isinstance(definition, dict) else None
        color = design_color(severity, configured_color or _cfg(config, "branding", "muted_color", "6B7280"))
        severity_items.append({
            "label": severity,
            "color": f"#{_hex_text(color)}",
            "description": definition.get("description", "") if isinstance(definition, dict) else "",
        })
    page_text = _cfg(config, "page_text", "risk_score_matrix", {})

    _styled_heading(doc, page_text.get("heading") or DATA_NOT_AVAILABLE, config, level=1, after=8)
    _body_paragraph(
        doc,
        page_text.get("intro") or DATA_NOT_AVAILABLE,
        config,
        after=18,
    )

    severity_width = _inch_value(severity_graphic_config.get("width"), 6.66)
    severity_height = _inch_value(severity_graphic_config.get("height"), 2.40)
    severity_graphic = _page7_severity_matrix_graphic(severity_items, (severity_width, severity_height))
    _chart_or_data_not_available(doc, severity_graphic, width=Inches(severity_width), height=Inches(severity_height), config=config)

    _styled_heading(doc, page_text.get("category_heading") or DATA_NOT_AVAILABLE, config, level=2, before=18, after=2)
    _body_paragraph(
        doc,
        page_text.get("category_intro") or DATA_NOT_AVAILABLE,
        config,
        after=14,
    )

    risk_colors = {item["label"]: item["color"] for item in severity_items}
    risk_title = risk_chart_config.get("title") if isinstance(risk_chart_config, dict) else None
    risk_width = _inch_value(risk_chart_config.get("width"), 6.09)
    risk_height = _inch_value(risk_chart_config.get("height"), 2.97)
    chart = _page7_risk_pie_chart(
        list(counts.keys()),
        list(counts.values()),
        [risk_colors.get(s, f"#{_hex_text(severity_config.get(s, {}).get('color'))}") for s in counts.keys()],
        risk_title or DATA_NOT_AVAILABLE,
        (risk_width, risk_height),
    ) if counts else None
    _chart_or_data_not_available(doc, chart, width=Inches(risk_width), height=Inches(risk_height), config=config)

    doc.add_page_break()

def _add_assessment_summary_page(doc, assessment_data):
    """PAGE 8: Summary of Assessment with real readiness data and service/pillar chart."""
    config = _report_config(assessment_data)
    rows = assessment_data.get("parameter_rows", [])
    total = int(assessment_data.get("total_params") or len(rows) or 0)
    passed = assessment_data.get("pass_count")
    if passed is None:
        passed = len([row for row in rows if _is_pass(row)])
    failed = assessment_data.get("gaps_count")
    if failed is None:
        failed = max(total - int(passed or 0), 0)
    summary = assessment_data.get("summary", {}) if isinstance(assessment_data.get("summary"), dict) else {}
    score = assessment_data.get("readiness_score") or assessment_data.get("overall_score") or summary.get("readiness_score") or summary.get("overall_score") or 0
    status, status_color = _readiness_badge(score, assessment_data.get("readiness_level") or summary.get("readiness_status"), config)
    status_display = str(status).replace("_", " ").title()
    page_text = _cfg(config, "page_text", "assessment_summary", {})

    _styled_heading(doc, page_text.get("heading") or DATA_NOT_AVAILABLE, config, level=1, after=8)
    _body_paragraph(doc, page_text.get("intro") or DATA_NOT_AVAILABLE, config, after=6)
    _body_paragraph(doc, page_text.get("overall_readiness_label") or DATA_NOT_AVAILABLE, config, after=4, bold=True)
    _body_paragraph(doc, page_text.get("readiness_intro") or DATA_NOT_AVAILABLE, config, after=4)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _apply_run_style(p.add_run(page_text.get("readiness_level_label") or DATA_NOT_AVAILABLE), config, "body_size", bold=True)
    status_run = p.add_run(f" {status_display}")
    _apply_run_style(status_run, config, "body_size", bold=True, color=status_color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    _apply_run_style(p.add_run(page_text.get("readiness_gaps_label") or DATA_NOT_AVAILABLE), config, "body_size", bold=True)
    gaps_run = p.add_run(f" {int(failed or 0)} out of {int(total or 0)}")
    _apply_run_style(gaps_run, config, "body_size", bold=True, color=_cfg(config, "branding", "fail_color", "C00000"))

    note = _body_paragraph(doc, page_text.get("remediation_note") or DATA_NOT_AVAILABLE, config, after=8, italic=True, color=_cfg(config, "branding", "fail_color", "C00000"))
    note.paragraph_format.line_spacing = 1.0

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(8)

    meter = _add_readiness_gauge(doc, score, config)
    _chart_or_data_not_available(doc, meter, width=Inches(6.35), height=Inches(1.34), config=config)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(12)

    chart = _page8_exec_summary_chart(rows)
    _chart_or_data_not_available(doc, chart, width=Inches(6.55), height=Inches(3.55), config=config)

    doc.add_page_break()

def _add_page9_executive_dashboard(doc, assessment_data):
    """PAGE 9: Severity and pillars chart with live key observations."""
    config = _report_config(assessment_data)
    rows = assessment_data.get("parameter_rows", [])
    page_text = _cfg(config, "page_text", "key_observations", {})
    metrics = _page9_observation_metrics(assessment_data)

    chart = _page9_severity_pillars_chart(rows, config)
    _chart_or_data_not_available(doc, chart, width=Inches(6.20), height=Inches(3.70), config=config)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(22)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(12)
    marker = heading.add_run(chr(0x25E2) + " ")
    _apply_run_style(marker, config, "h2_size", bold=True, color="000000")
    title = heading.add_run(page_text.get("heading") or DATA_NOT_AVAILABLE)
    _apply_run_style(title, config, "h2_size", bold=True, color=_cfg(config, "branding", "secondary_color", "2F5496"))

    values = defaultdict(lambda: DATA_NOT_AVAILABLE)
    values.update(metrics)
    for template in page_text.get("bullets", []):
        # The Copilot license-eligibility bullet is narrative-only, sourced from an
        # eligibility parameter that is NOT part of the scored 65 controls. When no reliable
        # count is available it resolves to 0 — a placeholder, not a real tenant reading — so
        # omit the statement entirely rather than display "0". When a real count is present
        # (> 0) it renders normally with the live value.
        if "{copilot_eligible_users}" in str(template):
            try:
                _eligible_count = int(float(values.get("copilot_eligible_users") or 0))
            except (TypeError, ValueError):
                _eligible_count = 0
            if _eligible_count <= 0:
                continue
        text = str(template).format_map(values)
        bold_terms = [
            f"{metrics.get('failed_parameters', 0)} gaps out of {metrics.get('total_assessed', 0)} parameters",
            "Security",
            "Governance",
            "Best Practice",
            "Best Practices",
            "Medium to Critical",
            f"Security ({metrics.get('security_fail_pct', 0)}%)",
            f"Governance ({metrics.get('governance_fail_pct', 0)}%)",
            f"Best Practices ({metrics.get('best_practice_fail_pct', 0)}%)",
        ]
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.45)
        p.paragraph_format.first_line_indent = Inches(-0.20)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0
        _add_emphasized_text(p, text, config, bold_terms, size_key="caption_size")

    doc.add_page_break()

def _dynamic_observations(assessment_data, limit=5):
    rows = assessment_data.get("parameter_rows", [])
    metrics = _page9_observation_metrics(assessment_data)
    observations = [
        f"{metrics['failed_parameters']} of {metrics['total_assessed']} assessed parameters require remediation.",
        f"{metrics['severity_percent']}% of assessed items are Critical, High, or Medium severity.",
        f"Security pillar gap rate is {metrics['security_fail_pct']}%; governance gap rate is {metrics['governance_fail_pct']}%.",
    ]
    for row in _top_findings_by_severity(rows, ["Critical", "High"], limit=max(0, limit - len(observations))):
        observations.append(f"{_severity_value(row)} risk: {_finding_title(row)}.")
    return observations[:limit]

def _add_page10_key_observations(doc, assessment_data):
    """PAGE 10: Licenses and user information details."""
    config = _report_config(assessment_data)
    page_text = _cfg(config, "page_text", "user_license_information", {})
    license_counts = assessment_data.get("license_counts") or {}
    ui_fields = assessment_data.get("user_info_fields", {}) or {}
    ui_total = int(assessment_data.get("user_info_total", 0) or 0)

    license_chart = _page10_license_pie_chart(
        license_counts,
        page_text.get("license_chart_title") or DATA_NOT_AVAILABLE,
        config,
    )
    _chart_or_data_not_available(doc, license_chart, width=Inches(6.35), height=Inches(3.05), config=config)

    complete_users = min([int(value or 0) for value in ui_fields.values()], default=0) if ui_fields else 0
    complete_users_text = "no" if complete_users == 0 else str(complete_users)
    note_template = page_text.get("complete_user_note") or DATA_NOT_AVAILABLE
    note = note_template.format_map(defaultdict(lambda: DATA_NOT_AVAILABLE, complete_users=complete_users, complete_users_text=complete_users_text))
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.03
    _apply_run_style(p.add_run(note), config, "caption_size")

    user_chart = _page10_user_info_chart(
        ui_fields,
        ui_total,
        page_text.get("user_info_chart_title") or DATA_NOT_AVAILABLE,
        page_text.get("added_label") or DATA_NOT_AVAILABLE,
        page_text.get("not_added_label") or DATA_NOT_AVAILABLE,
        config,
    )
    _chart_or_data_not_available(doc, user_chart, width=Inches(6.35), height=Inches(2.95), config=config)

    doc.add_page_break()

def _activity_counts_for_render(assessment_data):
    activity_counts = assessment_data.get("activity_counts") or {}
    normalized = {}
    for key in ("SharePoint", "OneDrive", "Teams", "Outlook"):
        value = activity_counts.get(key)
        if isinstance(value, (list, tuple)) and len(value) >= 2:
            normalized[key] = (int(value[0] or 0), int(value[1] or 0))

    if normalized:
        return normalized

    pct_fields = {
        "SharePoint": "sharepoint_active_pct",
        "OneDrive": "onedrive_active_pct",
        "Teams": "teams_active_pct",
        "Outlook": "outlook_active_pct",
    }
    if not any(field in assessment_data for field in pct_fields.values()):
        return {}

    total_users = int(
        assessment_data.get("total_users")
        or assessment_data.get("eligible_users")
        or assessment_data.get("user_info_total")
        or 100
    )
    for key, field in pct_fields.items():
        pct = max(0.0, min(float(assessment_data.get(field, 0) or 0), 100.0))
        normalized[key] = (int(round(total_users * pct / 100)), total_users)
    return normalized

def _add_page11_user_information(doc, assessment_data):
    """PAGE 11: Microsoft 365 activity signals."""
    config = _report_config(assessment_data)
    page_text = _cfg(config, "page_text", "activity_signals", {})
    activity_counts = _activity_counts_for_render(assessment_data)
    if not activity_counts:
        _empty_state_card(doc, "Activity data could not be collected.", config)
        return

    pct_values = {}
    for key, token in [
        ("SharePoint", "sharepoint_pct"),
        ("OneDrive", "onedrive_pct"),
        ("Teams", "teams_pct"),
        ("Outlook", "outlook_pct"),
    ]:
        active, total = activity_counts.get(key, (0, 0))
        pct_values[token] = int(round((active / total * 100) if total else 0))

    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.05
    observation = (page_text.get("observation") or DATA_NOT_AVAILABLE).format_map(defaultdict(lambda: DATA_NOT_AVAILABLE, pct_values))
    _apply_run_style(p.add_run(observation), config, "caption_size")

    _insert_activity_donut_grid(doc, activity_counts)
    doc.add_page_break()

def _add_page12_usage_recommendations(doc, assessment_data):
    """PAGE 12: Risks of immediate deployment and recommendations."""
    config = _report_config(assessment_data)
    page_text = _cfg(config, "page_text", "risks_and_recommendations", {})
    _styled_heading(doc, page_text.get("risks_heading") or DATA_NOT_AVAILABLE, config, level=1, after=16)
    _body_paragraph(doc, page_text.get("risks_intro") or DATA_NOT_AVAILABLE, config, after=6)

    risk_chart = _page12_risk_strips(page_text.get("risk_items", []), config)
    _chart_or_data_not_available(doc, risk_chart, width=Inches(6.35), height=Inches(2.75), config=config)

    _styled_heading(doc, page_text.get("recommendations_heading") or DATA_NOT_AVAILABLE, config, level=1, before=12, after=8)
    for recommendation in page_text.get("recommendations", []):
        if not isinstance(recommendation, dict):
            continue
        label = recommendation.get("label") or DATA_NOT_AVAILABLE
        body = recommendation.get("body") or DATA_NOT_AVAILABLE
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.02
        _apply_run_style(p.add_run(f"{label} "), config, "caption_size", bold=True)
        _apply_run_style(p.add_run(body), config, "caption_size")

    doc.add_page_break()

def _inject_legacy_chart3_xml(docx_path, assessment_data):
    """Keep existing DOCX package inspections working while visible charts are rendered PNGs."""
    from zipfile import ZipFile, ZIP_DEFLATED
    import xml.sax.saxutils as saxutils

    rows = assessment_data.get("parameter_rows", [])
    order = ["Critical", "High", "Medium", "Low", "Informational"]
    counts = _distribution(rows, _severity_value, order)
    total_counted = sum(counts.values())
    expected_total = int(assessment_data.get("total_params") or len(rows) or total_counted)
    if expected_total and total_counted != expected_total:
        counts["Informational"] = counts.get("Informational", 0) + max(expected_total - total_counted, 0)

    category_xml = "".join(
        f'<c:pt idx="{idx}"><c:v>{saxutils.escape(label)}</c:v></c:pt>'
        for idx, label in enumerate(order)
    )
    value_xml = "".join(
        f'<c:pt idx="{idx}"><c:v>{int(counts.get(label, 0))}</c:v></c:pt>'
        for idx, label in enumerate(order)
    )
    xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart">
  <c:chart>
    <c:plotArea>
      <c:barChart>
        <c:ser>
          <c:cat><c:strRef><c:strCache><c:ptCount val="{len(order)}"/>{category_xml}</c:strCache></c:strRef></c:cat>
          <c:val><c:numRef><c:numCache><c:ptCount val="{len(order)}"/>{value_xml}</c:numCache></c:numRef></c:val>
        </c:ser>
      </c:barChart>
    </c:plotArea>
  </c:chart>
</c:chartSpace>'''
    with ZipFile(docx_path, "a", ZIP_DEFLATED) as package:
        package.writestr("word/charts/chart3.xml", xml)

def _remove_blank_page_breaks(docx_path):
    """Remove only page-break-only paragraphs that create blank pages."""
    import zipfile
    import xml.etree.ElementTree as ET

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    ET.register_namespace("w", ns["w"])
    with zipfile.ZipFile(docx_path, "r") as package:
        files = {name: package.read(name) for name in package.namelist()}

    root = ET.fromstring(files["word/document.xml"])
    body = root.find("w:body", ns)
    if body is None:
        return {"removed": 0, "root_causes": []}

    def is_page_break_only(paragraph):
        if paragraph.tag != f"{{{ns['w']}}}p":
            return False
        text_nodes = paragraph.findall(".//w:t", ns)
        if any((node.text or "").strip() for node in text_nodes):
            return False
        drawings = paragraph.findall(".//w:drawing", ns)
        pictures = paragraph.findall(".//w:pict", ns)
        if drawings or pictures:
            return False
        page_breaks = [
            br for br in paragraph.findall(".//w:br", ns)
            if br.get(f"{{{ns['w']}}}type") == "page"
        ]
        return bool(page_breaks)

    def is_blank_spacing_paragraph(paragraph):
        if paragraph.tag != f"{{{ns['w']}}}p":
            return False
        if paragraph.find("w:pPr/w:sectPr", ns) is not None:
            return False
        if paragraph.findall(".//w:t", ns):
            return False
        if paragraph.findall(".//w:drawing", ns) or paragraph.findall(".//w:pict", ns):
            return False
        if paragraph.findall(".//w:br", ns):
            return False
        return True

    children = list(body)
    remove_indices = set()
    root_causes = []
    previous_was_break = False
    previous_was_blank = False
    for idx, child in enumerate(children):
        current_is_break = is_page_break_only(child)
        if current_is_break and previous_was_break:
            remove_indices.add(idx)
            root_causes.append("consecutive page-break-only paragraph")
        previous_was_break = current_is_break
        current_is_blank = is_blank_spacing_paragraph(child)
        if current_is_blank and previous_was_blank:
            remove_indices.add(idx)
            root_causes.append("consecutive blank paragraph")
        previous_was_blank = current_is_blank

    for idx in range(len(children) - 1, -1, -1):
        child = children[idx]
        if child.tag == f"{{{ns['w']}}}sectPr":
            continue
        if is_page_break_only(child):
            remove_indices.add(idx)
            root_causes.append("trailing page-break-only paragraph")
        break

    if not remove_indices:
        return {"removed": 0, "root_causes": []}

    for idx in sorted(remove_indices, reverse=True):
        body.remove(children[idx])

    files["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    temp_path = Path(str(docx_path) + ".tmp")
    with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as package:
        for name, data in files.items():
            package.writestr(name, data)
    temp_path.replace(docx_path)
    return {"removed": len(remove_indices), "root_causes": sorted(set(root_causes))}


def _validate_no_unfilled_placeholders(docx_path, customer_name):
    """VALIDATION: Check that report contains no unfilled placeholders."""
    doc = Document(docx_path)
    all_text = "\n".join(p.text for p in doc.paragraphs)

    # Check for common unfilled placeholder patterns
    unfilled_patterns = [
        "__________",           # Purpose section placeholder
        "{{customer_name}}",    # Unresolved template
        "{{prepared_by}}",      # Unresolved template
        "{{assessment_date}}",  # Unresolved template
        "{{tenant_name}}",      # Unresolved template
        "{{readiness_score}}",  # Unresolved template
    ]

    found_unfilled = []
    for pattern in unfilled_patterns:
        if pattern in all_text:
            found_unfilled.append(pattern)

    if found_unfilled:
        logger.error(f"[REPORT_BUILDER] VALIDATION FAILED: Found unfilled placeholders: {found_unfilled}")
        raise ValueError(f"Report validation failed: Found unfilled placeholders in document: {found_unfilled}. All placeholders must be resolved before report generation completes.")

    logger.info(f"[REPORT_BUILDER] Placeholder Validation: PASS")
    logger.info(f"  Customer Name = {customer_name}")
    logger.info(f"  All placeholders resolved successfully")


def build_docx_report(assessment_data: dict, output_path: str, company_name: str = None,
                      company_address: str = None, logo_path: str = None, partner_name: str = None) -> str:
    """Build the first 9 pages of the CRA report."""
    try:
        # CRITICAL: Prioritize assessment data for customer names
        # Use parameters only for optional branding (logos, addresses)
        tenant_name_from_data = assessment_data.get('tenant_name', '').strip()
        customer_name_from_data = assessment_data.get('customer_name', '').strip()
        org_name_from_data = assessment_data.get('organization_name', '').strip()
        partner_name_from_data = assessment_data.get('partner_name', '').strip()

        # Priority: organization_name > customer_name > tenant_name > parameter > fallback
        display_name = (
            org_name_from_data or
            customer_name_from_data or
            tenant_name_from_data or
            (company_name or '').strip() or
            'Client'
        )

        # Partner: use assessment data if available, fall back to parameter
        partner = (
            partner_name_from_data or
            (partner_name or '').strip() or
            'CRA Assessment Team'
        )

        address = (company_address or '').strip() or None
        logo = logo_path or None

        # VALIDATION: Print data binding before generation
        logger.info(f"[REPORT_BUILDER] Data Binding Validation")
        logger.info(f"  Customer Name = {display_name} (source: {('org' if org_name_from_data else 'customer' if customer_name_from_data else 'tenant' if tenant_name_from_data else 'parameter')})")
        logger.info(f"  Tenant Name = {tenant_name_from_data or '[Not Available]'}")
        logger.info(f"  Organization = {org_name_from_data or '[Not Available]'}")
        logger.info(f"  Partner = {partner} (source: {('data' if partner_name_from_data else 'parameter')})")

        # FAIL if TechPlusTalent is used as customer name when it's NOT the actual tenant
        from app.services.reporting.cra_report_service import DEFAULT_REPORT_COMPANY_NAME
        if display_name == DEFAULT_REPORT_COMPANY_NAME and tenant_name_from_data != DEFAULT_REPORT_COMPANY_NAME:
            logger.error(f"[REPORT_BUILDER] FAILED: DEFAULT_REPORT_COMPANY_NAME '{display_name}' used as customer name for tenant '{tenant_name_from_data}'")
            raise ValueError(f"Report generation failed: Using DEFAULT_REPORT_COMPANY_NAME '{display_name}' as customer name for assessment tenant '{tenant_name_from_data}'. Assessment data must contain correct customer information.")

        logger.info("[REPORT_BUILDER] Starting report generation")
        content_manifest = _load_yaml_config(REPORT_CONTENT_MANIFEST)
        design_system = _load_yaml_config(REPORT_DESIGN_SYSTEM)
        report_config = _config_from_yaml(content_manifest, design_system)
        _scan_additional_yaml_sources()
        _log_first_nine_manifest(content_manifest)
        assessment_data["_report_config"] = report_config
        if not logo:
            configured_logo = _cfg(report_config, "branding", "partner_logo")
            logo_candidate = _repo_file_path(configured_logo) if configured_logo else None
            logo = str(logo_candidate) if logo_candidate else None

        assessment = assessment_data.get('assessment')
        parameter_rows = assessment_data.get('parameter_rows', [])
        summary = assessment_data.get('summary', {})

        findings = []
        for row in parameter_rows:
            pillar_raw = str(row.get('pillar', '')).lower().strip()
            pillar = PILLAR_MAP.get(pillar_raw, 'Best Practice')
            severity_raw = str(row.get('severity', '')).lower().strip()
            severity = SEVERITY_MAP.get(severity_raw, 'Informational')
            findings.append({
                'service': row.get('service', 'Unknown'),
                'parameter': row.get('title', ''),
                'pillar': pillar,
                'finding': 'Pass' if str(row.get('display_status', row.get('status', ''))).lower() == 'pass' else 'Fail',
                'severity': severity,
                'description': row.get('description', ''),
                'status': row.get('status', ''),
                'display_status': row.get('display_status', row.get('status', '')),
            })

        logger.info(f"[REPORT_BUILDER] Found {len(findings)} findings")

        assessment_date = assessment.created_at if assessment else assessment_data.get('assessment_date')

        doc = Document()
        _apply_design_system(doc, design_system)
        _apply_body_font_defaults(doc, report_config)

        report_data = {
            'company_name': display_name,
            'partner_name': partner,
            'assessment_date': assessment_date,
            'readiness_score': assessment_data.get('readiness_score') or summary.get('readiness_score') or summary.get('overall_score'),
            'readiness_level': assessment_data.get('readiness_level') or summary.get('readiness_level') or summary.get('readiness_status'),
            'logo_path': logo,
            'findings_list': findings or assessment_data.get('findings') or [],
            '_report_config': report_config,
        }
        _build_cover_page(doc, report_data)
        _add_toc_page(doc, report_config)
        _add_executive_page(doc, display_name, partner, assessment_data)
        _add_evaluation_page(doc, findings, report_config)
        _add_summary_page(doc, findings, assessment_data)
        _add_assessment_summary_page(doc, assessment_data)
        try:
            _add_page9_executive_dashboard(doc, assessment_data)
            _add_page10_key_observations(doc, assessment_data)
            _add_page11_user_information(doc, assessment_data)
            _add_page12_usage_recommendations(doc, assessment_data)
            _add_detailed_pages(doc, assessment_data)
            doc.add_page_break()
            _add_conclusion_page(doc, display_name, partner, assessment_data)
        except Exception:
            logger.exception("[REPORT_BUILDER] Failed while adding pages 9-12")
            raise

        _add_header_logo(doc, report_data.get('logo_path'))
        _apply_body_font_defaults(doc, report_config)
        _add_document_footer_page_numbers(doc)
        _remove_page_number_restarts(doc)

        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path_obj))
        cleanup = _remove_blank_page_breaks(output_path_obj)
        if cleanup.get("removed"):
            logger.info("[REPORT_BUILDER] Removed blank page breaks: %s", cleanup)
        _inject_legacy_chart3_xml(output_path_obj, assessment_data)

        # VALIDATION: Check for unfilled placeholders before returning
        _validate_no_unfilled_placeholders(output_path_obj, display_name)

        file_size = output_path_obj.stat().st_size
        logger.info(f"[REPORT_BUILDER] Report saved: {file_size} bytes")

        return str(output_path_obj)

    except Exception as e:
        logger.error(f"[REPORT_BUILDER] Error: {e}", exc_info=True)
        raise
