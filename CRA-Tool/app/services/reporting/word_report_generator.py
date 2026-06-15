"""
CRA Word Report Generator
Uses docxtpl to fill cra_template.docx with real assessment data.
"""
import io
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import Patch
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Inches

logger = logging.getLogger(__name__)

TEMPLATE_PATH = Path(__file__).parent / 'templates' / 'cra_template.docx'

SEVERITY_ORDER = {
    'critical': 1, 'high': 2, 'medium': 3,
    'low': 4, 'info': 5, 'informational': 5
}

SEVERITY_DISPLAY = {
    'critical': 'Critical',
    'high': 'High',
    'medium': 'Medium',
    'low': 'Low',
    'info': 'Informational',
    'informational': 'Informational',
}


def _format_date(raw) -> str:
    try:
        s = str(raw)
        dt = (datetime.fromisoformat(s.split('.')[0])
              if 'T' in s else datetime.strptime(s, '%Y-%m-%d'))
        return dt.strftime('%B %d, %Y')
    except Exception:
        return datetime.utcnow().strftime('%B %d, %Y')


def _sort_by_severity(findings: list) -> list:
    return sorted(
        findings,
        key=lambda f: SEVERITY_ORDER.get(
            str(f.get('severity') or 'info').lower(), 5)
    )


def _build_params(findings: list) -> list:
    result = []
    for i, f in enumerate(_sort_by_severity(findings), 1):
        sev = str(f.get('severity') or 'info').lower()
        status = str(f.get('status') or 'fail').lower()
        finding_text = 'Pass' if status == 'pass' else 'Fail'
        result.append({
            'seq':         f'{i:02d}',
            'title':       (f.get('title') or
                            f.get('parameter_title') or ''),
            'pillar':      (f.get('pillar') or
                            f.get('category') or 'Best Practice'),
            'finding':     finding_text,
            'severity':    SEVERITY_DISPLAY.get(sev, sev.capitalize()),
            'description': (f.get('description') or
                            f.get('finding_description') or
                            f.get('evaluated_value') or ''),
            'risk':        (f.get('risk_description') or
                            f.get('parameter_risk_description') or ''),
            'doc_url':     f.get('documentation_url') or '',
            'rating_line': f"{SEVERITY_DISPLAY.get(sev, sev.capitalize())} - {finding_text}",
        })
    return result


def _calc_pillar_pct(findings_by_service: dict) -> dict:
    """Calculate pillar failure percentages correctly."""
    pillar_counts = {'Security': 0, 'Governance': 0, 'Best Practice': 0}
    total_fail = 0

    for service_findings in findings_by_service.values():
        for f in service_findings:
            if str(f.get('status') or '').lower() == 'fail':
                total_fail += 1
                pillar = str(f.get('pillar') or
                             f.get('category') or '').strip()
                if pillar in pillar_counts:
                    pillar_counts[pillar] += 1

    denom = total_fail or 1
    return {
        'security_pct':     str(round(pillar_counts['Security'] / denom * 100)),
        'governance_pct':   str(round(pillar_counts['Governance'] / denom * 100)),
        'bestpractice_pct': str(round(pillar_counts['Best Practice'] / denom * 100)),
    }


def _get_activity_stats(findings_by_service: dict,
                        report_data: dict) -> dict:
    """Get activity stats from report_data or derive from findings."""
    act = report_data.get('activity_stats') or {}
    meta = report_data.get('metadata') or {}
    usage = report_data.get('usage_stats') or {}

    return {
        'onedrive_active_pct':   str(int(
            act.get('onedrive') or
            usage.get('onedrive_active_pct') or
            meta.get('onedrive_active_pct') or 0)),
        'teams_active_pct':      str(int(
            act.get('teams') or
            usage.get('teams_active_pct') or
            meta.get('teams_active_pct') or 0)),
        'outlook_active_pct':    str(int(
            act.get('outlook') or
            usage.get('outlook_active_pct') or
            meta.get('outlook_active_pct') or 0)),
        'sharepoint_active_pct': str(int(
            act.get('sharepoint') or
            usage.get('sharepoint_active_pct') or
            meta.get('sharepoint_active_pct') or 0)),
    }


def _get_user_stats(report_data: dict) -> dict:
    """Get user eligibility stats from multiple possible sources."""
    meta = report_data.get('metadata') or {}
    summary = report_data.get('summary') or {}

    eligible = (
        report_data.get('eligible_users') or
        meta.get('eligible_users') or
        summary.get('eligible_users') or
        report_data.get('copilot_eligible_users') or 0
    )
    total = (
        report_data.get('total_users') or
        meta.get('total_users') or
        summary.get('total_users') or
        report_data.get('licensed_users') or 0
    )
    return {
        'eligible_users': str(int(eligible)),
        'total_users':    str(int(total)),
    }


def _get_sp_stats(report_data: dict) -> dict:
    """Get SharePoint stats."""
    sp = report_data.get('sharepoint_stats') or {}
    meta = report_data.get('metadata') or {}
    return {
        'sp_active_sites':  str(sp.get('active_sites') or
                                meta.get('sp_active_sites') or 0),
        'sp_total_sites':   str(sp.get('total_sites') or
                                meta.get('sp_total_sites') or 0),
        'sp_active_users':  str(sp.get('active_users') or
                                meta.get('sp_active_users') or 0),
        'sp_total_users':   str(sp.get('total_users') or
                                meta.get('sp_total_users') or 0),
        'sp_storage_used':  str(sp.get('storage_used_gb') or
                                meta.get('sp_storage_used') or '0'),
        'sp_storage_total': str(sp.get('storage_total_tb') or
                                meta.get('sp_storage_total') or '1'),
    }


def _generate_readiness_gauge(score: float, tpl) -> Optional[InlineImage]:
    """Generate horizontal gauge chart showing readiness score."""
    try:
        fig, ax = plt.subplots(figsize=(8, 1.5))
        fig.patch.set_facecolor('white')

        # Background bar (fail)
        ax.barh([0], [100], color='#DC2626', height=0.5)
        # Pass bar (score width)
        ax.barh([0], [score], color='#16A34A', height=0.5)

        # Score label
        ax.text(score, 0.35, f'{score:.2f}%',
                ha='center', fontsize=14, fontweight='bold', color='#2563EB')

        ax.set_xlim(0, 100)
        ax.set_xticks([0, 10, 20, 30, 40, 50, 60, 70, 80, 90, 100])
        ax.set_xticklabels(['0%','10%','20%','30%','40%','50%','60%','70%','80%','90%','100%'])
        ax.set_yticks([])
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_visible(False)

        legend_elements = [
            Patch(facecolor='#16A34A', label='Pass'),
            Patch(facecolor='#DC2626', label='Fail'),
        ]
        ax.legend(handles=legend_elements, loc='lower right', fontsize=9, framealpha=0)

        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        buf.seek(0)
        plt.close(fig)
        return InlineImage(tpl, buf, width=Inches(6.0))
    except Exception as e:
        logger.warning(f"Gauge chart failed: {e}")
        return None


def _generate_m365_services_chart(findings_by_service: dict, tpl) -> Optional[InlineImage]:
    """Generate stacked bar chart: M365 Services."""
    try:
        PILLAR_ORDER = ['Best Practice', 'Governance', 'Security']
        pillar_data = {p: {} for p in PILLAR_ORDER}

        for svc_name, findings in findings_by_service.items():
            for f in findings:
                pillar = f.get('pillar') or f.get('category') or 'Best Practice'
                status = str(f.get('status') or 'fail').lower()

                if pillar not in pillar_data:
                    pillar_data[pillar] = {}
                if svc_name not in pillar_data[pillar]:
                    pillar_data[pillar][svc_name] = {'pass': 0, 'fail': 0}

                pillar_data[pillar][svc_name][status] = pillar_data[pillar][svc_name].get(status, 0) + 1

        groups = []
        labels = []
        all_pass = []
        all_fail = []

        for pillar in PILLAR_ORDER:
            services = pillar_data.get(pillar, {})
            for svc, counts in sorted(services.items()):
                groups.append(pillar)
                labels.append(svc.replace('Microsoft ', 'M. ').replace(' for Business', ''))
                all_pass.append(counts.get('pass', 0))
                all_fail.append(counts.get('fail', 0))

        if not labels:
            labels = ['No Data']
            all_pass = [0]
            all_fail = [0]

        x = np.arange(len(labels))
        fig, ax = plt.subplots(figsize=(10, 5))
        fig.patch.set_facecolor('white')

        ax.bar(x, all_fail, 0.6, label='Fail', color='#DC2626')
        ax.bar(x, all_pass, 0.6, bottom=all_fail, label='Pass', color='#16A34A')

        ax.set_title('Executive Summary - M365 Services and 3 Pillars',
                     fontsize=12, fontweight='bold', pad=10)
        ax.set_xticks(x)
        ax.set_xticklabels(labels, rotation=45, ha='right', fontsize=8)
        ax.legend(loc='upper right')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        buf.seek(0)
        plt.close(fig)
        return InlineImage(tpl, buf, width=Inches(6.5))
    except Exception as e:
        logger.warning(f"M365 services chart failed: {e}")
        return None


def _generate_severity_pillars_chart(findings_by_service: dict, tpl) -> Optional[InlineImage]:
    """Generate stacked bar chart: Severity and Pillars."""
    try:
        SEVERITY_COLORS = {
            'Critical': '#DC2626', 'High': '#EA580C', 'Medium': '#D97706',
            'Low': '#65A30D', 'Informational': '#16A34A',
        }
        SEVERITY_ORDER = ['Critical', 'High', 'Medium', 'Low', 'Informational']
        PILLAR_ORDER = ['Best Practice', 'Governance', 'Security']
        STATUS_ORDER = ['Fail', 'Pass']

        data = {}
        for svc_findings in findings_by_service.values():
            for f in svc_findings:
                pillar = f.get('pillar') or f.get('category') or 'Best Practice'
                status = 'Pass' if str(f.get('status','')).lower() == 'pass' else 'Fail'
                severity = f.get('severity') or 'info'
                sev_display = {
                    'critical': 'Critical', 'high': 'High', 'medium': 'Medium',
                    'low': 'Low', 'info': 'Informational', 'informational': 'Informational'
                }.get(severity.lower(), 'Informational')

                key = (pillar, status)
                if key not in data:
                    data[key] = {s: 0 for s in SEVERITY_ORDER}
                data[key][sev_display] = data[key].get(sev_display, 0) + 1

        groups = []
        labels = []
        for pillar in PILLAR_ORDER:
            for status in STATUS_ORDER:
                groups.append((pillar, status))
                labels.append(status)

        x = np.arange(len(groups))
        fig, ax = plt.subplots(figsize=(9, 5))
        fig.patch.set_facecolor('white')

        bottoms = np.zeros(len(groups))
        for sev in SEVERITY_ORDER:
            values = [data.get(g, {}).get(sev, 0) for g in groups]
            ax.bar(x, values, 0.6, bottom=bottoms, label=sev, color=SEVERITY_COLORS[sev])
            bottoms += np.array(values)

        ax.set_title('Executive Summary - Severity and 3 Pillars',
                     fontsize=12, fontweight='bold', pad=10)
        ax.set_xticks(x)
        ax.set_xticklabels(labels, fontsize=9)
        ax.legend(loc='upper right', fontsize=8)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        buf.seek(0)
        plt.close(fig)
        return InlineImage(tpl, buf, width=Inches(6.0))
    except Exception as e:
        logger.warning(f"Severity chart failed: {e}")
        return None


def render_word_report(
    report_data: Dict[str, Any],
    logo_path: Optional[str] = None,
    company_name: Optional[str] = None,
    address: Optional[str] = None,
) -> io.BytesIO:

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"Template not found: {TEMPLATE_PATH}\n"
            f"Copy sample.docx to {TEMPLATE_PATH}"
        )

    tpl = DocxTemplate(str(TEMPLATE_PATH))

    summary  = report_data.get('summary') or {}
    fbs      = report_data.get('findings_by_service') or {}

    # Core metrics
    score       = float(summary.get('overall_score') or
                        report_data.get('score') or 0)
    fail_count  = int(summary.get('fail_count') or 0)
    total_count = int(summary.get('total_parameters') or 65)

    # Build per-service params
    def svc(key):
        for k, v in fbs.items():
            if key.lower() in k.lower():
                return v or []
        return []

    entra      = svc('entra')
    exchange   = svc('exchange')
    purview    = svc('purview')
    teams      = svc('teams')
    onedrive   = svc('onedrive')
    sharepoint = svc('sharepoint')

    # Calculate percentages correctly
    pillar_pcts = _calc_pillar_pct(fbs)
    activity    = _get_activity_stats(fbs, report_data)
    user_stats  = _get_user_stats(report_data)
    sp_stats    = _get_sp_stats(report_data)

    # Logo
    logo_img = None
    logger.info(f"[LOGO] render_word_report received logo_path: {logo_path}")
    logger.info(f"[LOGO] logo_path type: {type(logo_path)}")
    logger.info(f"[LOGO] logo_path is None: {logo_path is None}")
    logger.info(f"[LOGO] logo_path is empty: {logo_path == ''}")

    if logo_path and str(logo_path).strip():
        logo_path_str = str(logo_path).strip()
        logger.info(f"[LOGO] Processing logo_path: {logo_path_str}")

        logo_file = Path(logo_path_str).resolve()
        logger.info(f"[LOGO] Resolved absolute path: {logo_file}")
        logger.info(f"[LOGO] Logo file exists: {logo_file.exists()}")

        if logo_file.exists():
            file_size = logo_file.stat().st_size
            logger.info(f"[LOGO] File size: {file_size} bytes")

            try:
                logger.info(f"[LOGO] Creating InlineImage with width=2.0 inches")
                logo_img = InlineImage(tpl, logo_path_str, width=Inches(2.0))
                logger.info(f"[LOGO] ✅ Logo image created successfully")
            except Exception as e:
                logger.error(f"[LOGO] ❌ Logo load failed: {e}", exc_info=True)
        else:
            logger.error(f"[LOGO] ❌ Logo file does not exist: {logo_path_str}")
    else:
        logger.warning(f"[LOGO] ⚠️ No logo_path provided or empty string")

    # Charts
    readiness_chart = _generate_readiness_gauge(score, tpl)
    m365_chart = _generate_m365_services_chart(fbs, tpl)
    severity_chart = _generate_severity_pillars_chart(fbs, tpl)

    ctx = {
        # Identity
        'company_name':    (company_name or
                            report_data.get('tenant_name') or
                            report_data.get('company_name') or
                            'Client'),
        'partner_name':    (report_data.get('partner_name') or
                            'our team'),
        'assessment_date': _format_date(
                            report_data.get('assessment_date')),

        # Logo
        'logo_image': logo_img,

        # Charts
        'readiness_chart': readiness_chart,
        'm365_services_chart': m365_chart,
        'severity_chart': severity_chart,

        # Summary of Assessment
        'readiness_level': 'Ready' if score >= 70 else 'Not Ready',
        'readiness_score': f'{score:.2f}',
        'fail_count':      str(fail_count),
        'total_count':     str(total_count),

        # Pillar percentages
        **pillar_pcts,

        # User stats
        **user_stats,

        # Activity percentages
        **activity,

        # SharePoint stats
        **sp_stats,

        # Service findings lists
        'entra_findings':      _build_params(entra),
        'exchange_findings':   _build_params(exchange),
        'purview_findings':    _build_params(purview),
        'teams_findings':      _build_params(teams),
        'onedrive_findings':   _build_params(onedrive),
        'sharepoint_findings': _build_params(sharepoint),
    }

    tpl.render(ctx, autoescape=False)

    out = io.BytesIO()
    tpl.save(out)
    out.seek(0)

    # If logo was provided but not in template, add it to the document
    if logo_path and str(logo_path).strip():
        resolved_logo_path = Path(str(logo_path).strip()).resolve()
        if resolved_logo_path.exists():
            logger.info(f"[LOGO] Adding logo to document after template rendering...")
            logger.info(f"[LOGO] Using path: {resolved_logo_path}")
            try:
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                # Load the generated document
                doc = Document(out)

                # Insert logo at the beginning (before all content)
                # Add a paragraph with the logo at position 0
                new_para = doc.add_paragraph()
                new_para._element.getparent().insert(0, new_para._element)

                # Add logo to this new paragraph
                logo_run = new_para.add_run()
                logo_run.add_picture(str(resolved_logo_path), width=Inches(2.0))
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                logger.info(f"[LOGO] ✅ Logo added to document successfully")

                # Save back to BytesIO
                out = io.BytesIO()
                doc.save(out)
                out.seek(0)

            except Exception as e:
                logger.error(f"[LOGO] ❌ Failed to add logo after rendering: {e}", exc_info=True)
                # If adding logo fails, return the document without it
                pass
        else:
            logger.warning(f"[LOGO] ⚠️ Logo file does not exist at: {resolved_logo_path}")

    logger.info(
        f"Report generated for: {ctx['company_name']} | "
        f"Score: {score:.2f}% | Fail: {fail_count}/{total_count}"
    )
    return out
