"""
Prepare Word template by replacing hardcoded values with Jinja2 placeholders.
Handles Word's run-splitting, robust table detection, and risk rating numbering.
"""
import sys
import re
from pathlib import Path
from docx import Document

# Find source file
source_paths = [
    'app/services/reporting/templates/sample.docx',
    'app/services/reporting/templates/cra_template.docx',
    'sample.docx',
    './sample.docx',
]

source_file = None
for path in source_paths:
    if Path(path).exists():
        source_file = path
        print(f"Found source: {source_file}")
        break

if not source_file:
    print("ERROR: source docx not found!")
    sys.exit(1)

output_file = 'app/services/reporting/templates/cra_template.docx'

# Load document
doc = Document(source_file)
print(f"Loaded: {source_file}")
print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}\n")

# ============ HELPER FUNCTIONS ============

def replace_in_paragraph(para, find, replace):
    """Replace text in paragraph, handling split runs."""
    full = ''.join(r.text for r in para.runs)
    if find not in full:
        return False
    new_text = full.replace(find, replace)
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ''
    return True

def replace_in_doc(doc, find, replace):
    """Replace text everywhere in document."""
    count = 0
    # Paragraphs
    for para in doc.paragraphs:
        if replace_in_paragraph(para, find, replace):
            count += 1
    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_paragraph(para, find, replace):
                        count += 1
    return count

def set_table_cell_text(cell, text):
    """Set table cell text."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    if cell.paragraphs:
        cell.paragraphs[0].add_run(text)
    else:
        cell.add_paragraph(text)

# ============ STEP 1: FIX COVER PAGE ============

print("=" * 70)
print("STEP 1: Fix cover page")
print("=" * 70)

for i, para in enumerate(doc.paragraphs[:30]):
    if para.text.strip() == '.':
        if para.runs:
            para.runs[0].text = '{{ company_name }}'
            for r in para.runs[1:]:
                r.text = ''
        print(f"  Fixed cover page at paragraph {i}")
        break

# ============ STEP 2: CRITICAL REPLACEMENTS ============

print("\n" + "=" * 70)
print("STEP 2: Fix hardcoded values and double text")
print("=" * 70)

critical_fixes = [
    ('46.15%', '{{ readiness_score }}%'),
    ('April 20, 2026', '{{ assessment_date }}'),
    ('April 20,2026', '{{ assessment_date }}'),
    ('35 {{ fail_count }}', '{{ fail_count }}'),
    ('Ready/Not Ready', '{{ readiness_level }}'),
]

for find, replace in critical_fixes:
    count = replace_in_doc(doc, find, replace)
    if count > 0:
        print(f"  {find:30s} > {replace:30s} ({count})")

# ============ STEP 3: COMPANY/PARTNER NAMES ============

print("\n" + "=" * 70)
print("STEP 3: Company and partner names")
print("=" * 70)

name_fixes = [
    ('XYZ.', '{{ company_name }}.'),
    ('XYZfor', '{{ company_name }} for'),
    ('XYZ ', '{{ company_name }} '),
    ('xyz ', '{{ company_name }} '),
    ('xyz.', '{{ company_name }}.'),
    ('xyz,', '{{ company_name }},'),
    ('**xyz**', '{{ company_name }}'),
]

for find, replace in name_fixes:
    count = replace_in_doc(doc, find, replace)
    if count > 0:
        print(f"  {find:30s} > {replace:30s} ({count})")

# Fix partner name (after company replacements)
count = replace_in_doc(doc,
    '{{ company_name }} strongly recommends',
    '{{ partner_name }} strongly recommends')
if count > 0:
    print(f"  Company > Partner in 'strongly recommends': {count}")

# ============ STEP 4: READINESS GAPS AND PERCENTAGES ============

print("\n" + "=" * 70)
print("STEP 4: Readiness gaps and percentages")
print("=" * 70)

gaps_fixes = [
    (' out of 65', ' {{ fail_count }} out of {{ total_count }}'),
    ('out of 65', '{{ fail_count }} out of {{ total_count }}'),
    ('gaps out of 65 parameters', '{{ fail_count }} gaps out of {{ total_count }} parameters'),
    ('Security (51%)', 'Security ({{ security_pct }}%)'),
    ('Governance (26%)', 'Governance ({{ governance_pct }}%)'),
    ('Best Practices (23%)', 'Best Practices ({{ bestpractice_pct }}%)'),
    ('4 user accounts out of 14', '{{ eligible_users }} user accounts out of {{ total_users }}'),
]

for find, replace in gaps_fixes:
    count = replace_in_doc(doc, find, replace)
    if count > 0:
        print(f"  {find[:40]:40s} ({count})")

# ============ STEP 5: ACTIVITY STATS ============

print("\n" + "=" * 70)
print("STEP 5: Activity statistics")
print("=" * 70)

activity_fixes = [
    ('100% of OneDrive users, 0% of Microsoft Teams users, 100% of outlook users and 100% of SharePoint users',
     '{{ onedrive_active_pct }}% of OneDrive users, {{ teams_active_pct }}% of Microsoft Teams users, {{ outlook_active_pct }}% of outlook users and {{ sharepoint_active_pct }}% of SharePoint users'),
]

for find, replace in activity_fixes:
    count = replace_in_doc(doc, find, replace)
    if count > 0:
        print(f"  Activity stats updated: {count}")

# ============ STEP 6: SERVICE DESCRIPTIONS ============

print("\n" + "=" * 70)
print("STEP 6: Service description placeholders")
print("=" * 70)

descriptions = [
    ('4 users out of 4(100%) are active.', '{{ exchange_d01 }}'),
    ('Additional storage providers are allowed.', '{{ exchange_d02 }}'),
    ('All mailboxes are in good condition.', '{{ exchange_d03 }}'),
    ('Policy for individual sharing is set to least information shared.', '{{ exchange_d04 }}'),
    ('4 out of 4(100%) have read 70% of their mail.', '{{ exchange_d05 }}'),
    ('4 out of 4(100%) have sent more than 30 mails.', '{{ exchange_d06 }}'),
    ('Copilot integration is enabled.', '{{ teams_d01 }}'),
    ('Third-party apps are allowed.', '{{ teams_d02 }}'),
    ('No teams are inactive.', '{{ teams_d03 }}'),
    ('There are no teams that have less than 2 owners.', '{{ teams_d04 }}'),
    ('There are no teams with external users.', '{{ teams_d05 }}'),
    ('Recommended settings are configured.', '{{ teams_d06 }}'),
    ('There are no orphan teams.', '{{ teams_d07 }}'),
    ('No teams have external user assigned as owner.', '{{ teams_d08 }}'),
    ('Meeting transcription is enabled.', '{{ teams_d09 }}'),
    ('Guest access on teams is enabled.', '{{ teams_d10 }}'),
    ('Team lobby bypass set to Everyone in the Organization.', '{{ teams_d11 }}'),
    ('File Storage options are enabled: Dropbox, Box, GoogleDrive, ShareFile and Egnyte.', '{{ teams_d12 }}'),
    ('0 out of 4 team users are active (0%).', '{{ teams_d13 }}'),
    ('Meeting chats are enabled on global policy.', '{{ teams_d14 }}'),
    ('It is enabled and meeting recordings are set to automatically expire after 120 days.', '{{ teams_d15 }}'),
    ('Teams-Channel Email Address is enabled.', '{{ teams_d16 }}'),
    ('Audit logs are not currently enabled.', '{{ purview_d01 }}'),
    ('Secure score is only 66.84% which is less than the recommended industry standard (80%).', '{{ purview_d02 }}'),
    ('Sensitivity labels are not applied.', '{{ purview_d03 }}'),
    ('No sensitivity labels have been applied to Teams.', '{{ purview_d04 }}'),
    ('Compliance score is only 51% which is less than the recommended industry standard (80%).', '{{ purview_d05 }}'),
    ('Information Protection labels have not been applied.', '{{ purview_d06 }}'),
    ('Data Loss Prevention (DLP) rules are not configured.', '{{ purview_d07 }}'),
    ('No Audit log retention duration.', '{{ purview_d08 }}'),
    ('External sharing is set to only people in your organization.', '{{ onedrive_d01 }}'),
    ('OneDrive retention for deleted users is set to 3650 days.', '{{ onedrive_d02 }}'),
    ('4 out of 4 (100%) users have shown activity in the last 2 months.', '{{ onedrive_d03 }}'),
    ('Permission settings for anyone links is set to edit for both files and folders.', '{{ sharepoint_d01 }}'),
    ('No sites have been excluded from Copilot (Part of SharePoint advanced management).', '{{ sharepoint_d02 }}'),
    ('External sharing is disabled', '{{ sharepoint_d03 }}'),
    ('Expiration Policy for SharePoint and OneDrive guest access links is enabled and set for 90 days.', '{{ sharepoint_d04 }}'),
    ('Expiration policy is set for 30 days.', '{{ sharepoint_d05 }}'),
    ('No policy is set for inactive sites (Part of SharePoint advanced management).', '{{ sharepoint_d06 }}'),
    ('out of  sites are active.', '{{ sharepoint_d07 }}'),
    ('No policy is set for site ownership (Part of SharePoint advanced management).', '{{ sharepoint_d08 }}'),
    (' out of users are active.', '{{ sharepoint_d09 }}'),
    ('Apps using legacy authentication are disabled.', '{{ sharepoint_d10 }}'),
    ('Total storage consumption of the tenant is  out of TB.', '{{ sharepoint_d11 }}'),
]

count_total = 0
for find, replace in descriptions:
    count = replace_in_doc(doc, find, replace)
    count_total += count

print(f"Total description replacements: {count_total}")

# ============ STEP 7: PROCESS ALL TABLES ============

print("\n" + "=" * 70)
print("STEP 7: Update table Finding columns")
print("=" * 70)

SERVICE_TABLES = [
    {
        'name': 'EXCHANGE',
        'identify_by': ['Mailbox status', 'External Storage', 'mailbox'],
        'prefix': 'exchange',
        'count': 6,
    },
    {
        'name': 'PURVIEW',
        'identify_by': ['Audit Logs', 'Secure Score', 'Sensitivity Labels', 'DLP'],
        'prefix': 'purview',
        'count': 8,
    },
    {
        'name': 'TEAMS',
        'identify_by': ['Copilot Integration', 'Third Party Apps', 'Orphan Teams'],
        'prefix': 'teams',
        'count': 16,
    },
    {
        'name': 'ONEDRIVE',
        'identify_by': ['External Sharing Settings', 'deleted user', 'Active users on OneDrive'],
        'prefix': 'onedrive',
        'count': 3,
    },
    {
        'name': 'SHAREPOINT',
        'identify_by': ['Permission Settings for anyone', 'Sensitive SharePoint', 'Storage Quota'],
        'prefix': 'sharepoint',
        'count': 11,
    },
    {
        'name': 'ENTRA',
        'identify_by': ['Custom Banned', 'Emergency Access', 'Global Administrator'],
        'prefix': 'entra',
        'count': 21,
    },
]

for table in doc.tables:
    # Get all text from table to identify it
    table_text = ' '.join(
        cell.text for row in table.rows
        for cell in row.cells
    )

    matched_service = None
    for svc in SERVICE_TABLES:
        if any(keyword.lower() in table_text.lower()
               for keyword in svc['identify_by']):
            matched_service = svc
            break

    if not matched_service:
        continue

    prefix = matched_service['prefix']
    count = matched_service['count']
    print(f"  {matched_service['name']} table: updating {count} rows")

    # Skip header row (row 0), process data rows
    data_rows = table.rows[1:]

    for row_idx, row in enumerate(data_rows):
        if row_idx >= count:
            break

        n = row_idx + 1
        placeholder = f'{{{{ {prefix}_f{n:02d} }}}}'

        # Finding column is index 3
        if len(row.cells) >= 4:
            cell = row.cells[3]
            set_table_cell_text(cell, placeholder)

# ============ STEP 8: FIX RISK RATINGS ============

print("\n" + "=" * 70)
print("STEP 8: Fix Risk Rating placeholders")
print("=" * 70)

SERVICE_SECTIONS = {
    'entra id': ('entra', 21),
    'exchange online': ('exchange', 6),
    'microsoft purview': ('purview', 8),
    'microsoft teams': ('teams', 16),
    'onedrive for business': ('onedrive', 3),
    'sharepoint online': ('sharepoint', 11),
}

current_service = None
current_prefix = None
current_max = 0
param_counter = 0
risk_count = 0

for para in doc.paragraphs:
    text_lower = para.text.strip().lower()

    # Detect section changes
    for section_name, (prefix, max_count) in SERVICE_SECTIONS.items():
        if section_name in text_lower and len(para.text.strip()) < 30:
            current_service = section_name
            current_prefix = prefix
            current_max = max_count
            param_counter = 0
            break

    # Fix Risk Rating lines
    full_text = ''.join(r.text for r in para.runs)
    if 'Risk Rating:' in full_text and current_prefix:
        param_counter += 1
        if param_counter > current_max:
            param_counter = current_max

        new_placeholder = f'{{{{ {current_prefix}_r{param_counter:02d} }}}}'

        # Replace status after dash
        new_text = re.sub(
            r'(Risk Rating:[^-–]+-\s*[-–]?\s*)'
            r'(\{\{[^}]+\}\}|Pass|Fail|AB|BC|\s*)',
            rf'\1{new_placeholder}',
            full_text
        )

        if new_text != full_text and para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ''
            risk_count += 1

print(f"  Updated {risk_count} Risk Rating lines")

# ============ STEP 9: DELETE ARTIFACTS ============

print("\n" + "=" * 70)
print("STEP 9: Delete artifact paragraphs")
print("=" * 70)

to_delete = []
for para in doc.paragraphs:
    if para.text.strip() in ['a', '100%', '0%']:
        to_delete.append(para)

for para in to_delete:
    p = para._element
    p.getparent().remove(p)

print(f"  Deleted {len(to_delete)} artifacts")

# ============ SAVE ============

print("\n" + "=" * 70)
print("SAVING")
print("=" * 70)

doc.save(output_file)
print(f"Saved: {output_file}")
print(f"Size: {Path(output_file).stat().st_size} bytes")

print("\n" + "=" * 70)
print("SUCCESS: Template prepared!")
print("=" * 70)
print("\nNEXT STEP: Open cra_template.docx in Word and:")
print("  1. Delete the gauge chart")
print("  2. Delete the M365 Services bar chart")
print("  3. Delete the Severity bar chart")
print("  4. In place of each deleted chart, type the placeholder:")
print("     {{ readiness_chart }}")
print("     {{ m365_services_chart }}")
print("     {{ severity_chart }}")
print("  5. Save the file")
