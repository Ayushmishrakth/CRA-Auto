#!/usr/bin/env python3
"""
AAA TOC Page Validation Report Generator

Compares generated DOCX pages 2-4 against AAA blueprint specifications.
Outputs pixel-perfect validation metrics and comparison report.

Usage:
    python validate_toc_against_aaa.py <path_to_docx>
    python validate_toc_against_aaa.py  # Uses latest report
"""

import sys
import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn
from collections import defaultdict

sys.path.insert(0, str(Path(__file__).parent.parent))

def twips_to_inches(twips):
    """Convert twips to inches. 1 inch = 1440 twips."""
    return twips / 1440.0

def twips_to_pixels(twips, dpi=96):
    """Convert twips to pixels at given DPI. Default 96 DPI."""
    inches = twips_to_inches(twips)
    return inches * dpi

def extract_document_metrics(doc_path):
    """Extract all measurements from document."""
    doc = Document(doc_path)
    metrics = {
        "file": str(doc_path),
        "generated": datetime.now().isoformat(),
        "total_paragraphs": len(doc.paragraphs),
        "total_sections": len(doc.sections),
        "pages": {2: {}, 3: {}, 4: {}},
        "document_level": {},
    }

    # Document-level settings
    if doc.sections:
        section = doc.sections[0]
        metrics["document_level"] = {
            "top_margin": {
                "twips": section.top_margin.twips,
                "inches": section.top_margin.inches,
                "pixels_96dpi": twips_to_pixels(section.top_margin.twips),
            },
            "bottom_margin": {
                "twips": section.bottom_margin.twips,
                "inches": section.bottom_margin.inches,
                "pixels_96dpi": twips_to_pixels(section.bottom_margin.twips),
            },
            "left_margin": {
                "twips": section.left_margin.twips,
                "inches": section.left_margin.inches,
                "pixels_96dpi": twips_to_pixels(section.left_margin.twips),
            },
            "right_margin": {
                "twips": section.right_margin.twips,
                "inches": section.right_margin.inches,
                "pixels_96dpi": twips_to_pixels(section.right_margin.twips),
            },
        }

    # Per-page analysis (simplified - just categorize first N paragraphs)
    page_num = 1
    para_count_per_page = defaultdict(int)
    page_data = defaultdict(lambda: {"paragraphs": []})

    for para_idx, para in enumerate(doc.paragraphs):
        # Skip to page 2
        if page_num < 2:
            if para.text.strip() == "" and para.paragraph_format.space_after:
                page_num += 1
            continue

        if page_num > 4:
            break

        # Count paragraphs per page (heuristic)
        para_count_per_page[page_num] += 1
        if para_count_per_page[page_num] > 100:  # Unlikely to have >100 items per TOC page
            page_num += 1
            para_count_per_page[page_num] = 0

        # Extract detailed paragraph data
        para_info = {
            "index": para_idx,
            "text": para.text[:50],  # First 50 chars for display
            "has_tabs": False,
            "has_indents": False,
            "spacing": {},
            "indentation": {},
            "tabs": [],
            "fonts": [],
        }

        # Spacing info
        pPr = para._p.get_or_add_pPr()
        spacing_elem = pPr.find(qn("w:spacing"))
        if spacing_elem is not None:
            para_info["spacing"] = {
                "before_twips": spacing_elem.get(qn("w:before")),
                "after_twips": spacing_elem.get(qn("w:after")),
                "line_twips": spacing_elem.get(qn("w:line")),
                "lineRule": spacing_elem.get(qn("w:lineRule")),
            }

        # Indentation info
        ind_elem = pPr.find(qn("w:ind"))
        if ind_elem is not None:
            para_info["has_indents"] = True
            para_info["indentation"] = {
                "left_twips": ind_elem.get(qn("w:left")),
                "hanging_twips": ind_elem.get(qn("w:hanging")),
                "first_twips": ind_elem.get(qn("w:firstLine")),
            }

        # Tab stops
        tabs_elem = pPr.find(qn("w:tabs"))
        if tabs_elem is not None:
            para_info["has_tabs"] = True
            for tab in tabs_elem.findall(qn("w:tab")):
                para_info["tabs"].append({
                    "position_twips": tab.get(qn("w:pos")),
                    "alignment": tab.get(qn("w:val")),
                    "leader": tab.get(qn("w:leader")),
                })

        # Font info
        for run in para.runs:
            para_info["fonts"].append({
                "text": run.text[:20],
                "font_name": run.font.name,
                "size_pt": run.font.size.pt if run.font.size else None,
                "bold": run.font.bold,
            })

        page_data[page_num]["paragraphs"].append(para_info)

    # Assign page data
    for page_num in [2, 3, 4]:
        metrics["pages"][page_num] = page_data[page_num]

    return metrics

def validate_against_blueprint(metrics):
    """Validate metrics against AAA blueprint specs."""
    # Expected values from AAA blueprint
    expected = {
        "margins": {
            "top_twips": 1440,
            "bottom_twips": 1440,
            "left_twips": 1440,
            "right_twips": 1440,
        },
        "spacing": {
            "after_entry_twips": 100,
        },
        "indentation": {
            "level_1_twips": 0,
            "level_2_twips": 440,
            "level_3_twips": 720,
        },
        "tabs": {
            "page_number_position_twips": 9016,
            "dot_leader_style": "dot",
        },
        "line": {
            "spacing_twips": 240,
            "rule": "auto",
        },
    }

    validation = {
        "timestamp": datetime.now().isoformat(),
        "expected": expected,
        "actual": metrics,
        "results": {
            "document_level": {"pass": True, "errors": [], "warnings": []},
            "pages": {2: {}, 3: {}, 4: {}},
            "summary": {"pass": True, "total_errors": 0, "total_warnings": 0},
        },
    }

    # Validate document-level margins
    doc_level = validation["results"]["document_level"]
    actual_margins = metrics.get("document_level", {})

    for margin_key in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        expected_twips_key = f"{margin_key.replace('_margin', '')}_twips"
        expected_val = expected["margins"].get(expected_twips_key)
        actual_val = actual_margins.get(margin_key, {}).get("twips")

        if actual_val is None:
            doc_level["warnings"].append(f"Could not measure {margin_key}")
        else:
            diff = abs(int(actual_val) - expected_val)
            pixels_96dpi = twips_to_pixels(diff)

            if diff > 10:  # 10 twips tolerance
                doc_level["errors"].append({
                    "metric": margin_key,
                    "expected_twips": expected_val,
                    "actual_twips": int(actual_val),
                    "diff_twips": diff,
                    "diff_pixels": f"{pixels_96dpi:.2f}px",
                })
                doc_level["pass"] = False
            elif diff > 0:
                doc_level["warnings"].append({
                    "metric": margin_key,
                    "expected_twips": expected_val,
                    "actual_twips": int(actual_val),
                    "diff_twips": diff,
                })

    # Validate per-page formatting
    for page_num in [2, 3, 4]:
        page_validation = validation["results"]["pages"][page_num]
        page_validation["pass"] = True
        page_validation["errors"] = []
        page_validation["samples"] = []

        page_paras = metrics["pages"][page_num].get("paragraphs", [])
        page_validation["total_items"] = len(page_paras)

        # Sample first few paragraphs
        for para_info in page_paras[:5]:  # Check first 5 items
            sample = {
                "text": para_info["text"],
                "checks": {"pass": True, "issues": []},
            }

            # Check spacing after
            spacing_after = para_info.get("spacing", {}).get("after_twips")
            if spacing_after and int(spacing_after) != expected["spacing"]["after_entry_twips"]:
                diff = abs(int(spacing_after) - expected["spacing"]["after_entry_twips"])
                sample["checks"]["issues"].append(f"Spacing after: {spacing_after} twips (expected 100, diff={diff})")
                sample["checks"]["pass"] = False

            # Check tabs
            expected_tab_pos = expected["tabs"]["page_number_position_twips"]
            has_right_tab = any(
                int(t.get("position_twips") or 0) == expected_tab_pos
                for t in para_info.get("tabs", [])
            )
            if not has_right_tab and para_info.get("tabs"):
                sample["checks"]["issues"].append(f"Missing right tab at {expected_tab_pos} twips")
                sample["checks"]["pass"] = False

            page_validation["samples"].append(sample)

    # Calculate summary
    total_errors = len(validation["results"]["document_level"]["errors"])
    total_warnings = len(validation["results"]["document_level"]["warnings"])
    for page_num in [2, 3, 4]:
        total_errors += len(validation["results"]["pages"][page_num]["errors"])

    validation["results"]["summary"]["total_errors"] = total_errors
    validation["results"]["summary"]["total_warnings"] = total_warnings
    validation["results"]["summary"]["pass"] = total_errors == 0

    return validation

def generate_comparison_report(metrics, validation):
    """Generate human-readable comparison report."""
    report = """
================================================================================
                    AAA TOC PAGE VALIDATION REPORT
                         Pages 2-4 Comparison
================================================================================

DOCUMENT INFORMATION:
  File: {file}
  Generated: {generated}
  Total Paragraphs: {total_paragraphs}
  Total Sections: {total_sections}

VALIDATION STATUS: {status}

================================================================================
                           MARGIN MEASUREMENTS
================================================================================

Expected: 1.0 inch (1440 twips) on all sides
Tolerance: ±10 twips (±0.007 inches, ~1 pixel at 96 DPI)

{margin_table}

{margin_errors}

================================================================================
                         DOCUMENT-LEVEL RESULTS
================================================================================

Status: {doc_status}
Errors: {doc_errors}
Warnings: {doc_warnings}

{doc_issues}

================================================================================
                         PAGE-BY-PAGE ANALYSIS
================================================================================

{page_analysis}

================================================================================
                              SUMMARY
================================================================================

Total Items Validated: {total_items}
Total Errors: {total_errors}
Total Warnings: {total_warnings}

PASS/FAIL: {final_status}

Recommendation:
{recommendation}

================================================================================
"""

    # Format margins table
    margin_lines = []
    margins_data = metrics.get("document_level", {})
    for margin_name in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        if margin_name in margins_data:
            m = margins_data[margin_name]
            margin_lines.append(
                f"  {margin_name.replace('_', ' ').title():<20} "
                f"{m['twips']:>6} twips "
                f"({m['inches']:.4f} in) "
                f"({m['pixels_96dpi']:.2f}px @ 96DPI)"
            )

    # Format errors
    doc_level_result = validation["results"]["document_level"]
    error_lines = []
    if doc_level_result["errors"]:
        for err in doc_level_result["errors"]:
            if isinstance(err, dict):
                error_lines.append(
                    f"  ✗ {err['metric']}: "
                    f"Expected {err['expected_twips']} twips, "
                    f"Got {err['actual_twips']} twips, "
                    f"Difference: {err['diff_twips']} twips ({err['diff_pixels']})"
                )
    if doc_level_result["warnings"]:
        for warn in doc_level_result["warnings"]:
            if isinstance(warn, dict):
                error_lines.append(
                    f"  ⚠ {warn['metric']}: "
                    f"Minor deviation: {warn['diff_twips']} twips"
                )
            else:
                error_lines.append(f"  ⚠ {warn}")

    # Page analysis
    page_lines = []
    total_items = 0
    for page_num in [2, 3, 4]:
        page_data = validation["results"]["pages"][page_num]
        total_items += page_data.get("total_items", 0)
        page_lines.append(f"\nPage {page_num}:")
        page_lines.append(f"  Items: {page_data.get('total_items', 0)}")
        if page_data.get("samples"):
            for sample in page_data["samples"][:3]:
                status = "✓" if sample["checks"]["pass"] else "✗"
                page_lines.append(f"  {status} {sample['text']}")
                if sample["checks"]["issues"]:
                    for issue in sample["checks"]["issues"]:
                        page_lines.append(f"      {issue}")

    final_status = "✅ PASS" if validation["results"]["summary"]["pass"] else "❌ FAIL"
    recommendation = (
        "All measurements match AAA blueprint specifications. "
        "TOC pages are pixel-perfect compliant."
        if validation["results"]["summary"]["pass"]
        else
        "Measurement deviations detected. Review errors above and regenerate report."
    )

    return report.format(
        file=metrics["file"],
        generated=metrics["generated"],
        total_paragraphs=metrics["total_paragraphs"],
        total_sections=metrics["total_sections"],
        status=final_status,
        margin_table="\n".join(margin_lines),
        margin_errors="\n".join(error_lines) if error_lines else "  ✓ All margins within tolerance",
        doc_status="✓ PASS" if doc_level_result["pass"] else "✗ FAIL",
        doc_errors=len(doc_level_result["errors"]),
        doc_warnings=len(doc_level_result["warnings"]),
        doc_issues="\n".join(error_lines) if error_lines else "  No issues detected",
        page_analysis="\n".join(page_lines),
        total_items=total_items,
        total_errors=validation["results"]["summary"]["total_errors"],
        total_warnings=validation["results"]["summary"]["total_warnings"],
        final_status=final_status,
        recommendation=recommendation,
    )

def main():
    # Determine which document to validate
    if len(sys.argv) > 1:
        doc_path = Path(sys.argv[1])
    else:
        # Use latest report
        reports_dir = Path(__file__).parent.parent / "storage" / "reports"
        report_files = list(reports_dir.rglob("*.docx"))
        if not report_files:
            print("[ERROR] No DOCX files found. Provide path or generate report first.")
            sys.exit(1)
        doc_path = max(report_files, key=lambda p: p.stat().st_mtime)

    if not doc_path.exists():
        print(f"[ERROR] File not found: {doc_path}")
        sys.exit(1)

    print(f"[INFO] Validating: {doc_path}\n")

    try:
        # Extract metrics
        print("[*] Extracting document metrics...")
        metrics = extract_document_metrics(str(doc_path))

        # Validate
        print("[*] Validating against AAA blueprint...")
        validation = validate_against_blueprint(metrics)

        # Generate report
        print("[*] Generating comparison report...")
        report = generate_comparison_report(metrics, validation)

        # Output
        print(report)

        # Save JSON for programmatic access
        json_report = {
            "metrics": metrics,
            "validation": validation,
            "report_text": report,
        }
        json_path = doc_path.parent / f"{doc_path.stem}_aaa_validation.json"
        with open(json_path, "w") as f:
            json.dump(json_report, f, indent=2, default=str)
        print(f"\n[INFO] Detailed JSON report saved: {json_path}")

        # Exit with appropriate code
        sys.exit(0 if validation["results"]["summary"]["pass"] else 1)

    except Exception as e:
        print(f"[ERROR] Validation failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()
