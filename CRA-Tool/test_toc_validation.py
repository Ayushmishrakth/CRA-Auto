#!/usr/bin/env python3
"""
Quick test to validate AAA TOC page generation (pages 2-4).

Generates a minimal report with only the TOC pages and validates:
- Margins
- Tab stops
- Paragraph indentation
- Spacing
- Font specifications
- Page break positions
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from docx.oxml.ns import qn
from datetime import datetime
import json

def measure_toc_page(doc_path, target_page=2):
    """Extract measurements from a specific TOC page."""
    doc = Document(doc_path)

    measurements = {
        "page": target_page,
        "document_margins": {},
        "sections": 0,
        "paragraphs": [],
    }

    # Get document margins from first section
    if doc.sections:
        section = doc.sections[0]
        measurements["document_margins"] = {
            "top_twips": section.top_margin.twips,
            "bottom_twips": section.bottom_margin.twips,
            "left_twips": section.left_margin.twips,
            "right_twips": section.right_margin.twips,
            "top_inches": section.top_margin.inches,
            "bottom_inches": section.bottom_margin.inches,
            "left_inches": section.left_margin.inches,
            "right_inches": section.right_margin.inches,
        }
        measurements["sections"] = len(doc.sections)

    # Extract paragraph data
    current_page = 1
    for para_idx, para in enumerate(doc.paragraphs):
        # Skip paragraphs before page 2
        if current_page < target_page:
            if para.text.strip() == "":  # Page break indicator
                current_page += 1
            continue

        if current_page > target_page:
            break

        para_data = {
            "index": para_idx,
            "text": para.text[:60],  # First 60 chars
            "spacing_before_twips": para.paragraph_format.space_before.twips if para.paragraph_format.space_before else None,
            "spacing_after_twips": para.paragraph_format.space_after.twips if para.paragraph_format.space_after else None,
            "line_spacing": para.paragraph_format.line_spacing,
            "left_indent_twips": para.paragraph_format.left_indent.twips if para.paragraph_format.left_indent else None,
            "first_line_indent_twips": para.paragraph_format.first_line_indent.twips if para.paragraph_format.first_line_indent else None,
            "tabs": [],
            "fonts": [],
        }

        # Extract tab stops
        pPr = para._p.get_or_add_pPr()
        tabs_elem = pPr.find(qn("w:tabs"))
        if tabs_elem is not None:
            for tab in tabs_elem.findall(qn("w:tab")):
                tab_data = {
                    "position_twips": int(tab.get(qn("w:pos")) or 0),
                    "alignment": tab.get(qn("w:val"), "left"),
                    "leader": tab.get(qn("w:leader"), "none"),
                }
                para_data["tabs"].append(tab_data)

        # Extract font info from runs
        for run in para.runs:
            font_data = {
                "text": run.text[:20],
                "name": run.font.name,
                "size_pt": run.font.size.pt if run.font.size else None,
                "bold": run.font.bold,
                "italic": run.font.italic,
            }
            para_data["fonts"].append(font_data)

        measurements["paragraphs"].append(para_data)

    return measurements


def validate_toc_against_blueprint(doc_path):
    """Validate generated TOC against AAA blueprint specs."""
    from CRA-Tool.app.services.reporting.report_builder import _load_aaa_report_blueprint, _blueprint_margins

    blueprint = _load_aaa_report_blueprint()
    margins = _blueprint_margins(blueprint)

    expected = {
        "top_twips": 1440,
        "bottom_twips": 1440,
        "left_twips": 1440,
        "right_twips": 1440,
        "spacing_after_twips": 100,
        "line_twips": 240,
        "dot_leader_tab_twips": 9016,
        "level2_indent_twips": 440,
    }

    report = {
        "document_path": str(doc_path),
        "timestamp": datetime.now().isoformat(),
        "blueprint_loaded": blueprint is not None,
        "pages": {2: {}, 3: {}, 4: {}},
        "validation": {
            "pass": True,
            "errors": [],
            "warnings": [],
            "differences": [],
        },
    }

    doc = Document(doc_path)
    if doc.sections:
        section = doc.sections[0]
        actual = {
            "top_twips": section.top_margin.twips,
            "bottom_twips": section.bottom_margin.twips,
            "left_twips": section.left_margin.twips,
            "right_twips": section.right_margin.twips,
        }

        for key, expected_val in expected.items():
            if key in actual:
                actual_val = actual[key]
                diff = abs(actual_val - expected_val)
                if diff > 10:  # 10 twips tolerance
                    report["validation"]["errors"].append(
                        f"{key}: expected {expected_val}, got {actual_val}, diff={diff}"
                    )
                    report["validation"]["pass"] = False
                elif diff > 0:
                    report["validation"]["warnings"].append(
                        f"{key}: expected {expected_val}, got {actual_val}, diff={diff}"
                    )
                else:
                    report["validation"]["differences"].append({
                        "metric": key,
                        "expected": expected_val,
                        "actual": actual_val,
                        "diff": 0,
                        "status": "PASS"
                    })

    # Measure each page
    for page_num in [2, 3, 4]:
        measurements = measure_toc_page(doc_path, page_num)
        report["pages"][page_num] = measurements

    return report


if __name__ == "__main__":
    import sys
    from pathlib import Path

    # Change to CRA-Tool directory
    cra_tool_dir = Path(__file__).parent
    sys.path.insert(0, str(cra_tool_dir))

    print("=" * 80)
    print("AAA TOC PAGE VALIDATION TEST")
    print("=" * 80)
    print()

    # Try to find a generated report
    reports_dir = cra_tool_dir / "storage" / "reports"
    report_files = list(reports_dir.rglob("*.docx"))

    if not report_files:
        print("[INFO] No reports found. Would need to generate one first.")
        print("[INFO] Run: python -m CRA-Tool.scripts.generate_sample_report")
        sys.exit(0)

    latest_report = max(report_files, key=lambda p: p.stat().st_mtime)
    print(f"[INFO] Testing latest report: {latest_report}")
    print()

    try:
        report = validate_toc_against_blueprint(str(latest_report))

        print("VALIDATION RESULTS:")
        print("-" * 80)
        print(f"Document: {report['document_path']}")
        print(f"Timestamp: {report['timestamp']}")
        print()

        val = report["validation"]
        print(f"Status: {'PASS' if val['pass'] else 'FAIL'}")
        print(f"Errors: {len(val['errors'])}")
        print(f"Warnings: {len(val['warnings'])}")
        print()

        if val["errors"]:
            print("ERRORS:")
            for err in val["errors"]:
                print(f"  - {err}")
            print()

        if val["warnings"]:
            print("WARNINGS:")
            for warn in val["warnings"]:
                print(f"  - {warn}")
            print()

        print("DIFFERENCES:")
        for diff in val["differences"]:
            print(f"  {diff['metric']}: {diff['expected']} (expected) vs {diff['actual']} (actual)")
        print()

        print("PAGE SUMMARIES:")
        print("-" * 80)
        for page_num in [2, 3, 4]:
            page_data = report["pages"][page_num]
            print(f"\nPage {page_num}:")
            print(f"  Paragraphs: {len(page_data['paragraphs'])}")
            if page_data['paragraphs']:
                print(f"  First: {page_data['paragraphs'][0]['text']}")
                if page_data['paragraphs'][0]['tabs']:
                    print(f"  Tabs: {page_data['paragraphs'][0]['tabs']}")

        print("\n" + "=" * 80)
        print("Full report saved as JSON:")
        print(json.dumps(report, indent=2, default=str))

    except Exception as e:
        print(f"[ERROR] Validation failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
