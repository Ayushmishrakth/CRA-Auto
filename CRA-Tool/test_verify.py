from docx import Document

doc = Document('storage/test_header.docx')
header = doc.sections[0].header

print('Header paragraphs:', len(header.paragraphs))
print('Header first para runs:', len(header.paragraphs[0].runs))

if header.paragraphs[0].runs:
    for i, run in enumerate(header.paragraphs[0].runs):
        print(f'Run {i}:')
        print(f'  Text: {repr(run.text)}')
        # Check for drawing elements (images)
        has_drawing = False
        for elem in run._r:
            if 'drawing' in str(elem.tag).lower():
                has_drawing = True
                print(f'  Drawing element found!')
        if not has_drawing:
            print(f'  No drawing element')

print('\n=== RESULT ===')
if header.paragraphs[0].runs and any('drawing' in str(elem.tag).lower() for run in header.paragraphs[0].runs for elem in run._r):
    print('YES - Logo is in the header')
else:
    print('NO - Logo is NOT in the header')
